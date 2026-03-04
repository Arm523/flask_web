from flask import current_app
from datetime import date, datetime, timedelta
from docx import Document
import mysql.connector
from mysql.connector import Error
from dateutil.relativedelta import relativedelta
from num2words import num2words
from master_modbus import read_meter_unit_read
import calendar
import os
import json
from werkzeug.utils import secure_filename

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
CONFIG_DIR = os.path.join(BASE_DIR, 'config_meter')
PATH_MODEL_ELEC = os.path.join(CONFIG_DIR, 'model.json')
PATH_MODEL_WATER = os.path.join(CONFIG_DIR, 'model_water.json')

if not os.path.exists(CONFIG_DIR):
    os.makedirs(CONFIG_DIR, exist_ok=True)

def get_file_path(category):
    return PATH_MODEL_ELEC if category == 'electricity' else PATH_MODEL_WATER

def read_config(path):
    if not os.path.exists(path):
        # สร้างไฟล์เปล่าถ้ายังไม่มี
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, 'w', encoding='utf-8') as f:
            json.dump({}, f)
        return {}
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_config(path, data):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

# ดึงวันเวลาปัจจุบัน (หรือ mock)
def get_now(mocked=False):
    if mocked:
        # return datetime.now()
        return datetime(2026, 10,5, 9, 0, 0)
    else:
        return datetime.now()

today = get_now(mocked=True).date()

def generate_slip_filename(file, prefix, ref_id):
    """
    ฟังก์ชันสำหรับสร้างชื่อไฟล์ใหม่ รองรับทุกนามสกุลไฟล์
    รูปแบบ: prefix_id_DD-MM-YYYY_HHMMSS.extension
    """
    if not file or file.filename == '':
        return None
    _, ext = os.path.splitext(file.filename)
    
    ext = ext.lower()

    timestamp = datetime.now().strftime('%d-%m-%Y_%H%M%S')
    new_filename = f"{prefix}_{ref_id}_{timestamp}{ext}"
    
    # ล้างอักขระพิเศษเพื่อความปลอดภัย
    return secure_filename(new_filename)

def get_db_connection():
    try:
        conn = mysql.connector.connect(
            host='192.168.137.1',
            user='flask_user',
            password='nit.P#1234',
            database='apartment',
        )
        return conn
    except Error as e:
        print(f"DB Connection error: {e}")
        return None


def record_transaction(cursor, amount, t_type, category, ref_invoice_id=None, ref_expense_id=None, note="", created_by=None):
    """
    ฟังก์ชันกลางสำหรับบันทึกรายรับ-รายจ่ายลงตาราง transactions
    """
    try:
        sql = """
            INSERT INTO transactions (
                transaction_date, type, category, amount, 
                ref_invoice_id, ref_expense_id, note, created_by
            ) VALUES (NOW(), %s, %s, %s, %s, %s, %s, %s)
        """
        cursor.execute(sql, (
            t_type,           # 'income' หรือ 'expense'
            category,         # เช่น 'ค่าเช่า', 'ค่าซ่อม'
            amount,           # ยอดเงิน
            ref_invoice_id,   # เชื่อมโยงบิล (ถ้ามี)
            ref_expense_id,   # เชื่อมโยงใบจ่ายเงิน (ถ้ามี)
            note,             # บันทึกเพิ่มเติม
            created_by        # ID ผู้ทำรายการ
        ))
        return True
    except Exception as e:
        print(f"Error recording transaction: {e}")
        return False


def mark_contracts_expiring(today, cursor):
    """
    ฟังก์ชันรันรายวัน (เที่ยงคืน): 
    สร้างบิลสุดท้าย (Final Bill) อัตโนมัติ พร้อมดึง Options จากสัญญาลง invoice_items
    """
    print(f"\n🔍 [System Check: {today}] Starting maintenance job...")

    # --- STEP 1: ค้นหาสัญญาที่กำลังจะหมดใน 30 วัน และยังไม่มีบิล Final ---
    # แก้ไข Query ใน STEP 1 ของพี่
    cursor.execute("""
        SELECT c.*, u.unit_id 
        FROM contracts c
        JOIN unit u ON c.room_id = u.unit_id
        LEFT JOIN invoices i ON c.contract_id = i.contract_id 
             AND i.invoice_type = 'final' 
             AND i.status != 'cancelled'
        WHERE c.status IN (2, 3, 4)
          AND c.contract_end IS NOT NULL
          AND c.contract_end <= DATE_ADD(%s, INTERVAL 5 DAY) 
          AND i.invoice_id IS NULL
    """, (today,))
    
    contracts_to_bill = cursor.fetchall()

    # ดึงเรทน้ำไฟจาก Settings
    elec_rate = float(get_setting('electricity_rate', 0))
    water_rate = float(get_setting('water_rate', 0))

    for contract in contracts_to_bill:
        contract_id = contract['contract_id']
        rent_amount = contract['price']
        reimburse = contract['premiums'] or 0
        
        # 1.1 หาเลชมิเตอร์ครั้งก่อน (จากบิลล่าสุดของสัญญานี้)
        cursor.execute("""
            SELECT current_electricity_reading, current_water_reading, billing_period_end
            FROM invoices 
            WHERE contract_id = %s AND status != 'cancelled'
            ORDER BY billing_period_end DESC LIMIT 1
        """, (contract_id,))
        last_inv = cursor.fetchone()

        if last_inv and last_inv['billing_period_end']:
            new_billing_start = last_inv['billing_period_end'] + timedelta(days=1)
            prev_elec = last_inv['current_electricity_reading'] or 0
            prev_water = last_inv['current_water_reading'] or 0
        else:
            new_billing_start = contract['contract_start']
            prev_elec = contract['electricity_start'] or 0
            prev_water = contract['water_start'] or 0

        # 1.2 สร้างบิลหลัก (Draft)
        initial_total = rent_amount - reimburse

        cursor.execute("""
            INSERT INTO invoices (
                unit_id, tenant_id, contract_id, invoice_type, 
                billing_period_start, billing_period_end, 
                issue_date, due_date, rent_amount,
                electricity_rate, water_rate,
                previous_electricity_reading, previous_water_reading,
                total_amount, status, created_at, created_by, premiums, reimburse
            ) VALUES (
                %s, %s, %s, 'final', 
                %s, %s, 
                %s, %s, %s, 
                %s, %s, 
                %s, %s, 
                %s, 'draft', NOW(), NULL, 0, %s
            )
        """, (
            contract['room_id'], contract['tenant_id'], contract_id,
            new_billing_start, contract['contract_end'], # billing start/end
            today, contract['contract_end'],                   # issue/due date
            rent_amount,
            elec_rate, water_rate,
            prev_elec, prev_water,
            initial_total, reimburse                           # total และเงินประกัน
        ))
        invoice_id = cursor.lastrowid

        # 1.3 เพิ่ม Options จากสัญญาลงใน invoice_items
        cursor.execute("""
            SELECT o.* FROM `option` o
            JOIN contract_option co ON o.id = co.option_id
            WHERE co.contract_id = %s
        """, (contract_id,))
        options = cursor.fetchall()

        option_total = 0
        for opt in options:
            if opt['name'] != 'ค่าเช่าห้อง':
                cursor.execute("""
                    INSERT INTO invoice_items (invoice_id, description, unit_price, quantity, total_price, option_id)
                    VALUES (%s, %s, %s, 1, %s, %s)
                """, (invoice_id, opt['name'], opt['price'], opt['price'], opt['id']))
                option_total += float(opt['price'])

        # 1.4 อัปเดต total_amount (ยอดเดิม + ยอดจาก options)
        cursor.execute("""
            UPDATE invoices 
            SET total_amount = total_amount + %s
            WHERE invoice_id = %s
        """, (option_total, invoice_id))

        # 1.5 อัปเดตสถานะห้องเป็น 6 (ตรวจสอบห้อง เครียร์บิลทั้งหมด)
        print(f"✅ บิลสุดท้ายสร้างเสร็จ (ID: {invoice_id}) สำหรับห้อง {contract['room_id']} (รวม Options: {option_total} บาท)")

    # --- STEP 2 & 3: ปักธงสถานะสัญญาและรายวัน (คงเดิม) ---
    cursor.execute("UPDATE contracts SET status = 4 WHERE status IN (2, 3) AND contract_end <= DATE_ADD(%s, INTERVAL 30 DAY)", (today,))
    
    cursor.execute("""
        UPDATE unit u JOIN (SELECT unit_id, MAX(billing_period_end) as latest_end FROM invoices WHERE invoice_type='daily' AND status='paid' GROUP BY unit_id) i ON i.unit_id = u.unit_id
        SET u.status_id = 6 WHERE i.latest_end <= %s AND u.status_id = 2
    """, (today,))

    print("✅ เสร็จสิ้นการรันระบบบำรุงรักษาประจำวัน")


def mark_invoices_overdue(mocked_date=None):
    """
    อัปเดต status ของ invoices เป็น 'overdue'
    เฉพาะบิลที่ยังไม่จ่ายและหมดช่วง billing_period_end
    """
    today = mocked_date or date.today()
    conn = get_db_connection()
    if not conn:
        print("❌ เชื่อมต่อ DB ไม่สำเร็จ")
        return

    cursor = conn.cursor(dictionary=True)
    try:
        # อัปเดตบิลที่หมดช่วงแล้ว แต่ยังไม่จ่าย
        cursor.execute("""
            UPDATE invoices
            SET status = 'overdue'
            WHERE billing_period_end < %s
              AND status NOT IN ('paid','overdue','cancelled','void') AND invoice_type not in ('final','daily','extra_bill') 
        """, (today,))
        affected = cursor.rowcount
        conn.commit()
        print(f"✅ Mark invoices overdue: {affected} รายการ (as of {today})")

    except Exception as e:
        conn.rollback()
        print("❌ Error mark_invoices_overdue:", e)
    finally:
        cursor.close()
        conn.close()


def allowed_file(filename):
    allowed_extensions = {'docx', 'pdf', 'jpg', 'png'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def get_setting(key, default_value=None):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute(
        "SELECT setting_value FROM settings WHERE setting_key=%s", (key,))
    row = cursor.fetchone()
    cursor.close()
    conn.close()
    return row['setting_value'] if row else default_value


def auto_generate_all_invoices(mocked_date=None):
    """
    รวมการทำงานทั้งหมด:
    1) update late penalty ของทุกบิลที่ยังไม่จ่าย
    2) generate monthly invoices ถ้าถึงวัน
    """
    today = mocked_date 
    print(f"Today = {today}\n")

    conn = get_db_connection()
    if not conn:
        print("❌ DB connection failed")
        return
    cursor = conn.cursor(dictionary=True)

    try:
        
        mark_invoices_overdue(mocked_date=today)
        # -----------------------------------------
        # 1) UPDATE LATE PENALTY ทุก invoice
        # -----------------------------------------
        cursor.execute("""
            SELECT invoice_id,invoice_type
            FROM invoices
            WHERE status IN ('draft','pending')
              AND billing_period_end >= %s
        """, (today,))
        all_invoices = cursor.fetchall()

        print(f"🔧 Update late penalty for {len(all_invoices)} invoices")
        for inv in all_invoices:
            update_late_penalty(cursor, inv['invoice_id'])

        # 2) MARK CONTRACTS EXPIRING THIS MONTH
        mark_contracts_expiring(today, cursor)
        conn.commit()
        print("\n=== AUTO BILLING ENGINE DONE ===")

    except Exception as e:
        conn.rollback()
        print(f"❌ Error auto_generate_all_invoices: {e}")
        import traceback
        traceback.print_exc()

    finally:
        cursor.close()
        conn.close()

    # 3) AUTO MONTHLY BILL
    generate_monthly_invoices_if_due(mocked_date=today)


def calculate_prorated_rent(monthly_price, target_date):
    """
    คำนวณค่าเช่าตามจำนวนวันจริงในเดือนนั้นๆ
    """
    if not monthly_price:
        return 0
        
    # หาจำนวนวันทั้งหมดในเดือนนั้น (เช่น 28, 30, 31)
    _, total_days_in_month = calendar.monthrange(target_date.year, target_date.month)
    
    # หาราคาต่อวัน
    price_per_day = monthly_price / total_days_in_month
    
    # คำนวณตามจำนวนวันที่อยู่จริง (นับถึงวันที่ย้ายออก)
    stayed_days = target_date.day
    prorated_amount = price_per_day * stayed_days
    
    return round(prorated_amount, 2)


def prepare_placeholder_data(contract, business, options=None, settings=None):
    full_renter = f"{contract.get('fname','')} {contract.get('lname','')}".strip()
    bd = contract.get('bd')
    renew_count = contract.get('renew_count') or 0
    renew_text = f"(ต่ออายุครั้งที่ {renew_count})" if renew_count > 0 else "(สัญญาใหม่)"
    start = contract.get('contract_start')
    end = contract.get('contract_end')
    late_price = settings.get('late_fee_per_day', 0)

    # คำนวณจำนวนเดือน
    if start and end:
        delta = relativedelta(end, start)
        months = delta.years * 12 + delta.months
        renter_duration = f"{months} เดือน"
    else:
        renter_duration = "ไม่ระบุ"

    pay_date = contract.get("pay_date")  # เอาตรงจาก contracts table
    if pay_date:
        pay_first = pay_date.strftime("%d/%m/%Y")
    else:
        pay_first = ""

    price = contract.get('price') or 0
    advance_rent = contract.get('advance_rent') or 0
    premiums = contract.get('premiums') or 0
    late_price = float(settings.get('late_fee_per_day', 0)) 
    invoice_due_day = settings.get('invoice_due_day', 3)

    placeholder = {
        "#owner#": business.get("owner_name", "ผู้ให้เช่า"),
        "#business_name#": business.get("name", ""),
        "#renter#": full_renter,
        "#renew_version#": renew_text,
        "#tenant_name#": full_renter,
        "#renter_room#": contract.get('unit_name', ""),
        "#renter_age#": contract.get("age", ""),
        "#renter_address#": contract.get("address", ""),
        "#renter_id#": contract.get("id_card", ""),
        "#renter_id_release#": bd.strftime("%d/%m/%Y") if bd else "",
        "#renter_room#": contract.get('unit_name', ""), 
        "#rent_floor#": f"ชั้น {contract.get('floor','')}",
        "#renter_duration#": renter_duration,
        "#movein_date#": start.strftime("%d/%m/%Y") if start else "",
        "#room_cost#": f"{price:.2f}",
        "#room_cost_text#": num2words(price, lang='th') if price else "ศูนย์",
        "#pay_first#": pay_first,
        "#invoice_due_day#": invoice_due_day,
        "#late_price#": f"{late_price:.2f}",
        "#advance#": contract.get('premiums', 0),
        "#advance_text#": num2words(contract.get('premiums', 0), lang='th') if contract.get('premiums', 0) else "ศูนย์",
        "#now#": datetime.now().strftime("%d/%m/%Y"),
        "#address#": f"{contract.get('building','')} ชั้น {contract.get('floor','')}",
        "#fcost#": f"{price + advance_rent + premiums:.2f}",
        "#note#": "ผู้เช่ายินยอมตามเงื่อนไขสัญญา",
        "#tel#": business.get("tel", ""),
        "#email#": business.get("email", ""),
        "#tax_id#": business.get("tax_id", "")
    }

    # Options สูงสุด 6
    options = options or []
    for i in range(6):
        try:
            placeholder[f"#option{i+1}name#"] = options[i]["name"]
            placeholder[f"#option{i+1}cost#"] = f"{options[i]['price']:.2f}"+ " บาท"
        except IndexError:
            placeholder[f"#option{i+1}name#"] = ""
            placeholder[f"#option{i+1}cost#"] = ""

    return placeholder


def replace_placeholders(doc_path, data_dict, output_path):
    doc = Document(doc_path)

    def replace_in_paragraph(paragraph):
        # รวมข้อความทั้งหมด
        full_text = ''.join(run.text for run in paragraph.runs)
        for key, value in data_dict.items():
            full_text = full_text.replace(key, str(value))
        # แก้แต่ละ run ทีละตัว โดยค่อย ๆ distribute ตัวอักษรกลับ
        i = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run.text = full_text[i:i+run_len]
            i += run_len

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    for paragraph in doc.paragraphs:
        if "#now#" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("#now#", datetime.now().strftime("%d/%m/%Y"))
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        replace_in_table(table)

    for section in doc.sections:
        for header_paragraph in section.header.paragraphs:
            replace_in_paragraph(header_paragraph)
        for footer_paragraph in section.footer.paragraphs:
            replace_in_paragraph(footer_paragraph)

    doc.save(output_path)


def update_late_penalty(cursor, invoice_id):
    """อัปเดต late_penalty และ overdue_days ของ invoice"""
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    today = get_now(mocked=True).date()
    try:
        # แก้ไขจุดที่ 1: เพิ่มช่องว่างหน้า FROM และเช็คชื่อคอลัมน์
        cursor.execute(
            "SELECT due_date, total_amount, late_penalty, overdue_days, status, invoice_type "
            "FROM invoices WHERE invoice_id=%s",
            (invoice_id,)
        )
        inv = cursor.fetchone()
        
        if not inv:
            return 0
        
        # ป้องกันกรณี due_date ใน DB เป็น Null
        if not inv['due_date']:
            return 0
        
        if inv['invoice_type'] == 'first':
            return 0

        # คำนวณจำนวนวันที่เกินกำหนด
        overdue_days = max(0, (today - inv['due_date']).days)

        # ถ้าใบแจ้งหนี้สถานะไม่ต้องคำนวณใหม่ ให้คืนค่าเดิมที่มีอยู่
        if inv['status'] in ('paid', 'cancelled', 'overdue'):
            # ใช้ค่าจาก DB ถ้าไม่มีให้ใช้ที่คำนวณได้
            return inv.get('overdue_days') if inv.get('overdue_days') is not None else overdue_days

        old_penalty = float(inv['late_penalty'] or 0)
        late_fee_per_day = float(get_setting('late_fee_per_day', 100))

        # คำนวณค่าปรับส่วนเพิ่ม (เฉพาะวันที่ยังไม่เคยคิดเงิน)
        already_charged_days = old_penalty / late_fee_per_day
        pending_days = max(0, overdue_days - already_charged_days)

        if pending_days > 0:
            additional_penalty = pending_days * late_fee_per_day
            new_late_penalty = old_penalty + additional_penalty
            new_total = float(inv['total_amount'] or 0) + additional_penalty

            cursor.execute(
                "UPDATE invoices SET late_penalty=%s, total_amount=%s, overdue_days=%s WHERE invoice_id=%s",
                (new_late_penalty, new_total, overdue_days, invoice_id)
            )
            conn.commit()

        refresh_invoice_total(cursor, invoice_id)
        
        return overdue_days

    except Exception as e:
        print(f"Error update_late_penalty: {e}")
        if conn:
            conn.rollback()
        return 0  # สำคัญ: ต้องคืนค่าเป็นตัวเลข (0) เพื่อไม่ให้ Template พัง
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


def calculate_next_billing_date(start_date, today, billing_day):
    """
    Args:
        start_date: วันเริ่มสัญญา
        today: วันปัจจุบัน
        billing_day: วันที่ 1-31 (default=1)
    
    Returns:
        date: วันที่ควรออกบิล
    """
    if billing_day is None or billing_day <= 0:
        billing_day = 1
    
    # หาวันที่ billing_day ของเดือนปัจจุบัน
    try:
        current_month_billing = today.replace(day=billing_day)
    except ValueError:
        # วันที่ไม่ valid (เช่น 31 ก.พ.) → ใช้วันสุดท้ายของเดือน
        current_month_billing = (today + relativedelta(months=1)).replace(day=1) - timedelta(days=1)
    
    # ถ้าผ่านไปแล้ว → เดือนถัดไป
    if current_month_billing == today:
        try:
            return (current_month_billing + relativedelta(months=1)).replace(day=billing_day)
        except ValueError:
            return (current_month_billing + relativedelta(months=1)).replace(day=1) + relativedelta(months=1) - timedelta(days=1)
    else:
        return current_month_billing


def create_monthly_invoice(cursor, unit_id, billing_month, created_by):
    """
    สร้าง invoice รายเดือน (วิธีแยกบิล) รองรับ:
    - overdue invoice อัตโนมัติ
    - รวมค่าเช่า + options
    - ดึงค่าต่างๆ จาก settings และ contract
    """
    try:
        cursor.execute("""
            SELECT c.contract_id, c.tenant_id, c.price, c.contract_start, c.contract_end, c.billing_day,
                c.electricity_start, c.water_start, c.status
            FROM contracts c
            JOIN unit u ON c.room_id = u.unit_id  
            WHERE c.room_id=%s AND c.status IN (2,3)
            AND u.status_id != 7
            AND u.is_deleted = 0
            ORDER BY c.contract_start DESC
            LIMIT 1
        """, (unit_id,))
        contract = cursor.fetchone()
        if not contract:
            print("ไม่มีสัญญาเช่าปัจจุบัน")
            return None
    

        contract_id = contract['contract_id']
        tenant_id = contract['tenant_id']
        rent_amount = contract['price']
        billing_day = contract['billing_day'] or 1

        # 2. ดึงค่า settings
        cursor.execute("""
            SELECT setting_key, setting_value
            FROM settings
            WHERE setting_key IN ('invoice_due_day','electricity_rate','water_rate')
        """)
        settings = {s['setting_key']: s['setting_value'] for s in cursor.fetchall()}

        invoice_due_day = int(settings.get('invoice_due_day', 3))
        electricity_rate = float(settings.get('electricity_rate', 0))
        water_rate = float(settings.get('water_rate', 0))

        # 3. คำนวณวันเริ่ม-สิ้นสุด billing
        next_billing_date = calculate_next_billing_date(
            start_date=contract['contract_start'],
            today=billing_month,
            billing_day=billing_day
        )

        billing_start = next_billing_date.replace(day=1)
        billing_end = (billing_start + relativedelta(months=1)) - timedelta(days=1)
        if contract['contract_end'] and billing_end > contract['contract_end']:
            billing_end = contract['contract_end']
        if billing_start < contract['contract_start']:
            billing_start = contract['contract_start']

        is_expiring_this_month = (
            contract['contract_end'] and 
            contract['contract_end'].month == billing_month.month and 
            contract['contract_end'].year == billing_month.year
        )
                
        # ถ้าสถานะเป็น 4 (รอต่ออายุ/กำลังจะหมดสัญญา) ให้ "ข้าม" เหมือนเดิม เพื่อไปทำบิล Final
        if is_expiring_this_month or contract.get('status') == 4:
            print(f"⏭️ ข้ามการสร้างบิลปกติ: ห้อง {unit_id} จะหมดสัญญาเดือนนี้ หรือ Status=4 (จะไปสร้างบิล Final แทน)")
            return None
        # 5. ตรวจ invoice เดือนนี้ (ไม่สร้างซ้ำ)
        cursor.execute("""
            SELECT invoice_id
            FROM invoices
            WHERE unit_id=%s
            AND contract_id=%s
            AND invoice_type in ('monthly','first','final')
            AND billing_period_start=%s
            AND status NOT IN ('cancelled','void')
            LIMIT 1
        """, (unit_id, contract_id, billing_start))
        existing = cursor.fetchone()
        if existing:
            print(f"มี invoice เดือนนี้อยู่แล้ว ID: {existing['invoice_id']}")
            return existing['invoice_id']

        # 6. คำนวณ due_date
        try:
            due_date = billing_start.replace(day=invoice_due_day)
        except ValueError:
            last_day = (billing_start + relativedelta(months=1)) - timedelta(days=1)
            due_date = last_day

        # 7. ดึง previous electricity/water (เฉพาะ contract เดียวกัน)
        cursor.execute("""
            SELECT current_electricity_reading, current_water_reading, billing_period_start
            FROM invoices
            WHERE unit_id=%s AND contract_id=%s AND invoice_type='monthly'
                AND status NOT IN ('cancelled','void')
            ORDER BY billing_period_start DESC
            LIMIT 1
        """, (unit_id, contract_id))
        last_invoice = cursor.fetchone()
        if last_invoice:
            prev_elec = last_invoice['current_electricity_reading'] or 0
            prev_water = last_invoice['current_water_reading'] or 0
            last_billing_period = last_invoice['billing_period_start']
        else:
            prev_elec = contract.get('electricity_start', 0) or 0
            prev_water = contract.get('water_start', 0) or 0
            last_billing_period = None

        # 9. สร้าง invoice draft
        cursor.execute("""
            INSERT INTO invoices (
                unit_id, tenant_id, contract_id, invoice_type,
                billing_period_start, billing_period_end,
                issue_date, due_date, rent_amount,
                previous_electricity_reading, previous_water_reading,
                electricity_rate, water_rate,
                current_electricity_reading, current_water_reading,
                total_amount, status, created_by, created_at, premiums, reimburse
            ) VALUES (%s,%s,%s,%s,%s,%s,NOW(),%s,%s,
                    %s,%s,%s,%s,%s,%s,%s,%s,NULL,NOW(),0,0)
        """, (
            unit_id, tenant_id, contract_id, 'monthly',
            billing_start, billing_end,
            due_date, rent_amount,
            prev_elec, prev_water,
            electricity_rate, water_rate,
            None, None, rent_amount, 'draft'
        ))
        invoice_id = cursor.lastrowid

        # 10. ดึง options ของสัญญา
        cursor.execute("""
            SELECT o.*
            FROM `option` o
            JOIN contract_option co ON o.id = co.option_id
            WHERE co.contract_id=%s AND o.is_deleted = 0
        """, (contract_id,))
        options = cursor.fetchall()

        # 11. ใส่ option ลง invoice_items
        for opt in options:
            if opt['name'] != 'ค่าเช่าห้อง':
                cursor.execute("""
                    INSERT INTO invoice_items (invoice_id, description, unit_price, quantity, total_price, option_id)
                    VALUES (%s,%s,%s,1,%s,%s)
                """, (
                    invoice_id, opt['name'], opt['price'], opt['price'], opt['id']
                ))

        # 12. อัปเดต total_amount (รวมค่าเช่า + options)
        cursor.execute("""
            UPDATE invoices
            SET total_amount = %s 
                + COALESCE((SELECT SUM(total_price) FROM invoice_items WHERE invoice_id=%s),0),
                created_by=%s, 
                created_at=NOW()
            WHERE invoice_id=%s
        """, (
            rent_amount,
            invoice_id,
            created_by,          
            invoice_id
        ))

        print("สร้าง invoice draft สำเร็จ ID:", invoice_id)
        return invoice_id

    except Exception as e:
        print(f"❌ Error ในห้อง {unit_id}: {e}")
        raise e


def generate_monthly_invoices_if_due(mocked_date=None):
    today = mocked_date
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    if not conn:
        print("❌ เชื่อมต่อ DB ไม่สำเร็จ")
        return
    
    try:
        bill_gen_auto = int(get_setting('auto_generate_bill', 1))

        if bill_gen_auto == 0:
            print("ℹ️ ระบบสร้างบิลอัตโนมัติถูกปิด")
            return 0

        # 1) ดึงสัญญา
        cursor.execute("""
            SELECT 
                c.contract_id, c.room_id AS unit_id, c.tenant_id, c.price,
                c.contract_start, c.contract_end,
                c.billing_day
            FROM contracts c
            WHERE c.status IN (2, 3)
        """)
        contracts = cursor.fetchall()

        if not contracts:
            print("❌ ไม่มีสัญญา active")
            return

        print(f"📋 พบสัญญา {len(contracts)} รายการ\n")

        created_count = 0

        # 2) Loop ทีละสัญญา
        for c in contracts:
            print(f"--- 🏠 ห้อง {c['unit_id']} (contract {c['contract_id']}) ---")

            # 2.1) เช็กบิลซ้ำ
            cursor.execute("""
                SELECT invoice_id
                FROM invoices
                WHERE contract_id = %s
                AND invoice_type in ('monthly','final','first')
                AND MONTH(issue_date) = %s
                AND YEAR(issue_date) = %s
                LIMIT 1
            """, (c['contract_id'], today.month, today.year))

            exist = cursor.fetchone()

            if exist:
                print(f"⏭️ ข้าม → มีบิลเดือนนี้แล้ว (ID {exist['invoice_id']})\n")
                conn.rollback()
                continue

            # 2.2) คำนวณวันออกบิล
            billing_day = c['billing_day'] or 1
            print(f"📅 billing_day = {billing_day}")

            next_billing_date = calculate_next_billing_date(
                start_date=c['contract_start'],
                today=today,
                billing_day=billing_day
            )
            print(f"📅 next_billing_date = {next_billing_date}")

            if today < next_billing_date:
                print("⏭️ ข้าม → ยังไม่ถึงวันออกบิล\n")
                conn.rollback()
                continue

            print("✔ ถึงวันออกบิลแล้ว (today >= next_billing_date)")

            # 2.3) สร้างบิล
            invoice_id = create_monthly_invoice(
                cursor=cursor,
                unit_id=c['unit_id'],
                billing_month=today,
                created_by=None
            )

            add_audit_log(
                cursor, 
                'INVOICE', 
                'CREATE', 
                f'สร้างบิลเดือน {today.month}/{today.year} สำหรับห้อง ID: {c["unit_id"]}', 
                None
            )

            print(f"🧾 create_monthly_invoice() return = {invoice_id}")

            if not invoice_id:
                print("❌ ERROR → create_monthly_invoice ไม่คืน invoice_id\n")
                conn.rollback()
                continue

            conn.commit()
            created_count += 1
            print(f"✅ สร้างบิลสำเร็จ ID {invoice_id}\n")

        print(f"🎉 รวมสร้างบิลใหม่ทั้งหมด = {created_count} ใบ\n")

    except Exception as e:
        conn.rollback()
        print("❌ ERROR generate_monthly_invoices_if_due:", e)
    finally:
        cursor.close()
        conn.close()


def auto_delete_contract_files():
    """
    ลบไฟล์สัญญาเมื่อผู้เช่าย้ายออก หรือสัญญาหมดอายุ (status = 5 หรือ 6)
    """
    today = date.today()
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        # ดึง contract ที่มีไฟล์ และหมดอายุ / ย้ายออกแล้ว
        cursor.execute("""
            SELECT contract_id, contracts_file
            FROM contracts
            WHERE contracts_file IS NOT NULL
              AND (
                    (actual_move_out_date IS NOT NULL AND actual_move_out_date <= %s)
                    OR status IN (5,6)
                  )
        """, (today,))

        rows = cursor.fetchall()

        # ลบไฟล์ทีละไฟล์
        for row in rows:
            file_path = os.path.join("/home/precise/flask_web/contracts_file", row['contracts_file'])

            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    print(f"✅ ลบไฟล์ {row['contracts_file']} ของ contract {row['contract_id']}")
                else:
                    print(f"⚠️ ไม่พบไฟล์ {row['contracts_file']}")
            except Exception as e:
                print(f"❌ ลบไฟล์ {row['contracts_file']} ไม่สำเร็จ: {e}")

        # อัปเดต DB ให้เป็น NULL
        cursor.execute("""
            UPDATE contracts
            SET contracts_file = NULL
            WHERE (actual_move_out_date IS NOT NULL AND actual_move_out_date <= %s)
               OR status IN (5,6)
        """, (today,))
        conn.commit()

        print("✅ เคลียร์ contracts_file ใน DB เรียบร้อย")

    except Exception as e:
        conn.rollback()
        print("❌ Error auto_delete_contract_files:", e)

    finally:
        cursor.close()
        conn.close()


def check_meter_save(cursor, invoice_id): 
    """
    เช็คสถานะและอัปเดต meter_saved ภายใต้ cursor เดิม
    """
    try:
        # 1. เช็คเงื่อนไข
        cursor.execute("""
            SELECT invoice_id FROM invoices  
            WHERE invoice_id = %s 
              AND current_electricity_reading IS NOT NULL 
              AND current_water_reading IS NOT NULL
        """, (invoice_id,))
        
        if cursor.fetchone():
            # 2. อัปเดต (ไม่ต้องสั่ง conn.commit ตรงนี้ เดี๋ยวไป commit ทีเดียวที่ตัวแม่)
            cursor.execute("""
                UPDATE invoices SET meter_saved = 1 WHERE invoice_id = %s
            """, (invoice_id,))
            print(f"✅ บิล ID {invoice_id} อัปเดต meter_saved = 1 เรียบร้อย")
            return True
        return False
    except Exception as e:
        print("❌ Error check_meter_save:", e)
        return False
    

def refresh_invoice_total(cursor, invoice_id):
    # 1. ดึงข้อมูลบิลและข้อมูลสัญญา
    cursor.execute("""
        SELECT i.*, c.price as contract_price 
        FROM invoices i 
        LEFT JOIN contracts c ON i.contract_id = c.contract_id 
        WHERE i.invoice_id = %s
    """, (invoice_id,))
    inv = cursor.fetchone()
    if not inv: return

    # 🚩 ตัวกั้น: จ่ายแล้วห้ามแก้เงิน
    if inv['status'] in ('paid', 'cancelled'):
        return

    # 2. ดึงเรทน้ำไฟ
    elec_rate = float(get_setting('electricity_rate', 0))
    water_rate = float(get_setting('water_rate', 0))

    # 3. สรุปยอดรายการย่อย
    cursor.execute("""
        SELECT 
            SUM(CASE WHEN type IN ('service', 'option') THEN total_price ELSE 0 END) as service_option_sum,
            SUM(CASE WHEN type = 'penalty' THEN total_price ELSE 0 END) as penalty_sum,
            SUM(CASE WHEN type = 'discount' THEN total_price ELSE 0 END) as discount_sum,
            SUM(CASE WHEN type = 'meter_adjustment' THEN total_price ELSE 0 END) as meter_adjustment_sum
        FROM invoice_items WHERE invoice_id = %s
    """, (invoice_id,))
    sums = cursor.fetchone()

    # เตรียมค่าตัวเลข
    premiums = float(inv['premiums'] or 0)
    reimburse = float(inv['reimburse'] or 0)
    # รวมค่าปรับจากทั้งตารางหลัก (late_penalty) และรายการย่อย (penalty_sum)
    penalty = float(sums['penalty_sum'] or 0) + float(inv['late_penalty'] or 0)
    discount = float(sums['discount_sum'] or 0)
    service_charge = float(sums['service_option_sum'] or 0)
    meter_adjustment = float(sums['meter_adjustment_sum'] or 0) # ดึงยอดค้างจากการเปลี่ยนมิเตอร์มาเตรียมไว้

    # --- [Logic แยกตามประเภทบิล] ---
    if inv['invoice_type'] == 'first':
        # เดือนแรก: จ่ายแค่ประกัน (และค่าปรับ/ส่วนลด/บริการพิเศษ ถ้ามีใส่มา)
        rent = 0.0
        elec_total = 0.0
        water_total = 0.0
        meter_adjustment = 0.0 # เดือนแรกยังไม่เก็บเงินเปลี่ยนมิเตอร์ (รอเก็บเดือนหน้าตามที่คุณต้องการ)

    elif inv['invoice_type'] == 'daily':
        elec_total = 0.0
        water_total = 0.0
        meter_adjustment = 0.0 
        rent = float(inv['rent_amount']) if inv['rent_amount'] is not None else 0.0
        
    else: # 'normal' หรือ 'final'
        rent = float(inv['rent_amount'] or inv['contract_price'] or 0)

        elec_total = (inv['electricity_usage'] or 0) * elec_rate
        water_total = (inv['water_usage'] or 0) * water_rate

    # 4. คำนวณยอดสุทธิ
    new_total = (rent + premiums + elec_total + water_total + service_charge + penalty + meter_adjustment) - (discount + reimburse)

    # 5. อัปเดตข้อมูลกลับ
    cursor.execute("""
        UPDATE invoices SET 
            electricity_rate = %s, water_rate = %s,
            electricity_total = %s, water_total = %s, 
            service_charge = %s, rent_amount = %s,
            total_amount = %s
        WHERE invoice_id = %s
    """, (elec_rate, water_rate, elec_total, water_total, service_charge, rent, new_total, invoice_id))


def log_meter_reading(cursor, unit_id, meter_type, reading_value, source='auto', invoice_id=None, created_by=None):
    """
    ฟังก์ชันกลางสำหรับบันทึกเลขมิเตอร์ลงฐานข้อมูล
    """
    try:
        table = "meter" if meter_type == "electricity" else "meter_water"
        cursor.execute(f"SELECT id, serial_meter FROM {table} WHERE unit_id = %s LIMIT 1", (unit_id,))
        meter_info = cursor.fetchone()

        if not meter_info:
            return False, f"ไม่พบมิเตอร์ประเภท {meter_type} สำหรับห้อง {unit_id}"

        insert_reading = """
            INSERT INTO meter_reading 
            (unit_id, meter_type, meter_id, serial_meter, current_reading, source, invoice_id, read_date,created_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s, NOW(),%s)
        """
        cursor.execute(insert_reading, (
            unit_id, 
            meter_type, 
            meter_info['id'], 
            meter_info['serial_meter'], 
            reading_value, 
            source, 
            invoice_id,
            created_by
        ))

        update_sql = f"UPDATE {table} SET current_reading = %s, updated_at = NOW() WHERE id = %s"
        cursor.execute(update_sql, (reading_value, meter_info['id']))

        return True, "บันทึกสำเร็จ"

    except Exception as e:
        return False, str(e)
    

def auto_read_all_systems():
    """
    งานอ่านมิเตอร์อัตโนมัติ (รันทุกชั่วโมง) 
    บันทึกค่าลง meter, meter_water และ meter_reading
    """
    print("--- เริ่มอ่านมิเตอร์อัตโนมัติ ---")
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    try:
        # อ่านทั้ง 2 ระบบ (electricity และ water)
        for m_type in ['electricity', 'water']:
            table = 'meter' if m_type == 'electricity' else 'meter_water'
            reg_key = 'total_kWh' if m_type == 'electricity' else 'total_m3'
            is_water = (m_type == 'water')

            query = f"""
                SELECT m.* FROM {table} m
                JOIN unit u ON m.unit_id = u.unit_id
                WHERE m.status = 'active' AND u.is_deleted = 0
            """
            cursor.execute(query)
            meters = cursor.fetchall()

            for m in meters:
                # 2. อ่านค่าจาก Hardware
                val = read_meter_unit_read(
                    model_name=m['module'],
                    register_key=reg_key,
                    serial_ports=m['comport'],
                    ip=m['ip'],
                    port=m['port'],
                    slave_id=m['slave_id'], # แก้จาก unit_id เป็น slave_id ตามโครงสร้างฟังก์ชันคุณ
                    is_water=is_water,
                    api_base_url=m['base_url'], 
                    api_token=m['api_auth_token']
                )

                if val is not None and not isinstance(val, str):
                    # ฟังก์ชันนี้จะ UPDATE current_reading และ INSERT ลง meter_reading ให้เอง
                    success, msg = log_meter_reading(
                        cursor, 
                        unit_id=m['unit_id'], 
                        meter_type=m_type, 
                        reading_value=val, 
                        source='auto', # ระบุที่มาเป็น auto
                        created_by=None
                    )
                    if success:
                        print(f"✅ [{m_type}] ห้อง {m['unit_id']}: {msg} ({val})")
                    else:
                        # ถ้าเข้าฟังก์ชัน log แต่บันทึกไม่สำเร็จ จะเห็นสาเหตุตรงนี้
                        print(f"❌ [{m_type}] ห้อง {m['unit_id']}: บันทึกไม่สำเร็จเพราะ {msg}")
                else:
                    print(f"[{m_type}] ห้อง {m['unit_id']}: อ่านค่าล้มเหลว")

        conn.commit() # ยืนยันการบันทึกทั้งหมด
    except Exception as e:
        print(f"Auto Job Error: {e}")
    finally:
        cursor.close()
        conn.close()
    print("--- จบการทำงานอัตโนมัติ ---")


def action_sync_latest_meter_to_invoices():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    warnings = []
    updated_count = 0

    try:
        # 1. ค้นหาบิลที่เป็น Draft ทั้งหมด (ไม่ต้อง JOIN meter_reading แล้ว เพราะเราจะกวาดทุกบิล Draft)
        cursor.execute("SELECT invoice_id, unit_id, previous_electricity_reading, previous_water_reading, invoice_type FROM invoices WHERE status = 'draft'")
        invoices_to_sync = cursor.fetchall()

        if not invoices_to_sync:
            return True, "ไม่มีบิล Draft ให้ Sync", []

        for inv in invoices_to_sync:
            u_id = inv['unit_id']
            inv_id = inv['invoice_id']

            # 🚩 2. เช็คประวัติเปลี่ยนมิเตอร์ (ถ้า is_billed = 0 ใน history คือห้าม Sync)
            cursor.execute("SELECT id FROM meter_history WHERE unit_id = %s AND is_billed = 0 LIMIT 1", (u_id,))
            if cursor.fetchone():
                warnings.append(f"ห้อง {u_id}: มีประวัติเปลี่ยนมิเตอร์ค้างอยู่ กรุณาบันทึกมือที่หน้าบิล")
                continue

            # 🚩 3. เช็ค Extra Bill (ถ้ามีรายการค่าใช้จ่ายเพิ่มเติมให้ข้าม)
            cursor.execute("""
                SELECT invoice_type FROM invoices 
                WHERE invoice_id = %s AND invoice_type = 'extra_bill'
                LIMIT 1
            """, (inv_id,))
            if cursor.fetchone():
                warnings.append(f"ห้อง {u_id}: มีค่าใช้จ่ายเพิ่มเติม (Extra Bill) ระบบข้ามการ Sync เพื่อให้ตรวจสอบมือ")
                continue

            # 3. อ่านค่าล่าสุดจริงๆ จาก meter_reading (ไม่สน invoice_id)
            # ดึงไฟฟ้า
            cursor.execute("""
                SELECT current_reading FROM meter_reading 
                WHERE unit_id = %s AND meter_type = 'electricity' 
                ORDER BY read_date DESC LIMIT 1
            """, (u_id,))
            elec_data = cursor.fetchone()

            # ดึงน้ำ
            cursor.execute("""
                SELECT current_reading FROM meter_reading 
                WHERE unit_id = %s AND meter_type = 'water' 
                ORDER BY read_date DESC LIMIT 1
            """, (u_id,))
            water_data = cursor.fetchone()

            # ถ้าไม่มีข้อมูลใน meter_reading เลย ให้ข้ามห้องนี้ไป
            if not elec_data and not water_data:
                continue

            # 4. เตรียมค่า
            new_elec = float(elec_data['current_reading']) if elec_data else float(inv['previous_electricity_reading'] or 0)
            new_water = float(water_data['current_reading']) if water_data else float(inv['previous_water_reading'] or 0)

            # 5. ตรวจสอบเลขถอยหลัง (กันพลาด)
            # ถ้าเลขใหม่น้อยกว่าเลขเก่า จะข้าม (ยกเว้นบิลเดือนแรก)
            if inv['invoice_type'] != 'first':
                if (elec_data and new_elec < float(inv['previous_electricity_reading'] or 0)) or \
                   (water_data and new_water < float(inv['previous_water_reading'] or 0)):
                    warnings.append(f"ห้อง {u_id}: เลขใหม่น้อยกว่าเดิม ระบบข้ามการ Sync")
                    continue

            # 6. อัปเดตลงบิล
            elec_usage = max(0, new_elec - float(inv['previous_electricity_reading'] or 0))
            water_usage = max(0, new_water - float(inv['previous_water_reading'] or 0))

            cursor.execute("""
                UPDATE invoices SET 
                    current_electricity_reading = %s, 
                    current_water_reading = %s,
                    electricity_usage = %s,
                    water_usage = %s,
                    meter_saved = 1
                WHERE invoice_id = %s
            """, (new_elec, new_water, elec_usage, water_usage, inv_id))

            # 7. อัปเดตเลขเริ่มของห้อง (Unit)
            cursor.execute("UPDATE unit SET electricity_start = %s, water_start = %s WHERE unit_id = %s", (new_elec, new_water, u_id))

            # 8. คำนวณยอดเงินรวมใหม่
            refresh_invoice_total(cursor, inv_id)
            updated_count += 1

        conn.commit()
        return True, updated_count, warnings

    except Exception as e:
        if conn: conn.rollback()
        return False, str(e), []
    finally:
        if cursor: cursor.close()
        if conn: conn.close()


def add_audit_log(cursor, category, action, description, user_id=None):
    """
    บันทึกประวัติกิจกรรมลงในตาราง audit_logs
    """
    sql = """
        INSERT INTO audit_log (category, action, description, created_by, created_at)
        VALUES (%s, %s, %s, %s, %s)
    """
    now = datetime.now()
    
    try:
        cursor.execute(sql, (
            category, 
            action, 
            description, 
            user_id, 
            now
        ))
    except Exception as e:
        print(f"❌ Error while adding audit log: {e}")