from flask import Flask, render_template, request, redirect, url_for, session, flash, send_from_directory, send_file, abort, jsonify, current_app
from functools import wraps
from datetime import datetime, date, timedelta
import os
from dotenv import load_dotenv
import re
from werkzeug.utils import secure_filename
from docx import Document
from utils import (allowed_file, auto_generate_all_invoices, prepare_placeholder_data, replace_placeholders, 
get_db_connection, get_now,update_late_penalty, record_transaction, create_monthly_invoice,
get_setting,check_meter_save,log_meter_reading,refresh_invoice_total,auto_read_all_systems,
action_sync_latest_meter_to_invoices,read_config,PATH_MODEL_ELEC,PATH_MODEL_WATER,add_audit_log,
generate_slip_filename,cleanup_expired_contracts)

from dateutil.relativedelta import relativedelta
import serial.tools.list_ports
import json
from master_modbus import read_meter_tool, write_meter_tool, read_meter_unit_read
from apscheduler.schedulers.background import BackgroundScheduler
from api import api
import fcntl
from werkzeug.security import check_password_hash,generate_password_hash

today = get_now(mocked=False).date()

scheduler = BackgroundScheduler()

BASE_DIR = os.path.abspath(os.path.dirname(__file__))

UPLOAD_FOLDER = 'uploaded_docs'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'jpg', 'png'}

CONTRACTS_FOLDER = os.path.join(BASE_DIR, 'contracts_file')
BASE_UPLOAD_PATH = os.path.join(BASE_DIR, 'uploads')
EXPENSE_UPLOAD_PATH = os.path.join(BASE_UPLOAD_PATH, 'expense_slips')
INCOME_UPLOAD_PATH = os.path.join(BASE_UPLOAD_PATH, 'income_slips')
UPLOAD_ID_CARD = os.path.join(BASE_DIR, 'uploads', 'id_card')    
UPLOAD_PROFILE = os.path.join(BASE_DIR, 'static', 'profile_user')
UPLOAD_ID_CARD_TENANTS = os.path.join(BASE_DIR, 'uploads', 'id_card_tenants')

PATHS = [UPLOAD_ID_CARD, EXPENSE_UPLOAD_PATH, INCOME_UPLOAD_PATH, UPLOAD_PROFILE, UPLOAD_ID_CARD_TENANTS, CONTRACTS_FOLDER, UPLOAD_ID_CARD_TENANTS]
for path in PATHS:
    os.makedirs(path, exist_ok=True)

load_dotenv()    

app = Flask(__name__)
app.register_blueprint(api)
app.secret_key = os.getenv('FLASK_SECRET_KEY')
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=15) # ตั้งเวลาที่ต้องการ
app.config['SESSION_REFRESH_EACH_REQUEST'] = True 

auto_generate_all_invoices(mocked_date=today)
cleanup_expired_contracts(mocked_date=today,contracts_folder_path=CONTRACTS_FOLDER)

app.config['UPLOAD_FOLDER_DOCS'] = UPLOAD_FOLDER

def job_read_meters_task():
    now = datetime.now()
    print(f"Reading Meters at {now.strftime('%H:%M:%S')}")
    auto_read_all_systems() 

def job_invoices_task():
    now = datetime.now()
    print(f"Generating Invoices for: {now()}")
    auto_generate_all_invoices(mocked_date=today)

def log_activity(category, action, description):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            response = f(*args, **kwargs)
            try:
                conn = get_db_connection()
                cursor = conn.cursor()
                current_user_id = session.get('user', {}).get('user_id')
                
                add_audit_log(
                    cursor, 
                    category, 
                    action, 
                    description, 
                    user_id=current_user_id
                )
                
                conn.commit()
                cursor.close()
                conn.close()
            except Exception as e:
                print(f"Log Error: {e}")
                
            return response
        return decorated_function
    return decorator

@app.before_request
def check_session_timeout():
    allowed_endpoints = ['login', 'static'] 
    
    if request.endpoint not in allowed_endpoints:
        if 'user' not in session:
            return redirect(url_for('login'))

def role_required(allowed_roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user' not in session:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'status': 'error', 'message': 'กรุณาเข้าสู่ระบบใหม่'}), 401
                return redirect(url_for('login'))
            
            user_role = session.get('role') 
            if user_role not in allowed_roles:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'status': 'error', 'message': 'คุณไม่มีสิทธิ์ดำเนินการ (Admin/Manager Only)'}), 403
                flash("คุณไม่มีสิทธิ์เข้าถึงหน้านี้", "danger")
                return redirect(url_for('dashboard'))
            
            return f(*args, **kwargs)
        return decorated_function
    return decorator

lock_file = open(".scheduler.lock", "wb")

def start_scheduled_jobs():
    try:
        fcntl.flock(lock_file, fcntl.LOCK_EX | fcntl.LOCK_NB)
        
        scheduler = BackgroundScheduler()
        scheduler.add_job(job_read_meters_task, 'cron', hour=0, minute=0, id='read_meter_hourly')
        scheduler.add_job(job_invoices_task, 'cron', hour=0, minute=0, id='create_invoices')
        scheduler.start()
        
        print("✅ [MASTER] Scheduler started successfully on this worker.")
        
    except IOError:
        # 4. ถ้าจองไม่สำเร็จ (Slave Workers) ให้ข้ามไป ไม่ต้องรัน Scheduler
        print("❌ [SLAVE] Scheduler is already running on another worker. Skipping...")

start_scheduled_jobs()

# สร้างโฟลเดอร์นี้ถ้ายังไม่มี
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

@app.context_processor
def inject_user():
    user_data = session.get('user', {})
    
    return dict(
        current_username=user_data.get('username'),
        current_user_gender=user_data.get('gender'),
        current_user_img=user_data.get('profile_img'), 
        user=user_data 
    )

# ---------------------- AUTH ----------------------
@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password'].strip()

        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        # ค้นหาด้วย username เพียงอย่างเดียว
        query = """
            SELECT u.*, r.r_name 
            FROM user u 
            JOIN role r ON u.role_id = r.role_id 
            WHERE u.username = %s AND u.is_deleted = 0
        """
        cursor.execute(query, (username,))
        user = cursor.fetchone()
        cursor.close()
        conn.close()

        # ตรวจสอบ: 1. มี user นี้ไหม 2. รหัสผ่านที่ Hash แล้วตรงกันไหม
        if user and check_password_hash(user['password'], password):
            session.permanent = True
            session['user'] = {
                'user_id': user['user_id'],
                'username': user['username'],
                'gender': user['gender'],
                'role_id': user['role_id'],
                'role_name': user['r_name'],
                'profile_img': user.get('profile_img')
            }
            session['role'] = user['r_name']
            flash('เข้าสู่ระบบสำเร็จ', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง หรือบัญชีถูกระงับ', 'danger')
            return render_template('login.html')

    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('user', None)
    flash('ออกจากระบบเรียบร้อย', 'success')
    return redirect(url_for('login'))

@app.route('/add_user', methods=['GET', 'POST'])
def add_user():

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT * FROM role")
    roles = cursor.fetchall()

    if request.method == 'POST':
        # รับข้อมูลจากฟอร์ม
        username = request.form['username'].strip()
        password = request.form['password'].strip()
        fname = request.form['fname'].strip()
        lname = request.form['lname'].strip()
        id_card = request.form['id_card'].strip()
        gender = request.form['gender']
        email = request.form['email'].strip()
        tel = request.form['tel'].strip()
        role_id = request.form['role_id']
        start_card = request.form['start_card']
        end_card = request.form['end_card']

        cursor.execute("SELECT user_id FROM user WHERE username = %s OR email = %s OR id_card = %s", 
                       (username, email, id_card))
        existing = cursor.fetchone()
        if existing:
            flash('ไม่สามารถเพิ่มได้: ชื่อผู้ใช้, อีเมล หรือเลขบัตรประชาชนนี้มีอยู่ในระบบแล้ว', 'danger')
            return render_template('add_user.html', roles=roles)

        if len(password) < 6:
            flash('รหัสผ่านต้องมีความยาวอย่างน้อย 6 ตัวอักษร', 'warning')
            return render_template('add_user.html', roles=roles)
        
        if not re.search("[a-z]", password) or not re.search("[A-Z]", password) or not re.search("[0-123456789]", password):
            flash('รหัสผ่านควรประกอบด้วยตัวพิมพ์ใหญ่ ตัวพิมพ์เล็ก และตัวเลข', 'warning')
            return render_template('add_user.html', roles=roles)

        hashed_password = generate_password_hash(password)

        try:
            query = """INSERT INTO user 
                       (username, password, fname, lname, id_card, gender, email, tel, role_id, start_card, end_card,created_at) 
                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())"""
            cursor.execute(query, (username, hashed_password, fname, lname, id_card, gender, email, tel, role_id, start_card, end_card))
            new_user_id = cursor.lastrowid

            print(f"New user created with ID: {new_user_id}")

            file_profile = request.files.get('profile_img')
            if file_profile and file_profile.filename != '':
                ext = file_profile.filename.rsplit('.', 1)[1].lower()
                filename = f"profile_{new_user_id}.{ext}" 
                file_profile.save(os.path.join(UPLOAD_PROFILE, filename))
                cursor.execute("UPDATE user SET profile_img=%s WHERE user_id=%s", (filename, new_user_id))

            file_id_card = request.files.get('id_card_file')
            if file_id_card and file_id_card.filename != '':
                ext = file_id_card.filename.rsplit('.', 1)[1].lower()
                filename = f"idcard_{new_user_id}.{ext}" 
                file_id_card.save(os.path.join(UPLOAD_ID_CARD, filename))
                cursor.execute("UPDATE user SET id_card_file=%s WHERE user_id=%s", (filename, new_user_id))

            add_audit_log(
                cursor, 
                'USER', 
                'INSERT', 
                f'เพิ่มผู้ใช้งานใหม่ ID: {new_user_id}', 
                session.get('user', {}).get('user_id')
            )
            conn.commit()
            cursor.close()
            conn.close()
            
            flash('เพิ่มผู้ใช้งานใหม่เรียบร้อยแล้ว', 'success')
            return redirect(url_for('user_settings')) # ปรับเป็น route ที่คุณใช้งาน

        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {str(e)}', 'danger')

    # 3. ปิดการเชื่อมต่อสำหรับกรณีที่เป็น GET หรือ POST แล้ว Error
    cursor.close()
    conn.close()
    
    return render_template('add_user.html', roles=roles)

@app.route('/uploads/<path:foldername>/<filename>')
def custom_uploads(foldername, filename):
    upload_path = os.path.join('uploads', foldername)
    return send_from_directory(upload_path, filename)

# ---------------------- DASHBOARD ----------------------
@app.route('/dashboard')
def dashboard():
    STATUS_MAPPING = {
        1: 'ว่าง', 2: 'มีผู้เช่า', 3: 'ทำสัญญาเช่า',
        4: 'ปิดปรับปรุง', 5: 'รอการชำระ',
        6: 'ชําระบิลสุดท้าย/ตรวจสอบสภาพห้อง'
    }
    # ------------------ get filters ------------------
    number = request.args.get('number', '').strip()
    building = request.args.get('building', '').strip()
    floor = request.args.get('floor', '').strip()
    status_id = request.args.get('status_id', '').strip()

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT setting_value FROM settings WHERE setting_key = 'min_stay_before_early_move'")
    min_stay = cursor.fetchone()
    min_stay_config = int(min_stay['setting_value']) if min_stay else 0

    # ------------------ fetch rooms based on filters ------------------
    base_query = "SELECT * FROM unit WHERE 1=1 and is_deleted=0"
    params = []

    if number:
        base_query += " AND name LIKE %s"
        params.append("%" + number + "%")

    if building:
        base_query += " AND building=%s"
        params.append(building)

    if floor:
        base_query += " AND floor=%s"
        params.append(floor)

    if status_id:
        base_query += " AND status_id=%s"
        params.append(status_id)

    cursor.execute(base_query, params)
    units = cursor.fetchall()

    rooms = []

    if not units:
        return render_template(
            "dashboard.html",
            rooms=rooms
        )

    # 2) ดึงสัญญาล่าสุดของห้องที่แสดง (เฉพาะผลค้นหา)
    unit_ids = [u['unit_id'] for u in units]

    cursor.execute(f"""
        SELECT c.*, t.fname AS tenant_fname, t.lname AS tenant_lname
        FROM contracts c
        LEFT JOIN tenants t ON t.tenant_id = c.tenant_id
        WHERE c.contract_id IN (
            SELECT MAX(contract_id)
            FROM contracts
            WHERE room_id IN ({','.join(['%s']*len(unit_ids))})
              AND status IN (1,2,3,4) and is_deleted=0
            GROUP BY room_id
        )
    """, tuple(unit_ids))

    rows = cursor.fetchall()
    contract_map = {r['room_id']: r for r in rows}

    # ------------------ contract options ------------------
    contract_ids = [r['contract_id'] for r in rows]
    options_map = {}
    pending_map = {}

    if contract_ids:
        cursor.execute(f"""
            SELECT co.contract_id, o.name
            FROM contract_option co
            JOIN `option` o ON o.id = co.option_id
            WHERE co.contract_id IN ({','.join(['%s']*len(contract_ids))})
        """, tuple(contract_ids))

        for r in cursor.fetchall():
            options_map.setdefault(r['contract_id'], []).append(r['name'])

    # ------------------ latest daily invoice ------------------
    cursor.execute("""
        SELECT i.unit_id, i.invoice_id, i.billing_period_start, i.billing_period_end,
            i.guest_fname, i.guest_lname, i.status, i.meter_saved, i.previous_electricity_reading , i.current_electricity_reading,
            i.previous_water_reading, i.current_water_reading
        FROM invoices i
        JOIN (
            SELECT unit_id, MAX(invoice_id) AS invoice_id
            FROM invoices
            WHERE invoice_type = 'daily'
            AND status IN ('draft','paid')
            GROUP BY unit_id
        ) li ON i.invoice_id = li.invoice_id
    """)
    daily_rows = cursor.fetchall()
    daily_map = {r['unit_id']: r for r in daily_rows}

    # ------------------ current monthly invoices ------------------
    if contract_ids:
        cursor.execute(f"""
            SELECT *
            FROM invoices
            WHERE invoice_id IN (
                SELECT MAX(invoice_id)
                FROM invoices
                WHERE invoice_type IN ('monthly','first','final')
                AND status IN ('draft','unpaid','overdue','paid')
                AND contract_id IN ({','.join(['%s']*len(contract_ids))})
                GROUP BY contract_id
            )
        """, tuple(contract_ids))
        invoice_rows = cursor.fetchall()
        invoice_map = {r['contract_id']: r for r in invoice_rows}

    # 2. --- ดึงบิลที่ค้างทั้งหมด (รวม extra_bill และบิลอื่นๆ เพื่อใช้นับเลขสรุป) ---
        cursor.execute(f"""
            SELECT contract_id, status 
            FROM invoices 
            WHERE status IN ('draft', 'unpaid', 'overdue')
            AND contract_id IN ({','.join(['%s']*len(contract_ids))})
        """, contract_ids)
        all_pending_rows = cursor.fetchall()
        for inv in all_pending_rows:
            pending_map.setdefault(inv['contract_id'], []).append(inv['status'])

    # ------------------ build rooms list ------------------
    rooms = []
    for u in units:
        contract = contract_map.get(u['unit_id'])
        daily = daily_map.get(u['unit_id'])
        expired_contract = None

        if contract and contract.get('status') == 4:
            expired_contract = contract

        if contract and contract.get('status') in (6,7):
            contract = None

        current_invoice = invoice_map.get(contract['contract_id']) if contract else None
        has_daily_booking = 1 if daily else 0

        if has_daily_booking:
            # รายวัน: ดึงจากบิลรายวันล่าสุด
            prev_elec = daily.get('previous_electricity_reading')
            prev_wat  = daily.get('previous_water_reading')
            curr_elec = daily.get('current_electricity_reading')
            curr_wat  = daily.get('current_water_reading')
            tenant_name = f"{daily.get('guest_fname','')} {daily.get('guest_lname','')}".strip()
            billing_period_start = daily.get('billing_period_start')
            billing_period_end = daily.get('billing_period_end')
        elif current_invoice:
            # รายเดือน: ดึงจากบิลรายเดือนล่าสุด
            prev_elec = current_invoice.get('previous_electricity_reading')
            prev_wat  = current_invoice.get('previous_water_reading')
            curr_elec = current_invoice.get('current_electricity_reading')
            curr_wat  = current_invoice.get('current_water_reading')
            tenant_name = f"{contract.get('tenant_fname','')} {contract.get('tenant_lname','')}".strip()
            billing_period_start = None
            billing_period_end = None
        else:
            # ไม่มีบิลเลย: ใช้ค่าเริ่มต้นจากห้อง (unit)
            prev_elec = u.get('electricity_start', 0)
            prev_wat  = u.get('water_start', 0)
            curr_elec = None
            curr_wat  = None
            tenant_name = f"{contract.get('tenant_fname','')} {contract.get('tenant_lname','')}".strip() if contract else '-'
            billing_period_start = None
            billing_period_end = None

        options_selected = options_map.get(contract['contract_id'], []) if contract else []
        pay_date = contract.get('pay_date') if contract else None
        contract_status = contract.get('status') if contract else None
        current_invoice = invoice_map.get(contract['contract_id']) if contract else None

        cursor.execute("""
            SELECT is_billed FROM meter_history 
            WHERE unit_id = %s ORDER BY id DESC LIMIT 1
        """, (u['unit_id'],))
        result = cursor.fetchone() 
        if result:
            is_billed_status = result['is_billed']
        else:
            is_billed_status = 1

        rooms.append({
            'has_pending_bill': True if (contract and contract.get('contract_id') in pending_map) else False,
            'today': today,
            'id': u['unit_id'],
            'unit_id': u['unit_id'],
            'number': u['name'],
            'building': u['building'],
            'floor': u['floor'],
            'status_id': u['status_id'],
            'status': STATUS_MAPPING.get(u['status_id'], 'ไม่ทราบสถานะ'),
            'bill' : is_billed_status,
            'tenant_name': tenant_name,
            'electricity_start': u.get('electricity_start'),
            'water_start': u.get('water_start'),
            'meter_id': u['meter_id'],
            'meter_water_id': u['meter_water_id'],
            'options_selected': options_selected,
            'contract_id': contract['contract_id'] if contract else None,
            'contract_status': contract_status,
            'expired_contract': expired_contract,
            'contract_file': contract.get('contracts_file') if contract else None,
            'contract_start': contract.get('contract_start') if contract else None,
            'contract_end': contract.get('contract_end') if contract else None,
            'pay_date': pay_date.isoformat() if isinstance(pay_date, date) else None,
            'guest_fname': daily.get('guest_fname') if daily else None,
            'guest_lname': daily.get('guest_lname') if daily else None,
            'has_daily_booking': has_daily_booking,
            'latest_daily_invoice_id': daily.get('invoice_id') if daily else None,
            'latest_daily_invoice_status': daily.get('status') if daily else None,
            'latest_daily_invoice_type': daily.get('invoice_type') if daily else None,  
            'invoice_id': daily.get('invoice_id') if daily else None, 
            'billing_period_start': billing_period_start,
            'billing_period_end': billing_period_end,
            'notice_move_out_date': contract.get('notice_move_out_date') if contract else None,
            'active_move_out_date': contract.get('move_out_date') if contract else None,
            'previous_electricity_reading': prev_elec,
            'previous_water_reading': prev_wat,
            'current_electricity_reading': curr_elec,
            'current_water_reading': curr_wat,
            'current_invoice': current_invoice,
            'min_stay_config': min_stay_config,
            'i_status': current_invoice.get('status') if current_invoice else None,
            'i_type': current_invoice.get('invoice_type') if contract else None,
            'meter_saved': (
                current_invoice['meter_saved'] if current_invoice and 'meter_saved' in current_invoice else
                (daily['meter_saved'] if daily and 'meter_saved' in daily else None)
            )
        })
    return render_template(
        'dashboard.html',
        rooms=rooms,today=today
    )

@app.route("/unit_pending/<int:contract_id>")
def unit_pending(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT i.*, u.name AS unit_name
        FROM invoices i
        JOIN contracts c ON i.contract_id = c.contract_id
        JOIN tenants t ON c.tenant_id = t.tenant_id
        JOIN unit u ON c.room_id = u.unit_id
        WHERE i.contract_id = %s
        ORDER BY invoice_id DESC
    """, (contract_id,))
    invoices = cursor.fetchall()

    return render_template(
        "unit_invoices_pending.html",
        invoices=invoices
    )

@app.route('/view_slip/<filename>')
def view_slip(filename):
    return send_from_directory(INCOME_UPLOAD_PATH, filename)

@app.route('/update_invoice', methods=['POST'])
def update_invoice():
    # รับค่าและจัดการประเภทข้อมูลให้เรียบร้อยก่อนเข้า DB
    d = request.form
    invoice_id = d.get('invoice_id')
    prev_elec = float(d.get('prev_elec') or 0)
    curr_elec = float(d.get('curr_elec') or 0)
    prev_water = float(d.get('prev_water') or 0)
    curr_water = float(d.get('curr_water') or 0)
    
    elec_usage = curr_elec - prev_elec
    water_usage = curr_water - prev_water

    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        sql_update = """
            UPDATE invoices 
            SET billing_period_start = %s, billing_period_end = %s, due_date = %s,
                previous_electricity_reading = %s, current_electricity_reading = %s,
                previous_water_reading = %s, current_water_reading = %s,
                electricity_usage = %s, water_usage = %s
            WHERE invoice_id = %s
        """
        cursor.execute(sql_update, (
            d.get('billing_period_start'), d.get('billing_period_end'), d.get('due_date'),
            prev_elec, curr_elec, prev_water, curr_water, 
            elec_usage, water_usage, invoice_id
        ))

        refresh_invoice_total(cursor, invoice_id)
        check_meter_save(cursor, invoice_id)

        conn.commit()
        
        # ใช้ยอดที่อัปเดตไปแล้วมาโชว์ (ไม่ต้อง SELECT ใหม่ถ้า refresh ทำงานถูกต้อง)
        flash(f"อัปเดตบิล #{invoice_id} สำเร็จ", "success")

    except Exception as e:
        if conn: conn.rollback()
        flash(f"Error: {str(e)}", "danger")
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

    return redirect(request.referrer)

@app.route('/contract/<int:contract_id>/options', methods=['GET', 'POST'])
def contract_options(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        # 1. ดึงข้อมูลสัญญา
        cursor.execute("""
            SELECT c.contract_id, u.name as room_name, u.building, u.floor
            FROM contracts c
            JOIN unit u ON c.room_id = u.unit_id
            WHERE c.contract_id = %s
        """, (contract_id,))
        contract = cursor.fetchone()
        
        if not contract:
            flash("ไม่พบสัญญา", "danger")
            return redirect(url_for('dashboard'))

        if request.method == 'POST':
            new_option_ids = request.form.getlist('options')

            cursor.execute("DELETE FROM contract_option WHERE contract_id = %s", (contract_id,))
            
            # เพิ่ม Option ใหม่เข้าไปในสัญญา
            for opt_id in new_option_ids:
                cursor.execute("INSERT INTO contract_option (contract_id, option_id) VALUES (%s, %s)",
                               (contract_id, opt_id))

            # หมายเหตุ: เราตัดส่วนวนลูป active_invoices และ refresh_invoice_total ออกทั้งหมด
            # เพื่อไม่ให้ไปกระทบกับบิลปัจจุบัน
            conn.commit()
            flash("อัปเดต Option ในสัญญาเรียบร้อยแล้ว (จะมีผลในบิลรอบถัดไป)", "success")
            return redirect(url_for('contract_options', contract_id=contract_id))

        # --- ส่วนของ GET Request ---
        cursor.execute("SELECT * FROM `option` WHERE is_deleted = 0")
        all_options = cursor.fetchall()

        cursor.execute("SELECT option_id FROM contract_option WHERE contract_id = %s", (contract_id,))
        selected_ids = [str(row['option_id']) for row in cursor.fetchall()]

        return render_template(
            "contract_options.html",
            contract=contract,
            all_options=all_options,
            selected_ids=selected_ids
        )

    except Exception as e:
        conn.rollback()
        flash("เกิดข้อผิดพลาด", "danger")
        return redirect(url_for('dashboard'))
    finally:
        cursor.close()
        conn.close()

@app.route('/cancel_move_out/<int:unit_id>', methods=['POST'])
def cancel_move_out(unit_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        # ล้างวันแจ้งย้ายออก และอัปเดตสถานะห้อง
        cursor.execute("""
            UPDATE unit u
            JOIN contracts c ON u.unit_id = c.room_id
            SET c.notice_move_out_date = NULL, u.status_id = 2
            WHERE u.unit_id = %s AND c.status IN (1,2,3)
        """, (unit_id,))
        
        if cursor.rowcount == 0:
            flash("ไม่พบห้องหรือสัญญาที่สามารถล้างวันแจ้งย้ายออกได้", "warning")
        else:
            # อัปเดต invoices ที่เกี่ยวข้องเป็น cancelled
            cursor.execute("""
                UPDATE invoices
                SET status = 'cancelled'
                WHERE unit_id = %s AND status = 'draft'
            """, (unit_id,))
            
            add_audit_log(
                cursor, 
                'UNIT', 
                'Notice to Moveout', 
                f'ยกเลิกเเจ้งย้ายออกห้อง ID : {unit_id}', 
                session.get('user', {}).get('user_id')
            )
            conn.commit()
            flash("ล้างวันแจ้งย้ายออกและยกเลิกบิลเรียบร้อยแล้ว", "success")

    except Exception as e:
        conn.rollback()
        flash(f"เกิดข้อผิดพลาด: {str(e)}", "danger")
    
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('dashboard'))

@app.route('/cancel_contract/<int:contract_id>', methods=['POST'])
def cancel_contract(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        cursor.execute("""
            UPDATE contracts
            SET status = 7
            WHERE contract_id = %s AND status = 1
        """, (contract_id,))
        
        if cursor.rowcount == 0:
            flash("ไม่พบสัญญาที่สามารถยกเลิกได้", "warning")
        else:
            cursor.execute("""
                UPDATE invoices
                SET status = 'cancelled'
                WHERE contract_id = %s AND status = 'draft'
            """, (contract_id,))

            cursor.execute("""
                UPDATE unit u
                JOIN contracts c ON u.unit_id = c.room_id
                SET u.status_id = 1
                WHERE c.contract_id = %s
            """, (contract_id,))

            add_audit_log(
            cursor, 
            'CONTRACT', 
            'Cancel Contract', 
            f'ยกเลิกสัญญา ID : {contract_id}', 
            session.get('user', {}).get('user_id')
            )

            conn.commit()
            flash("ยกเลิกสัญญาและบิลที่เกี่ยวข้องเรียบร้อยแล้ว", "success")

    except Exception as e:
        conn.rollback()
        flash(f"เกิดข้อผิดพลาด: {str(e)}", "danger")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('dashboard'))

@app.route('/add_invoice_item/<int:invoice_id>', methods=['GET', 'POST'])
def add_invoice_item(invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        description = request.form['description']
        quantity = float(request.form['quantity'] or 0)
        unit_price = float(request.form['unit_price'] or 0)
        item_type = request.form['type']
        total_price = unit_price * quantity

        # 1. เพิ่มรายการลง DB
        cursor.execute("""
            INSERT INTO invoice_items (invoice_id, description, unit_price, quantity, total_price, type)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (invoice_id, description, unit_price, quantity, total_price, item_type))

        # 2. เรียกใช้ฟังก์ชันกลาง (คำนวณใหม่ยกชุด)
        refresh_invoice_total(cursor, invoice_id)

        add_audit_log(
            cursor, 
            'INVOICE', 
            'INSERT', 
            f'เพิ่มรายการบิล ID: {invoice_id}', 
            session.get('user', {}).get('user_id')
        )

        conn.commit()
        conn.close()
        
        flash('เพิ่มรายการและคำนวณยอดบิลใหม่แล้ว', 'success')
        return redirect(url_for('add_invoice_item', invoice_id=invoice_id))

    # GET -> แสดงผลปกติ
    cursor.execute("SELECT * FROM invoices WHERE invoice_id = %s", (invoice_id,))
    invoice = cursor.fetchone()
    cursor.execute("SELECT * FROM invoice_items WHERE invoice_id = %s ORDER BY id ASC", (invoice_id,))
    items = cursor.fetchall()
    conn.close()
    return render_template('add_invoice_item.html', invoice=invoice, items=items)

@app.route('/delete_invoice_item/<int:item_id>/<int:invoice_id>', methods=['GET', 'POST'])
def delete_invoice_item(item_id, invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        # 1. ลบรายการ
        cursor.execute("DELETE FROM invoice_items WHERE id=%s AND invoice_id=%s", (item_id, invoice_id))

        # 2. เรียกฟังก์ชันกลาง (ยอดเงินจะหักออกหรือบวกกลับตามเงื่อนไขเป๊ะๆ)
        refresh_invoice_total(cursor, invoice_id)

        add_audit_log(
            cursor, 
            'USER', 
            'DELETE', 
            f'ลบรายการบิล ID: {item_id}', 
            session.get('user', {}).get('user_id')
        )

        conn.commit()
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.method == 'POST':
            return jsonify({"status": "success", "message": "ลบและปรับปรุงยอดเงินแล้ว"})
    except Exception as e:
        conn.rollback()
        return jsonify({"status": "error", "message": str(e)})
    finally:
        conn.close()
    return redirect(url_for('add_invoice_item', invoice_id=invoice_id))

@app.route("/daily_pending/<int:invoice_id>")
def daily_pending(invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT 
            i.invoice_id, i.invoice_type, i.issue_date, 
            i.billing_period_start, i.billing_period_end, i.due_date, 
            i.contract_id, i.total_amount, i.status, i.payment_date, 
            i.late_penalty, u.name AS unit_name,
            i.previous_electricity_reading, i.current_electricity_reading,
            i.previous_water_reading, i.current_water_reading
        FROM invoices i
        JOIN unit u ON i.unit_id = u.unit_id
        WHERE i.invoice_id = %s OR parent_invoice_id = %s
        ORDER BY invoice_id DESC
    """, (invoice_id, invoice_id))
    
    invoices = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        "daily_invoices_pending.html",
        invoices=invoices,
    )

@app.route('/finance')
def finance_page():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    # 1. เตรียมข้อมูลปี (ย้อนหลัง 5 ปี)
    current_year = datetime.now().year
    years_list = [y for y in range(current_year, current_year - 6, -1)]

    # 2. **ต้องเพิ่มสิ่งนี้** คือรายชื่อเดือน เพื่อให้ Dropdown ใน HTML มีข้อมูล
    months_list = [
        {'id': 'all', 'name': 'ทุกเดือน'},
        {'id': 1, 'name': 'มกราคม'}, {'id': 2, 'name': 'กุมภาพันธ์'},
        {'id': 3, 'name': 'มีนาคม'}, {'id': 4, 'name': 'เมษายน'},
        {'id': 5, 'name': 'พฤษภาคม'}, {'id': 6, 'name': 'มิถุนายน'},
        {'id': 7, 'name': 'กรกฎาคม'}, {'id': 8, 'name': 'สิงหาคม'},
        {'id': 9, 'name': 'กันยายน'}, {'id': 10, 'name': 'ตุลาคม'},
        {'id': 11, 'name': 'พฤศจิกายน'}, {'id': 12, 'name': 'ธันวาคม'}
    ]

    # 3. รับค่าจาก URL (ถ้ามี) เพื่อกำหนดค่าเริ่มต้นให้ Dropdown
    selected_year = request.args.get('year', default=current_year, type=int)
    selected_month = request.args.get('month', default='all')

    # 4. ดึงรายการล่าสุด (แสดงผลครั้งแรกตอนโหลดหน้า)
    cursor.execute("SELECT * FROM transactions ORDER BY transaction_date DESC LIMIT 15")
    transactions = cursor.fetchall()

    cursor.close()
    conn.close()

    # 5. ส่งค่าทั้งหมดกลับไปที่ Template
    return render_template('finance.html', 
                           years_list=years_list, 
                           months_list=months_list, 
                           selected_year=selected_year,
                           selected_month=selected_month,
                           transactions=transactions)

@app.route('/read_meter_unit', methods=['POST'])
def read_meter_unit():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    room_id = request.form.get('room_id') 
    read_type = request.form.get('read_type') # 'electricity' หรือ 'water'
    
    try:
        # 1. เลือกตารางและดึงข้อมูลมิเตอร์รายห้อง
        if read_type == 'electricity':
            table_name = "meter"
            is_water = False
        else:
            table_name = "meter_water"
            is_water = True

        # ดึงค่า configuration ทั้งหมดของมิเตอร์นั้นๆ
        sql = f"SELECT slave_id, module, unit_id, comport, ip, port, base_url, api_auth_token, status, unit_key FROM {table_name} WHERE unit_id = %s"
        cursor.execute(sql, [room_id])
        meter_data = cursor.fetchone()

        if not meter_data:
            return jsonify({"status": "error", "message": "ไม่พบข้อมูลมิเตอร์ในระบบ"})
        
        if meter_data['status'] == 'inactive':
            return jsonify({"status": "error", "message": "มิเตอร์ไม่ได้เปิดใช้งาน"})

        # 2. เรียกฟังก์ชันอ่านค่า (ส่งทั้งข้อมูล Hardware และ API Config เข้าไป)
        val = read_meter_unit_read(
            model_name=meter_data['module'],
            register_key=meter_data['unit_key'],   
            serial_ports=meter_data['comport'],
            ip=meter_data['ip'],                  
            port=meter_data['port'],              
            slave_id=meter_data['slave_id'],
            is_water=is_water,
            api_base_url=meter_data['base_url'],           
            api_token=meter_data['api_auth_token']                 
        )

        # ดึง user_id จาก session สำหรับการทำ Log
        user_id = session.get('user', {}).get('user_id', 0)

        # 3. ตรวจสอบผลลัพธ์การอ่าน
        if val is None or "Error" in str(val) or "Cannot connect" in str(val) or "Exception" in str(val):
            return jsonify({"status": "error", "message": str(val) or "ติดต่อมิเตอร์ไม่ได้"})

        # 🚀 4. อัปเดตค่าลง Database (unit, meter_table, log)
        if is_water:
            cursor.execute("UPDATE unit SET water_start = %s WHERE unit_id = %s", (val, room_id))
            cursor.execute("UPDATE meter_water SET updated_at = NOW() WHERE unit_id = %s", (room_id,)) 
            log_meter_reading(cursor, room_id, 'water', val, source='manual', created_by=user_id)
        else:
            cursor.execute("UPDATE unit SET electricity_start = %s WHERE unit_id = %s", (val, room_id))
            cursor.execute("UPDATE meter SET updated_at = NOW() WHERE unit_id = %s", (room_id,)) 
            log_meter_reading(cursor, room_id, 'electricity', val, source='manual', created_by=user_id)

        conn.commit()

        return jsonify({
            "status": "success", 
            "value": val, 
            "target_id": meter_data['unit_key'], # ส่งชื่อ key กลับไปให้ JS อัปเดตช่อง input
            "unit": "m3" if is_water else "kWh",
            "message": "อ่านค่าและบันทึกเรียบร้อย"
        })

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"status": "error", "message": str(e)})
    finally:
        cursor.close()
        conn.close()

@app.route('/api/sync-meters', methods=['POST'])
def sync_meters_route():
    try:
        # แก้ไขบรรทัดนี้ให้รับ 3 ค่าตามที่ฟังก์ชันส่งมา
        success, message, warnings = action_sync_latest_meter_to_invoices()
        
        if success:
            return jsonify({
                "status": "success", 
                "message": message,
                "warnings": warnings,  # ส่งรายการคำเตือนกลับไปด้วย
                "updated_count": message
            }), 200
        else:
            return jsonify({
                "status": "error", 
                "message": f"ฐานข้อมูลทำงานผิดพลาด: {message}"
            }), 500
    except Exception as e:
        return jsonify({
            "status": "error", 
            "message": str(e)
        }), 500

@app.route('/api/usage-dashboard', methods=['GET'])
def get_usage_dashboard():
    unit_id = request.args.get('unit_id', default='all')
    year = request.args.get('year', default=2024, type=int)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    try:
        where_clause = "WHERE YEAR(billing_period_start) = %s AND status != 'draft'"
        params = [year]
        
        if unit_id != 'all':
            where_clause += " AND unit_id = %s"
            params.append(unit_id)

        # SQL ที่แก้ปัญหา only_full_group_by
        summary_sql = f"""
            SELECT 
                DATE_FORMAT(billing_period_start, '%b %Y') as label,
                SUM(current_electricity_reading - previous_electricity_reading) as el_total,
                SUM(current_water_reading - previous_water_reading) as wt_total
            FROM invoices 
            {where_clause}
            GROUP BY label, YEAR(billing_period_start), MONTH(billing_period_start)
            ORDER BY YEAR(billing_period_start) ASC, MONTH(billing_period_start) ASC
        """
        cursor.execute(summary_sql, tuple(params))
        summary_res = cursor.fetchall()

        # ส่วน Ranking (อันนี้ไม่มีปัญหา GROUP BY เพราะเราใช้แค่ unit_id)
        ranking_sql = """
            SELECT unit_id, SUM(current_electricity_reading - previous_electricity_reading) as el_usage
            FROM invoices 
            WHERE YEAR(billing_period_start) = %s AND status != 'draft'
            GROUP BY unit_id 
            ORDER BY el_usage DESC 
            LIMIT 10
        """
        cursor.execute(ranking_sql, (year,))
        ranking_res = cursor.fetchall()

        return jsonify({
            "status": "success",
            "summary": {
                "labels": [r['label'] for r in summary_res],
                "el": [round(max(0, float(r['el_total'] or 0)), 2) for r in summary_res],
                "wt": [round(max(0, float(r['wt_total'] or 0)), 2) for r in summary_res]
            },
            "ranking": {
                "units": [r['unit_id'] for r in ranking_res],
                "el": [round(max(0, float(r['el_usage'] or 0)), 2) for r in ranking_res]
            }
        })
    except Exception as e:
        print(f"Error detail: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/audit_logs')
@role_required(['admin', 'manager'])
def view_audit_logs():
    selected_category = request.args.get('category', '')
    start_date = request.args.get('start_date', '')
    end_date = request.args.get('end_date', '')
    
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT DISTINCT category FROM audit_log")
    categories = [row['category'] for row in cursor.fetchall()]

    query = """
        SELECT al.*, u.username as staff_name, r.r_name, u.user_id as staff_id
        FROM audit_log al
        LEFT JOIN user u ON al.created_by = u.user_id
        LEFT JOIN role r ON u.user_id = r.role_id
        WHERE 1=1
    """
    params = []

    # ฟิลเตอร์ตามหมวดหมู่
    if selected_category:
        query += " AND al.category = %s"
        params.append(selected_category)

    # ฟิลเตอร์ตามวันที่
    if start_date:
        query += " AND DATE(al.created_at) >= %s"
        params.append(start_date)

    if end_date:
        query += " AND DATE(al.created_at) <= %s"
        params.append(end_date)
    
    query += " ORDER BY al.created_at DESC"

    # --- ส่วนที่เพิ่ม/แก้ไข: เช็คเงื่อนไขการ LIMIT ---
    # ถ้าไม่ได้เลือกวันที่เริ่มต้น และ ไม่ได้เลือกวันที่สิ้นสุด ให้แสดงแค่ 50 รายการล่าสุด
    if not start_date and not end_date:
        query += " LIMIT 50"
    # ------------------------------------------
    
    cursor.execute(query, tuple(params))
    logs = cursor.fetchall()
    conn.close()
    
    return render_template('audit_logs.html', 
                           logs=logs, 
                           categories=categories, 
                           selected_category=selected_category,
                           start_date=start_date,
                           end_date=end_date)

# ---------------------- DASHBOARD METER ----------------------
@app.route('/meter_analysis', methods=['GET', 'POST'])
def meter_analysis():
    current_year = datetime.now().year
    
    # 2. รับค่าจาก Filter (ใช้ default เป็นค่าว่างเพื่อตรวจสอบการกดล้างค่า)
    year = request.args.get('year', default=current_year, type=int)
    start_date = request.args.get('start_date', default='')
    end_date = request.args.get('end_date', default='')
    unit_id = request.args.get('unit_id', default='')
    
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # --- ส่วนที่ 1: กราฟ (ดึงข้อมูลรายเดือนตามปีที่เลือก) ---
    sql_usage = """
        SELECT m.month, 
            ROUND(COALESCE(SUM(i.current_electricity_reading - i.previous_electricity_reading), 0), 2) as electricity,
            ROUND(COALESCE(SUM(i.current_water_reading - i.previous_water_reading), 0), 2) as water
        FROM (
            SELECT 1 AS month UNION SELECT 2 UNION SELECT 3 UNION SELECT 4 
            UNION SELECT 5 UNION SELECT 6 UNION SELECT 7 UNION SELECT 8 
            UNION SELECT 9 UNION SELECT 10 UNION SELECT 11 UNION SELECT 12
        ) AS m
        LEFT JOIN invoices i ON m.month = MONTH(i.billing_period_start) 
            AND YEAR(i.billing_period_start) = %s
            AND i.status != 'cancelled'
        GROUP BY m.month ORDER BY m.month
    """
    cursor.execute(sql_usage, (year,))
    usage_res = cursor.fetchall()
    
    # --- ส่วนที่ 2: นับ Active Meters ---
    cursor.execute("SELECT COUNT(DISTINCT unit_id) as active_count FROM invoices WHERE YEAR(billing_period_start) = %s", (year,))
    active_meters = cursor.fetchone()['active_count'] or 0

    # --- ส่วนที่ 3: Log (ดึงตามช่วงวันที่ หรือ 100 รายการล่าสุด) ---
    log_sql = """
        SELECT u.name, mr.serial_meter, mr.meter_type, mr.meter_id, mr.source, usr.username, mr.created_by, mr.invoice_id,
               COALESCE(mr.current_reading, 0) as current_reading, 
               mr.read_date 
        FROM meter_reading mr
        INNER JOIN unit u ON mr.unit_id = u.unit_id
        LEFT JOIN user usr ON mr.created_by = usr.user_id
    """
    params = []

    if start_date and end_date:
        # กรณีมีการเลือกช่วงวันที่
        log_sql += " WHERE DATE(mr.read_date) BETWEEN %s AND %s"
        params.extend([start_date, end_date])
        if unit_id:
            log_sql += " AND mr.unit_id = %s"
            params.append(unit_id)
        log_sql += " ORDER BY mr.read_date DESC"
    else:
        # กรณีไม่ได้เลือกวันที่ หรือกด Clear ให้โชว์ 100 อันล่าสุด
        if unit_id:
            log_sql += " WHERE mr.unit_id = %s"
            params.append(unit_id)
        log_sql += " ORDER BY mr.read_date DESC LIMIT 100"

    cursor.execute(log_sql, tuple(params))
    recent_logs = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template('meter_analysis.html', 
                           year=year, 
                           start_date=start_date, 
                           end_date=end_date,
                           selected_unit=unit_id,
                           usage_data=usage_res, 
                           active_meters=active_meters,
                           logs=recent_logs)

# ---------------------- RENEW -----------------------
@app.route('/contracts_renew/<int:contract_id>', methods=['GET', 'POST'])
def renew_contracts(contract_id):

    today = get_now().date() if 'get_now' in globals() else date.today()
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # ตรวจสอบยอดค้างชำระ
    cursor.execute("""
        SELECT COUNT(*) as unpaid_count 
        FROM invoices 
        WHERE contract_id = %s AND status IN ('overdue')
    """, (contract_id,))
    if cursor.fetchone()['unpaid_count'] > 0:
        flash("⚠️ ไม่สามารถต่อสัญญาได้: กรุณาเคลียร์บิลค้างชำระก่อน", "danger")
        return redirect(url_for('dashboard'))

    # 1. ดึงข้อมูลสัญญาเดิม
    cursor.execute("""
        SELECT c.*, u.name as room_name, u.unit_id 
        FROM contracts c 
        JOIN unit u ON c.room_id = u.unit_id 
        WHERE c.contract_id = %s AND c.status = 4
    """, (contract_id,))
    contract = cursor.fetchone()

    if not contract:
        flash("ไม่พบสัญญาที่รอการต่ออายุ", "warning")
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        new_end = request.form.get('contract_end') 
        files = request.files.getlist('files[]')

        if not new_end:
            flash("กรุณาระบุวันสิ้นสุดสัญญาใหม่", "warning")
            return redirect(request.url)

        # 2. ตั้งค่าเวอร์ชันและชื่อไฟล์ใหม่
        current_v = (contract['renew_count'] or 0) + 1
        today_str = today.strftime("%d%m%Y")
        room_name = contract['room_name']
        base_filename = f"contract_{contract_id}_{today_str}_{room_name}_v{current_v}"
        
        saved_filenames = []
        target_folder = os.path.join(BASE_DIR, 'contracts_file')
        os.makedirs(target_folder, exist_ok=True)

        # 3. บันทึกไฟล์ใหม่ลง Server
        for i, file in enumerate(files):
            if file and file.filename != '':
                ext = os.path.splitext(file.filename)[1].lower()
                if ext[1:] in ALLOWED_EXTENSIONS:
                    full_filename = f"{base_filename}_{i+1}{ext}"
                    file.save(os.path.join(target_folder, full_filename))
                    saved_filenames.append(full_filename)

        # 4. จัดการชื่อไฟล์ใน Database (อัพเดตทับของเดิมตามที่คุณต้องการ)
        new_files_str = ",".join(saved_filenames)
        # ถ้าไม่อัพโหลดไฟล์ใหม่ ให้ใช้ค่าเดิมใน DB เพื่อป้องกันข้อมูลหาย
        final_files_in_db = new_files_str if new_files_str else contract['contracts_file']

        try:
            # 5. อัปเดตข้อมูลสัญญาเป็น Status 3 (Active)
            cursor.execute("""
                UPDATE contracts 
                SET contract_end = %s, 
                    status = 3, 
                    renew_count = %s, 
                    contracts_file = %s
                WHERE contract_id = %s
            """, (new_end, current_v, final_files_in_db, contract_id))

            # เคลียร์บิล Final เก่า (ถ้ามี)
            cursor.execute("UPDATE invoices SET status = 'cancelled' WHERE contract_id = %s AND invoice_type = 'final' AND status = 'draft'", (contract_id,))
            cursor.execute("UPDATE unit SET status_id = 2 WHERE unit_id = %s", (contract['unit_id'],))

            conn.commit() # สำคัญ: ต้อง Commit ก่อนสร้างบิล

            # 6. สร้างบิลรอบใหม่ทันที
            create_monthly_invoice(
                cursor=cursor,
                unit_id=contract['unit_id'], 
                billing_month=today, 
                created_by=session.get('user_id')
            )

            add_audit_log(
                cursor, 
                'CONTRACT', 
                'RENEW', 
                f'ต่อสัญญาเวอร์ชัน {current_v} สำเร็จ', 
                session.get('user', {}).get('user_id')
            )

            flash(f"✅ ต่อสัญญาเวอร์ชัน {current_v} เรียบร้อย", "success")
        except Exception as e:
            conn.rollback()
            flash(f"❌ Error: {str(e)}", "danger")
        finally:
            cursor.close()
            conn.close()
        return redirect(url_for('dashboard'))

    return render_template("renew_contracts.html", contract=contract)


# ---------------------- TENANT ----------------------
@app.route('/add_tenant/<int:unit_id>', methods=['GET', 'POST'])
def add_tenant(unit_id):

    if request.method == 'POST':
        fname = request.form.get('fname')
        lname = request.form.get('lname')
        id_card = request.form.get('id_card')
        phone = request.form.get('phone')
        address = request.form.get('address')
        email = request.form.get('email')
        gender = request.form.get('gender')
        bd_str = request.form.get('bd')

        try:
            bd = datetime.strptime(bd_str, '%Y-%m-%d').date()
        except ValueError:
            flash('รูปแบบวันเกิดไม่ถูกต้อง', 'danger')
            return redirect(url_for('add_tenant', unit_id=unit_id))

        today = date.today()
        age = today.year - bd.year - \
            ((today.month, today.day) < (bd.month, bd.day))

        if not id_card or not id_card.isdigit() or len(id_card) != 13:
            flash('กรุณากรอกเลขบัตรประชาชน 13 หลักให้ถูกต้อง', 'danger')
            return redirect(url_for('add_tenant', unit_id=unit_id))

        conn = get_db_connection()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO tenants (fname, lname, id_card, gender, bd, age, tel, address, email)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (fname, lname, id_card, gender, bd, age, phone, address, email))
            conn.commit()

            add_audit_log(
                cursor, 
                'TENANTS', 
                'INSERT', 
                f'เพิ่มผู้เช่าใหม่', 
                session.get('user', {}).get('user_id')
            )

            flash('เพิ่มผู้เช่าสำเร็จ', 'success')
            return redirect(url_for('create_lease', unit_id=unit_id))
        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {e}', 'danger')
        finally:
            cursor.close()
            conn.close()

    return render_template('add_tenant.html', unit_id=unit_id)


# --------------------- INVOICE ----------------------
@app.route("/unit_invoices/<int:contract_id>", methods=["GET", "POST"])
def unit_invoices(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # ดึงข้อมูล unit ของ contract
    cursor.execute("""
        SELECT u.* , c.contract_id
        FROM contracts AS c
        JOIN unit AS u ON c.room_id = u.unit_id
        WHERE c.contract_id = %s
    """, (contract_id,))
    unit = cursor.fetchone()

    if not unit:
        return "Unit not found", 404

    # ดึง invoices ของ contract
    cursor.execute("""
        SELECT i.invoice_id, i.issue_date, i.billing_period_start, i.billing_period_end, i.due_date, i.total_amount, i.status, i.payment_date, i.late_penalty
        FROM invoices i
        JOIN contracts c ON i.contract_id = c.contract_id
        JOIN tenants t ON c.tenant_id = t.tenant_id
        WHERE i.contract_id = %s
        ORDER BY invoice_id DESC
    """, (contract_id,))
    invoices = cursor.fetchall()

    # สรุปยอด
    total_all = sum(inv["total_amount"] for inv in invoices)
    total_paid = sum(inv["total_amount"] for inv in invoices if inv["status"] == "paid")
    total_unpaid = sum(inv["total_amount"] for inv in invoices if inv["status"] == "unpaid")
    total_overdue = sum(inv["total_amount"] for inv in invoices if inv["status"] == "overdue")

    return render_template(
        "invoice_contracts.html",
        unit=unit,
        contract_id=contract_id,
        invoices=invoices,
        total_all=total_all,
        total_paid=total_paid,
        total_unpaid=total_unpaid,
        total_overdue=total_overdue
    )

@app.route('/cancel_invoice/<int:invoice_id>', methods=['POST'])
def cancel_invoice(invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        # ตรวจสอบสถานะก่อนอัปเดต
        cursor.execute("SELECT * FROM invoices WHERE invoice_id = %s", (invoice_id,))
        inv = cursor.fetchone()
        
        if not inv:
            return jsonify({"status": "error", "message": "ไม่พบข้อมูลบิล"})
        
        if inv.get('slip_file'):
            file_to_delete = os.path.join(INCOME_UPLOAD_PATH, inv['slip_file'])
            try:
                if os.path.exists(file_to_delete):
                    os.remove(file_to_delete)
                    print(f"ลบไฟล์สำเร็จ: {file_to_delete}")
                else:
                    print(f"ไม่พบไฟล์บน Server: {file_to_delete}")
            except Exception as e:
                print(f"เกิดข้อผิดพลาดขณะลบไฟล์: {e}")
            
        cursor.execute("""
            UPDATE invoices 
            SET status = 'cancelled' , slip_file = NULL
            WHERE invoice_id = %s
        """, (invoice_id,))


        cursor.execute("""
            DELETE FROM transactions 
            WHERE ref_invoice_id = %s AND type = 'income'
        """, (invoice_id,))

        add_audit_log(
            cursor, 
            'INVOICE', 
            'CANCEL', 
            f'ยกเลิกบิล ID: {invoice_id}', 
            session.get('user', {}).get('user_id')
        )
        
        conn.commit()
        return jsonify({"status": "success", "message": "ยกเลิกบิลเรียบร้อยแล้ว"})
        
    except Exception as e:
        if conn: conn.rollback()
        return jsonify({"status": "error", "message": str(e)})
    finally:
        if conn: conn.close()

@app.route("/create_invoice_extra_bill/<int:contract_id>", methods=["GET", "POST"])
def create_invoice_extra_bill(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # 1. ดึงข้อมูลสัญญาและห้อง
    cursor.execute("""
        SELECT c.contract_id, c.tenant_id, c.room_id, u.name as room_name
        FROM contracts c 
        JOIN unit u ON u.unit_id = c.room_id
        WHERE c.contract_id = %s
    """, (contract_id,))
    contract = cursor.fetchone()

    if not contract:
        conn.close()
        return "Contract not found", 404

    if request.method == 'POST':
        try:
            due_date = request.form.get('due_date')
            total_amount = float(request.form.get('total_amount', 0))
            
            if conn.in_transaction:
                conn.rollback()

            # เริ่ม Transaction
            conn.start_transaction()

            # 2. INSERT ลงตาราง invoices (ตารางแม่)
            insert_invoice_query = """
                INSERT INTO invoices (
                    unit_id, tenant_id, contract_id, invoice_type, 
                    issue_date, due_date, rent_amount, electricity_total, water_total, total_amount, status, created_by, created_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s,  %s, %s, %s,%s, NOW())
            """
            invoice_values = (
                contract['room_id'], 
                contract['tenant_id'], 
                contract_id, 
                'extra_bill', 
                datetime.now().date(), 
                due_date,
                0,0,0,
                total_amount, 
                'draft',
                session.get('user', {}).get('user_id')
                
            )
            cursor.execute(insert_invoice_query, invoice_values)
            new_invoice_id = cursor.lastrowid
            
            item_index = 0
            while True:
                desc = request.form.get(f'items[{item_index}][description]')
                if not desc: break # ถ้าไม่มีรายการแล้วให้หยุดลูป
                
                amount = float(request.form.get(f'items[{item_index}][amount]', 0))
                item_type = request.form.get(f'items[{item_index}][type]', 'other')

                insert_item_query = """
                    INSERT INTO invoice_items (
                        invoice_id, description, unit_price, quantity, total_price, type
                    ) VALUES (%s, %s, %s, %s, %s, %s)
                """
                cursor.execute(insert_item_query, (new_invoice_id, desc, amount, 1, amount, item_type))
                item_index += 1

            add_audit_log(
                cursor, 
                'INVOICE', 
                'CREATE', 
                f'สร้างบิลใหม่พิเศษ ID: {new_invoice_id}', 
                session.get('user', {}).get('user_id')
            )

            conn.commit()
            return redirect(url_for('unit_pending', contract_id=contract_id))

        except Exception as e:
            conn.rollback()
            print(f"Error: {e}")
            return f"เกิดข้อผิดพลาด: {str(e)}", 500
        finally:
            cursor.close()
            conn.close()

    return render_template("create_invoice_extra_bill.html", 
                           contract_id=contract_id, 
                           room_name=contract['room_name'],
                           today_date=datetime.now().strftime('%Y-%m-%d'))

@app.route('/manual_create_bill/<int:contract_id>', methods=['GET', 'POST'])
def manual_create_bill(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        cursor.execute("""
            SELECT c.*, u.unit_id, u.name 
            FROM contracts c 
            JOIN unit u ON c.room_id = u.unit_id 
            WHERE c.contract_id = %s
        """, (contract_id,))
        contract = cursor.fetchone()

        if not contract:
            return "ไม่พบสัญญา", 404
        
        e_rate = float(get_setting('electricity_rate', 7))
        w_rate = float(get_setting('water_rate', 18))

        # ดึง Options จากสัญญา (เพื่อให้หน้า GET แสดงรายการเริ่มต้น และหน้า POST ใช้ตรวจสอบ)
        cursor.execute("""
            SELECT o.name, o.price 
            FROM contract_option co
            JOIN `option` o ON co.option_id = o.id
            WHERE co.contract_id = %s
        """, (contract_id,))
        default_options = cursor.fetchall()
        default_options_total = sum(float(opt['price']) for opt in default_options)

        # --- ส่วนที่ 2: ถ้ามีการส่งข้อมูลมา (POST) ---
        if request.method == 'POST':
            inv_type = request.form.get('invoice_type')
            b_start = request.form.get('billing_start')
            b_end = request.form.get('billing_end')
            
            def f(val): return float(val) if val else 0.0
            
            p_e = f(request.form.get('prev_elec'))
            c_e = f(request.form.get('curr_elec'))
            e_total = f(request.form.get('e_total'))
            
            p_w = f(request.form.get('prev_water'))
            c_w = f(request.form.get('curr_water'))
            w_total = f(request.form.get('w_total'))
            
            rent = f(request.form.get('rent_amount'))
            others = f(request.form.get('other_charges'))
            
            # รับรายการ Invoice Items (ทั้งที่แก้จาก Option เดิม และที่เพิ่มใหม่)
            item_names = request.form.getlist('item_name[]')
            item_amounts = request.form.getlist('item_amount[]')
            items_total = sum(f(amt) for amt in item_amounts)
            
            # Logic คืนประกัน
            reimburse = f(request.form.get('reimburse'))

            if inv_type == 'first':
                # คำนวณยอดสุทธิ: ค่าเช่า + ไฟ + น้ำ + อื่นๆ + รายการเสริม + (คืนประกัน)
                total_amount = rent + e_total + w_total + others + items_total - reimburse

            else:
                total_amount = rent + e_total + w_total + others + items_total

            # 2.1 บันทึกลงตาราง invoices
            sql_inv = """
                INSERT INTO invoices (
                    unit_id, tenant_id, contract_id, invoice_type, 
                    billing_period_start, billing_period_end, issue_date, due_date,
                    rent_amount, previous_electricity_reading, current_electricity_reading, 
                    electricity_total, previous_water_reading, current_water_reading, 
                    water_total, other_charges, reimburse, total_amount, status,meter_saved,created_at,premiums
                ) VALUES (%s, %s, %s, %s, %s, %s, NOW(), %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'draft',1,NOW(),0)
            """
            cursor.execute(sql_inv, (
                contract['room_id'],contract['tenant_id'],contract_id,inv_type,                     
                b_start,b_end,b_end,rent,p_e,c_e,e_total,p_w,                           
                c_w,w_total,others,reimburse,total_amount                  
            ))
            invoice_id = cursor.lastrowid

            add_audit_log(
                cursor, 
                'INVOICE', 
                'CREATE', 
                f'สร้างใบแจ้งหนี้ ID: {invoice_id}', 
                session.get('user', {}).get('user_id')
            )

            # 2.2 บันทึกลงตาราง invoice_items (วนลูปจากที่ส่งมาจากหน้าเว็บ)
            sql_item = """
                INSERT INTO invoice_items 
                (invoice_id, description, unit_price, quantity, total_price, option_id, type) 
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """

            if item_names:
                for name, amt in zip(item_names, item_amounts):
                    if name.strip():
                        price = float(amt) if amt else 0.0
                        # บันทึกเป็น service ตามโจทย์
                        cursor.execute(sql_item, (
                            invoice_id, name, price, 1, price, None, 'service'
                        ))
            
            conn.commit()
            return redirect(url_for('unit_pending', contract_id=contract_id))

        # --- ส่วนที่ 3: แสดงผลหน้าฟอร์ม (GET) ---
        # ดึงค่ามิเตอร์ล่าสุด
        cursor.execute("""
            SELECT current_electricity_reading, current_water_reading 
            FROM invoices WHERE contract_id = %s AND status = 'paid'
            ORDER BY billing_period_end DESC LIMIT 1
        """, (contract_id,))
        last_bill = cursor.fetchone()

        return render_template(
            "manual_bill_form.html",
            contract=contract,
            last_bill=last_bill,
            e_rate=e_rate,
            w_rate=w_rate,
            default_options=default_options,
            default_options_total=default_options_total,
            today=date.today().isoformat()
        )

    except Exception as e:
        if conn: conn.rollback()
        return f"Error: {str(e)}", 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()


# ---------------------- LEASE ----------------------
@app.route('/create_lease/<int:unit_id>', methods=['GET', 'POST'])
def create_lease(unit_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""SELECT setting_value FROM settings WHERE setting_key='min_lease_months'""")
    result = cursor.fetchone()
    try:
        min_months = int(result['setting_value']) if result and result['setting_value'] else 0
    except ValueError:
        min_months = 0

    if request.method == 'POST':
        try:
            tenant_id = int(request.form.get('tenant_id'))
            lease_start = datetime.strptime(request.form['lease_start'], '%Y-%m-%d')
            lease_end   = datetime.strptime(request.form['lease_end'], '%Y-%m-%d')
            premium = float(request.form.get('premium') or 0)
            selected_options = request.form.getlist('options')
            amount = int(request.form.get('amount') or 1)
            pay_date    = datetime.strptime(request.form['pay_date'], '%Y-%m-%d')

            # ดึงค่าเริ่มต้นมิเตอร์
            cursor.execute("""
                SELECT u.*, t.price_monthly, u.electricity_start, u.water_start
                FROM unit u
                LEFT JOIN type t ON u.type_unit_id = t.type_id
                WHERE u.unit_id=%s AND u.is_deleted = 0
            """, (unit_id,))
            room = cursor.fetchone()
            if not room:
                flash("ไม่พบข้อมูลห้อง", "danger")
                return redirect(url_for('dashboard'))
            
            elif room['meter_id'] is None or room['meter_water_id'] is None:
                flash("ไม่พบข้อมูลมิเตอร์ (กรุณาผูกมิเตอร์ก่อนสร้างสัญญา)", "danger")
                return redirect(url_for('dashboard'))
            
            room_price = float(room['price_monthly'] or 0)
            electricity_start = float(room['electricity_start'] or 0)
            water_start = float(room['water_start'] or 0)

            # ดึงเรทราคา
            cursor.execute("""
                SELECT setting_key, setting_value
                FROM settings
                WHERE setting_key IN ('electricity_rate', 'water_rate')
            """)
            rows = cursor.fetchall()
            rates = {row['setting_key']: float(
                row['setting_value']) for row in rows}
            electricity_rate = rates.get('electricity_rate', 0)
            water_rate = rates.get('water_rate', 0)

            # ตรวจสอบสัญญาเดิม
            cursor.execute("""
                SELECT * FROM contracts
                WHERE room_id=%s AND status IN (1,2,3)
            """, (unit_id,))
            if cursor.fetchone():
                flash('ห้องนี้มีสัญญายังไม่สิ้นสุด', 'danger')
                return redirect(url_for('dashboard'))

            # ✅ สร้างสัญญา
            cursor.execute("""
                INSERT INTO contracts (tenant_id, room_id, contract_start, contract_end,
                                        price, premiums, status,
                                        amount, pay_date, electricity_start, water_start)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (tenant_id, unit_id, lease_start, lease_end,
                room_price, premium, 1, amount, pay_date, electricity_start, water_start))
            contract_id = cursor.lastrowid

            add_audit_log(
                cursor, 
                'CONTRACT', 
                'CREATE', 
                f'สร้างสัญญา ID: {contract_id}', 
                session.get('user', {}).get('user_id')
            )

            # ✅ บันทึก options ลง contract_option + รวมราคา
            if selected_options:
                for opt_id in selected_options:
                    cursor.execute(
                        "INSERT INTO contract_option (contract_id, option_id) VALUES (%s,%s)",
                        (contract_id, opt_id)
                    )
                # รวมราคา option
                placeholders = ",".join(["%s"] * len(selected_options))
                sql = f"SELECT SUM(price) AS total FROM `option` WHERE id IN ({placeholders})"
                cursor.execute(sql, tuple(selected_options))
                opt_sum = cursor.fetchone()
                service_charge = float(opt_sum['total'] or 0)

            # ✅ สร้าง invoice รอบแรก
            first_day = lease_start.date()
            last_day = (first_day.replace(day=1) +
                        relativedelta(months=1)) - timedelta(days=1)

            cursor.execute("""
                INSERT INTO invoices (
                    unit_id, tenant_id, contract_id,
                    invoice_type,
                    billing_period_start, billing_period_end,
                    issue_date, due_date,
                    rent_amount, premiums, service_charge,
                    total_amount, status, created_by, created_at,
                    previous_electricity_reading, previous_water_reading,
                    electricity_usage, electricity_rate, electricity_total,
                    water_usage, water_rate, water_total
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,NOW(),
                        %s,%s,%s,%s,%s,%s,%s,%s)
            """, (
                unit_id, tenant_id, contract_id,
                'first',
                first_day, last_day,
                first_day, pay_date,
                0, premium, 0,
                premium, 'draft', session['user']['user_id'],
                electricity_start, water_start,
                electricity_start, electricity_rate, 0,
                water_start, water_rate, 0
            ))
    
            # อัปเดตสถานะห้อง
            cursor.execute(
                "UPDATE unit SET status_id=3 WHERE unit_id=%s", (unit_id,))

            conn.commit()
            flash('สร้างสัญญาและบิลรอบแรกเรียบร้อยแล้ว', 'success')
            return redirect(url_for('dashboard'))

        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {e}', 'danger')
            return redirect(request.url)
        finally:
            cursor.close()
            conn.close()

    # GET method
    cursor.execute("""
        SELECT u.*, t.price_monthly
        FROM unit u
        LEFT JOIN type t ON u.type_unit_id = t.type_id
        WHERE u.unit_id=%s
    """, (unit_id,))
    room = cursor.fetchone()
    if not room:
        flash("ไม่พบข้อมูลห้อง", "danger")
        return redirect(url_for('dashboard'))
    elif room['meter_id'] is None or room['meter_water_id'] is None:
        flash("ไม่พบข้อมูลมิเตอร์ (กรุณาผูกมิเตอร์ก่อนสร้างสัญญา)", "danger")
        return redirect(url_for('dashboard'))
    
    room_price = float(room['price_monthly'] or 0) if room else 0

    cursor.execute("""
       SELECT t.* FROM tenants t
        WHERE t.is_deleted = 0 
        AND NOT EXISTS (
            SELECT 1 
            FROM contracts c 
            WHERE c.tenant_id = t.tenant_id 
            AND c.status IN (1, 2, 3, 4)
        )
        ORDER BY t.fname;
    """)
    tenants = cursor.fetchall()

    cursor.execute("SELECT * FROM `option` WHERE is_deleted = 0 ORDER BY name")
    options = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'create_lease.html',
        room=room,
        tenants=tenants,
        options=options,
        room_price=room_price,
        min_months=min_months
    )


# ---------------------- DAILY BOOKING ----------------------
@app.route('/daily_booking/<int:unit_id>', methods=['GET', 'POST'])
def daily_booking(unit_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # ดึงข้อมูลห้องและราคาต่อวัน
    cursor.execute("""
        SELECT u.*, t.price_daily, u.electricity_start, u.water_start
        FROM unit u
        JOIN type t ON u.type_unit_id = t.type_id
        WHERE u.unit_id = %s
    """, (unit_id,))
    room = cursor.fetchone()
    if not room:
        flash("ไม่พบข้อมูลห้อง", "danger")
        return redirect(url_for('dashboard'))
    elif room['meter_id'] is None or room['meter_water_id'] is None:
        flash("ไม่พบข้อมูลมิเตอร์ (กรุณาผูกมิเตอร์ก่อนจองรายวัน)", "danger")
        return redirect(url_for('dashboard'))

    electricity_start = float(room['electricity_start'] or 0)
    water_start = float(room['water_start'] or 0)

    cursor.execute("""
        SELECT setting_key, setting_value
        FROM settings
        WHERE setting_key IN ('electricity_rate', 'water_rate')
    """)
    rows = cursor.fetchall()
    rates = {row['setting_key']: float(
        row['setting_value']) for row in rows}
    electricity_rate = rates.get('electricity_rate', 0)
    water_rate = rates.get('water_rate', 0)

    if request.method == 'POST':
        fname = request.form.get('guest_fname')
        lname = request.form.get('guest_lname')
        phone = request.form.get('guest_phone')
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')
        due_date = request.form.get('due_date')


        start = datetime.strptime(start_date, '%Y-%m-%d').date()
        end = datetime.strptime(end_date, '%Y-%m-%d').date()
        days = (end - start).days + 1

        # ตรวจสอบว่ามีการจองช่วงนั้นอยู่หรือยัง
        cursor.execute("""
        SELECT * FROM invoices
        WHERE unit_id = %s
            AND invoice_type = 'daily'
            AND status IN ('draft', 'paid')  
            AND NOT (billing_period_end < %s OR billing_period_start > %s)
        """, (unit_id, start, end))
        overlap = cursor.fetchone()
        if overlap:
            flash("ห้องนี้มีผู้เช่าในช่วงเวลาที่เลือกแล้ว", "danger")
            return redirect(request.url)

        rent_total = room['price_daily'] * days
        issue_date = datetime.today().date()

        # ✅ เพิ่ม invoice
        cursor.execute("""
            INSERT INTO invoices (
                unit_id, guest_fname, guest_lname, guest_phone, invoice_type,
                billing_period_start, billing_period_end, 
                previous_electricity_reading, previous_water_reading,
                electricity_rate, water_rate,
                issue_date, due_date, rent_amount,
                total_amount, status, created_by, created_at
            ) VALUES (
                %s, %s, %s, %s, 'daily',
                %s, %s, %s, %s,
                %s, %s, %s, %s, %s,
                %s, 'draft', %s, NOW()
            )
        """, (
            unit_id, fname, lname, phone,       # 4 ตัวแรก
            start, end,                        # ช่วงวันที่
            electricity_start, water_start,     # เลขมิเตอร์
            electricity_rate, water_rate,       # เรทราคา (เพิ่มคอมม่าตรงนี้ใน SQL)
            issue_date, due_date, rent_total,   # วันที่ออกบิล และค่าเช่า
            rent_total,                         # total_amount (ใช้ค่าเดียวกับค่าเช่าก่อนในดราฟ)
            session['user']['user_id']          # คนสร้าง
        ))
        daily_booking_id = cursor.lastrowid

        add_audit_log(
            cursor, 
            'DAILY_BOOKING', 
            'CREATE', 
            f'จองรายวันบิล ID: {daily_booking_id}', 
            session.get('user', {}).get('user_id')
        )

        # ✅ อัปเดตสถานะห้องเป็น “มีรอการชําระ” (status_id = 5)
        cursor.execute("""
            UPDATE unit SET status_id = 5 WHERE unit_id = %s
        """, (unit_id,))

        conn.commit()
        flash("สร้างใบแจ้งหนี้รายวันเรียบร้อย", "success")
        return redirect(url_for('dashboard'))

    cursor.close()
    conn.close()
    return render_template('daily_booking.html', room=room)

@app.route('/cancel_daily_booking/<int:invoice_id>', methods=['POST'])
def cancel_daily_booking(invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True) 

    # 1. เช็คข้อมูลบิล
    cursor.execute("SELECT unit_id, status FROM invoices WHERE invoice_id = %s", (invoice_id,))
    invoice = cursor.fetchone()

    if not invoice:
        flash("ไม่พบใบแจ้งหนี้", "danger")
        return redirect(url_for('dashboard'))

    if invoice['status'] != 'draft': # ใช้ชื่อคอลัมน์แทน index
        flash("ไม่สามารถยกเลิกได้ เนื่องจากสถานะไม่ใช่ draft", "warning")
        return redirect(url_for('dashboard'))

    try:
        # 2. อัปเดตสถานะบิล และ คืนสถานะห้องเป็น 'ว่าง' (status_id=1)
        cursor.execute("UPDATE invoices SET status = 'cancelled' WHERE invoice_id = %s", (invoice_id,))
        cursor.execute("UPDATE unit SET status_id = 1 WHERE unit_id = %s", (invoice['unit_id'],))
        
        # 3. บันทึก Audit Log (ควรทำก่อน commit)
        add_audit_log(
            cursor, 
            'DAILY_BOOKING', 
            'DELETE', 
            f'ยกเลิกการจองรายวัน ID: {invoice_id}', 
            session.get('user', {}).get('user_id')
        )
        
        conn.commit() # ยืนยันการเปลี่ยนแปลงทั้งหมด
        flash("ยกเลิกการจองและคืนสถานะห้องเรียบร้อย", "success")
        
    except Exception as e:
        conn.rollback()
        flash(f"เกิดข้อผิดพลาด: {str(e)}", "danger")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('dashboard'))

@app.route("/create_invoice_daily_extra/<int:invoice_id>", methods=["GET", "POST"])
def create_invoice_daily_extra(invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # 1. ดึงข้อมูลบิลหลัก (ลบ tenant_id ออกถ้าไม่ได้ใช้ เพื่อป้องกัน Error)
    cursor.execute("""
        SELECT i.invoice_id, i.unit_id, i.guest_fname, i.guest_lname, i.guest_phone,
               u.name as room_name
        FROM invoices i
        JOIN unit u ON i.unit_id = u.unit_id
        WHERE i.invoice_id = %s
    """, (invoice_id,))
    main_invoice = cursor.fetchone()

    if not main_invoice:
        cursor.close()
        conn.close()
        flash("ไม่พบข้อมูลบิลหลัก", "danger")
        return redirect(url_for('dashboard')) # ถ้าหาบิลไม่เจอ ให้กลับหน้าแรก

    if request.method == 'POST':
        try:
            if not conn.in_transaction:
                conn.start_transaction()

            due_date = request.form.get('due_date')
            total_amount = float(request.form.get('total_amount', 0))

            insert_invoice_query = """
                INSERT INTO invoices (
                    unit_id, invoice_type, 
                    guest_fname, guest_lname,
                    issue_date, due_date, 
                    rent_amount, electricity_total, water_total, total_amount, 
                    status,created_by, created_at,parent_invoice_id
                ) VALUES (%s, %s, %s, %s, %s, %s, 0, 0, 0, %s, %s, %s, NOW(), %s)
            """
            invoice_values = (
                main_invoice['unit_id'], 
                'extra_bill', 
                main_invoice['guest_fname'], 
                main_invoice['guest_lname'],
                datetime.now().date(), 
                due_date,
                total_amount, 
                'draft',
                session.get('user', {}).get('user_id'),
                main_invoice['invoice_id']
            )
            cursor.execute(insert_invoice_query, invoice_values)
            new_invoice_id = cursor.lastrowid

            # 3. บันทึกรายการย่อย (Loop แบบวนตาม Index จริงเพื่อความปลอดภัย)
            items_data = request.form.to_dict(flat=False)
            for key in items_data:
                if 'description' in key:
                    idx = key.split('[')[1].split(']')[0]
                    desc = request.form.get(f'items[{idx}][description]')
                    if desc:
                        amount = float(request.form.get(f'items[{idx}][amount]', 0))
                        item_type = request.form.get(f'items[{idx}][type]', 'other')
                        cursor.execute("""
                            INSERT INTO invoice_items (
                                invoice_id, description, unit_price, quantity, total_price, type
                            ) VALUES (%s, %s, %s, 1, %s, %s)
                        """, (new_invoice_id, desc, amount, amount, item_type))

            conn.commit()
            flash("สร้างบิลเพิ่มเติมสำเร็จ!", "success")
            
            # ✅ แก้ไขจุดนี้: ส่ง invoice_id (ตัวเดิม) กลับไปด้วยเพื่อให้ url_for ทำงานได้
            return redirect(url_for('daily_pending', invoice_id=invoice_id))

        except Exception as e:
            if conn: conn.rollback()
            flash(f"เกิดข้อผิดพลาด: {str(e)}", "danger")
            return redirect(request.url)
        finally:
            if cursor: cursor.close()
            if conn: conn.close()

    # สำหรับ GET: แสดงหน้าฟอร์ม
    return render_template("create_invoice_daily_extra.html", 
                           invoice_id=invoice_id,
                           room_name=main_invoice['room_name'],
                           guest_name=f"{main_invoice['guest_fname']} {main_invoice['guest_lname']}",
                           today_date=datetime.now().strftime('%Y-%m-%d'))

@app.route('/update_daily_invoice', methods=['POST'])
def update_daily_invoice():
    inv_id = request.form.get('invoice_id')
    bill_start = request.form.get('bill_start')
    bill_end = request.form.get('bill_end')
    due_date = request.form.get('due_date')
    cur_elec = request.form.get('cur_elec')
    cur_wat = request.form.get('cur_wat')

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            UPDATE invoices 
            SET billing_period_start=%s, billing_period_end=%s, due_date=%s,
                current_electricity_reading=%s, current_water_reading=%s,
                meter_saved=1
            WHERE invoice_id=%s
        """, (bill_start, bill_end, due_date, cur_elec, cur_wat, inv_id))
        conn.commit()
        flash(f"แก้ไขบิล #{inv_id} เรียบร้อยแล้ว", "success")
    except Exception as e:
        conn.rollback()
        flash(f"ผิดพลาด: {str(e)}", "danger")
    finally:
        cursor.close()
        conn.close()
    
    return redirect(request.referrer)

# --------------------- Checkout -----------------------------
@app.route('/confirm_checkout/<int:unit_id>', methods=['POST'])
def confirm_checkout(unit_id):

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    # ใช้เวลาปัจจุบันสำหรับการบันทึกวันออกจริง
    today = datetime.now().date()

    try:
        # --- 1. ตรวจสอบบิลค้างชำระทั่วไป (ห้ามมี Unpaid หรือ Overdue เด็ดขาด) ---
        cursor.execute("""
            SELECT COUNT(*) AS unpaid FROM invoices 
            WHERE unit_id = %s AND status IN ('draft', 'overdue')
        """, (unit_id,))
        if cursor.fetchone()['unpaid'] > 0:
            flash("❌ ไม่สามารถ Checkout ได้: ยังมียอดค้างชำระในระบบ (กรุณาเคลียร์บิลสีแดง)", "danger")
            return redirect(url_for('dashboard'))

        # --- 2. ตรวจสอบสถานะสัญญาเพื่อแยกประเภท รายเดือน/รายวัน ---
        cursor.execute("""
            SELECT contract_id, notice_move_out_date, contract_end 
            FROM contracts 
            WHERE room_id = %s AND status IN (3, 4)
            LIMIT 1
        """, (unit_id,))
        contract = cursor.fetchone()

        if contract:
            # === กรณีรายเดือน (Monthly Contract) ===
            # ต้องมีบิลประเภท 'final' ที่จ่าย (paid) เรียบร้อยแล้วเท่านั้น
            cursor.execute("""
                SELECT COUNT(*) AS final_paid FROM invoices 
                WHERE contract_id = %s 
                  AND invoice_type = 'final' 
                  AND status = 'paid'
            """, (contract['contract_id'],))
            
            if cursor.fetchone()['final_paid'] == 0:
                flash("❌ ต้องออกบิลปิดยอด และจัดเก็บเงินให้เรียบร้อยก่อน", "warning")
                return redirect(url_for('dashboard'))

            # กำหนดสถานะสัญญาหลัง Checkout (5=จบสัญญาปกติ, 6=จบสัญญาแบบแจ้งย้าย)
            final_contract_status = 6 if contract['notice_move_out_date'] else 5
            
            # อัปเดตสถานะสัญญา และบันทึกวันออกจริง
            cursor.execute("""
                UPDATE contracts 
                SET status = %s, actual_move_out_date = %s 
                WHERE contract_id = %s
            """, (final_contract_status, today, contract['contract_id']))

        else:
            # === กรณีรายวัน (Daily) ===
            # เช็คว่าบิลรายวันล่าสุดของห้องนี้ จ่ายเงินหรือยัง
            cursor.execute("""
                SELECT COUNT(*) AS daily_unpaid FROM invoices 
                WHERE unit_id = %s AND invoice_type = 'daily' 
                           AND status NOT IN ('paid', 'cancelled', 'void')
            """, (unit_id,))
            
            if cursor.fetchone()['daily_unpaid'] > 0:
                flash("❌ (รายวัน) กรุณาตรวจสอบการชำระเงินค่าห้องรายวันก่อน Checkout", "warning")
                return redirect(url_for('dashboard'))

            # เปลี่ยนประเภทบิลรายวันจาก 'daily' เป็น 'daily_checkout' เพื่อปิดยอดในทางบัญชี
            cursor.execute("""
                UPDATE invoices 
                SET invoice_type = 'daily_checkout' 
                WHERE unit_id = %s 
                  AND invoice_type = 'daily' 
                  AND status = 'paid'
            """, (unit_id,))

        # --- 3. ขั้นตอนสุดท้าย: คืนสถานะห้องว่าง (Status 1) ---
        cursor.execute("""
            UPDATE unit 
            SET status_id = 1 
            WHERE unit_id = %s
        """, (unit_id,))

        add_audit_log(
            cursor, 
            'UNIT', 
            'CHECKOUT', 
            f'ยืนยันคืนห้อง ID: {unit_id}', 
            session.get('user', {}).get('user_id')
        )

        # บันทึกข้อมูลทั้งหมดลง Database
        conn.commit()
        flash(f"✅ ดำเนินการคืนห้อง {unit_id} เรียบร้อยแล้ว สถานะห้องเป็น 'ว่าง'", "success")

    except Exception as e:
        conn.rollback()
        print(f"Error confirm_checkout: {e}")
        flash("เกิดข้อผิดพลาดรุนแรงในการอัปเดตข้อมูล กรุณาติดต่อผู้ดูแลระบบ", "danger")
    
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('dashboard'))


# ---------------------- BUSINESS INFORMATION ----------------------
@app.route('/business', methods=['GET', 'POST'])
@role_required(['admin','manager'])
def business():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        owner_name = request.form.get('owner_name')
        business_name = request.form.get('business_name')
        tax_id = request.form.get('tax_id')
        address = request.form.get('address')
        phone = request.form.get('phone')
        email = request.form.get('email')
        bank_name = request.form.get('bank_name')
        bank_account_no = request.form.get('bank_account_no')
        account_name = request.form.get('account_name')
        promptpay_id = request.form.get('promptpay_id', '').replace('-', '').strip()

        update_query = """
            UPDATE business
            SET name=%s, tax_id=%s, address=%s, tel=%s, email=%s , bank_name=%s,bank_account_no=%s,account_name=%s,promptpay_id=%s, owner_name=%s
            WHERE id=1
        """
        try:
            cursor.execute(update_query, (business_name,
                           tax_id, address, phone, email, bank_name, 
                           bank_account_no, account_name, promptpay_id, owner_name ))
            conn.commit()
            flash('บันทึกข้อมูลสำเร็จ!', 'success')
        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {e}', 'danger')
        cursor.close()
        conn.close()
        return redirect(url_for('business'))

    cursor.execute(
        "SELECT * FROM business WHERE id = 1")
    business_data = cursor.fetchone()
    cursor.close()
    conn.close()

    if business_data is None:
        # กำหนดค่าเริ่มต้นถ้ายังไม่มีข้อมูลใน DB
        business_data = {
            'name': '', 'tax_id': '', 'address': '', 'tel': '', 
            'email': '', 'bank_name': '', 'bank_account_no': '', 
            'account_name': '', 'promptpay_id': '', 'owner_name': ''
        }

    return render_template('business.html', biz=business_data)


# ---------------------- MANAGEMENT UNITS ----------------------
@app.route('/manage_units', methods=['GET', 'POST'])
def manage_units():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        name = request.form.get('name')
        building = request.form.get('building')
        floor = request.form.get('floor')
        zone = request.form.get('zone')
        type_unit_id = request.form.get('type_unit_id')

        try:
            cursor.execute("""
                INSERT INTO unit (name, created_at, status_id, type_unit_id, building, floor, zone)
                VALUES (%s, NOW(), %s, %s, %s, %s, %s)
            """, (name, 1, type_unit_id, building, floor, zone))
            conn.commit()
            flash('เพิ่มห้องเรียบร้อยแล้ว', 'success')
        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {e}', 'danger')

    # ดึงข้อมูลห้อง พร้อมราคาหลากหลายประเภท
    cursor.execute("""
        SELECT u.*, s.name_status AS status_name, t.n_type AS type_name,
               t.price_monthly, t.price_daily
        FROM unit u 
        LEFT JOIN s_unit s ON u.status_id = s.status_id
        LEFT JOIN type t ON u.type_unit_id = t.type_id
        WHERE u.is_deleted = 0
        ORDER BY u.building, u.floor, u.name
    """)
    units = cursor.fetchall()

    # ดึงประเภทห้อง
    cursor.execute("SELECT * FROM type t WHERE t.is_deleted = 0")
    types = cursor.fetchall()

    cursor.close()
    conn.close()

    if not units:
        flash('ไม่พบข้อมูลห้อง', 'warning')
    return render_template('manage_units.html', units=units, types=types)

@app.route('/add_type', methods=['POST'])
def add_type():
    n_type = request.form.get('n_type')
    price_monthly = request.form.get('price_monthly')
    price_daily = request.form.get('price_daily')

    if not n_type or not price_monthly or not price_daily:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", "danger")
        return redirect(url_for('manage_units'))

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            'INSERT INTO type (n_type, price_monthly, price_daily) VALUES (%s, %s, %s)',
            (n_type, price_monthly, price_daily)
        )
        new_id = cursor.lastrowid

        add_audit_log(
            cursor, 
            'TYPE', 
            'INSERT', 
            f'เพิ่มประเภทห้อง ID: {new_id}', 
            session.get('user', {}).get('user_id')
        )

        conn.commit()
        flash("เพิ่มประเภทห้องสำเร็จ", "success")
    except Exception as e:
        conn.rollback()
        flash(f"เกิดข้อผิดพลาด: {e}", "danger")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('manage_units'))

@app.route('/delete_unit/<int:unit_id>', methods=['POST'])
def delete_unit(unit_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        cursor.execute("SELECT status_id FROM unit WHERE unit_id = %s", (unit_id,))
        room = cursor.fetchone()

        if not room:
            flash("ไม่พบข้อมูลห้องนี้", "danger")
            return redirect(url_for('manage_units'))

        cursor.execute("""
            SELECT contract_id FROM contracts 
            WHERE room_id = %s 
            AND status IN (1, 2, 3, 4)
        """, (unit_id,))
        active_contract = cursor.fetchone()

        restricted_room_statuses = [2, 3, 4, 5, 6, 7]
        
        if room['status_id'] in restricted_room_statuses or active_contract:
            if active_contract:
                flash("ไม่สามารถลบได้: มีสัญญาที่อยู่ในสถานะ ร่าง/รอชำระ/ใช้งาน หรือ ใกล้หมดอายุ", "warning")
            else:
                flash("ไม่สามารถลบได้: สถานะห้องยังไม่ว่าง (ต้องเป็นสถานะ 'ว่าง' เท่านั้น)", "warning")
            return redirect(url_for('manage_units'))
        
        cursor.execute("UPDATE unit SET is_deleted = 1 WHERE unit_id=%s", (unit_id,))

        add_audit_log(
            cursor, 
            'UNIT', 
            'DELETE', 
            f'ลบห้อง ID: {unit_id}', 
            session.get('user', {}).get('user_id')
        )
        
        conn.commit()
        flash("ย้ายห้องไปที่ถังขยะเรียบร้อยแล้ว (Soft Delete)", "success")
        
    except Exception as e:
        conn.rollback()
        flash(f"เกิดข้อผิดพลาดในการลบห้อง: {e}", "danger")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('manage_units'))

@app.route('/edit_unit/<int:unit_id>', methods=['GET', 'POST'])
def edit_unit(unit_id):

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # 1. ดึงข้อมูลห้องปัจจุบันขึ้นมาก่อน เพื่อเช็คสถานะเดิม
    cursor.execute("SELECT * FROM unit WHERE unit_id = %s AND is_deleted = 0", (unit_id,))
    room = cursor.fetchone()

    if not room:
        flash('ไม่พบข้อมูลห้องพัก', 'danger')
        return redirect(url_for('manage_units'))

    if request.method == 'POST':
        name = request.form.get('name')
        building = request.form.get('building')
        floor = request.form.get('floor')
        zone = request.form.get('zone')
        type_unit_id = request.form.get('type_unit_id')
        
        # รับค่าจากฟอร์ม (hidden หรือ button)
        status_id_from_form = request.form.get('status_id')
        # รับค่าจาก switch (ถ้าติ๊กจะได้ 'on')
        mark_as_broken = request.form.get('mark_as_broken')

        # --- Logic จัดการสถานะ ---
        final_status_id = room['status_id']  # เริ่มต้นด้วยสถานะเดิมจาก DB

        if int(room['status_id']) == 1:  # ถ้าเดิมคือ 'ว่าง'
            if mark_as_broken == 'on':
                final_status_id = 4     # เปลี่ยนเป็น 'ปิดปรับปรุง'
            else:
                final_status_id = 1     # ยังคง 'ว่าง' เหมือนเดิม
        
        elif int(room['status_id']) == 4: # ถ้าเดิมคือ 'ปิดปรับปรุง'
            if status_id_from_form == '1':
                final_status_id = 1     # เปลี่ยนกลับเป็น 'ว่าง'
            else:
                final_status_id = 4     # ยังคง 'ปิดปรับปรุง'

        # สำหรับสถานะ 2, 3, 5, 6, 7 ระบบจะไม่แก้สถานะในหน้านี้ (ใช้ค่าเดิมใน DB)

        try:
            cursor.execute("""
                UPDATE unit SET 
                    name=%s, building=%s, floor=%s, zone=%s, type_unit_id=%s, status_id=%s
                WHERE unit_id=%s
            """, (name, building, floor, zone, type_unit_id, final_status_id, unit_id))

            add_audit_log(
                cursor, 
                'UNIT', 
                'UPDATE', 
                f'แก้ไขห้อง ID: {unit_id}', 
                session.get('user', {}).get('user_id')
            )
            conn.commit()
            flash(f'แก้ไขข้อมูลห้อง {name} สำเร็จ', 'success')
            return redirect(url_for('manage_units'))
        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาดในการบันทึก: {e}', 'danger')
        finally:
            cursor.close()
            conn.close()

    # GET: ดึงข้อมูลเพื่อแสดงผล
    cursor.execute("SELECT type_id, n_type, price_monthly FROM type WHERE is_deleted = 0 ORDER BY n_type")
    types = cursor.fetchall()
    
    cursor.close()
    conn.close()

    return render_template('edit_unit.html', room=room, types=types)

@app.route('/delete_type/<int:type_id>', methods=['POST'])
def delete_type(type_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("UPDATE type SET is_deleted = 1 WHERE type_id=%s", (type_id,))
        add_audit_log(
            cursor, 
            'TYPE', 
            'DELETE', 
            f'ลบประเภทห้อง ID: {type_id}', 
            session.get('user', {}).get('user_id')
        )
        conn.commit()
        
        flash("ลบประเภทห้องเรียบร้อย! อย่าลืมไปเปลี่ยน 'ประเภทห้อง' ในหน้าจัดการห้องพักให้เป็นประเภทอื่นด้วยนะครับ", "success")
    except Exception as e:
        conn.rollback()
        flash(f"เกิดข้อผิดพลาด: {e}", "danger")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('manage_units'))

@app.route('/edit_type/<int:type_id>', methods=['GET', 'POST'])
def edit_type(type_id):

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        n_type = request.form.get('n_type')
        price_monthly = request.form.get('price_monthly')
        price_daily = request.form.get('price_daily')

        try:
            cursor.execute("""
                UPDATE type SET
                    n_type=%s, price_monthly=%s, price_daily=%s
                WHERE type_id=%s
            """, (n_type, price_monthly, price_daily, type_id))

            add_audit_log(
                cursor, 
                'TYPE', 
                'UPDATE', 
                f'แก้ไขประเภทห้อง ID: {type_id}', 
                session.get('user', {}).get('user_id')
            )
            conn.commit()
            flash('แก้ไขประเภทห้องเรียบร้อยแล้ว', 'success')
            return redirect(url_for('manage_units'))
        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {e}', 'danger')

    cursor.execute("SELECT * FROM type WHERE type_id = %s", (type_id,))
    room_type = cursor.fetchone()

    cursor.close()
    conn.close()

    return render_template('edit_type.html', room_type=room_type)


# ---------------------- MANAGE Option ----------------------
@app.route('/manage_option')
def manage_option():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM `option` WHERE is_deleted = 0 ORDER BY name")
    options = cursor.fetchall()
    cursor.close()
    conn.close()
    return render_template('manage_option.html', options=options)

@app.route('/edit_option/<int:option_id>', methods=['GET', 'POST'])
def edit_option(option_id):
        
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        name = request.form.get('name')
        price = request.form.get('price')

        cursor.execute("""
            UPDATE `option`
            SET name = %s, price = %s
            WHERE id = %s
        """, (name, price, option_id))

        add_audit_log(
            cursor, 
            'OPTION', 
            'UPDATE', 
            f'แก้ไข Option ID: {option_id}', 
            session.get('user', {}).get('user_id')
        )
        
        conn.commit()
        cursor.close()
        conn.close()

        flash('อัปเดตข้อมูลสำเร็จแล้ว', 'success')
        return redirect(url_for('manage_option'))

    cursor.execute("SELECT * FROM `option` WHERE id = %s", (option_id,))
    option = cursor.fetchone()
    cursor.close()
    conn.close()

    if not option:
        flash('ไม่พบข้อมูล Option นี้', 'danger')
        return redirect(url_for('manage_option'))

    return render_template('edit_option.html', option=option)

@app.route('/add_option', methods=['GET', 'POST'])
def add_option():
    if request.method == 'POST':
        name = request.form.get('name')
        price = request.form.get('price')
        conn = get_db_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute("""
                INSERT INTO `option` (name, price)
                VALUES (%s, %s)
            """, (name, price))
            option_id = cursor.lastrowid
            add_audit_log(
                cursor, 
                'OPTION', 
                'INSERT', 
                f'เพิ่ม Option ID: {option_id}', 
                session.get('user', {}).get('user_id')
            )
            
            conn.commit()
            flash('เพิ่ม Option เรียบร้อยแล้ว', 'success')
        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {str(e)}', 'danger')
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('manage_option'))
    return render_template('add_option.html')

@app.route('/delete_option/<int:option_id>', methods=['POST'])
def delete_option(option_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True, buffered=True) 
    
    try:
        # 1. เช็คว่า "เคยมีการนำไปผูกกับสัญญาไหนบ้างไหม" (ไม่สน status)
        cursor.execute("SELECT id FROM contract_option WHERE option_id = %s LIMIT 1", (option_id,))
        has_history = cursor.fetchone()

        if not has_history:
            # --- กรณีที่ 1: ไม่เคยถูกนำไปใช้เลย -> ลบทิ้งจริงๆ (Hard Delete) ---
            cursor.execute("DELETE FROM `option` WHERE id = %s", (option_id,))
            add_audit_log(
                cursor, 
                'OPTION', 
                'DELETE', 
                f'ลบ Option ID: {option_id}', 
                session.get('user', {}).get('user_id')
            )
            conn.commit()
            flash('ลบข้อมูล Option ออกจากระบบเรียบร้อยแล้ว (Hard Delete)', 'success')
            return redirect(url_for('manage_option'))

        # 2. ถ้าเคยถูกใช้ (มีประวัติ) -> ต้องเช็คต่อว่า "ตอนนี้ยังติดสัญญาที่ใช้งานอยู่ไหม"
        query_active = """
            SELECT co.id FROM contract_option co
            INNER JOIN contracts c ON co.contract_id = c.contract_id
            WHERE co.option_id = %s AND c.status IN (1, 2, 3, 4)
            LIMIT 1
        """
        cursor.execute(query_active, (option_id,))
        is_still_using = cursor.fetchone()

        if is_still_using:
            # --- กรณีที่ 2: ยังติดสัญญาที่ใช้งานอยู่ (1-4) -> ห้ามลบ ---
            flash('ไม่สามารถลบได้: Option นี้กำลังถูกใช้งานในสัญญาปัจจุบัน', 'warning')
        else:
            # --- กรณีที่ 3: เคยใช้ แต่สัญญาเหล่านั้นจบหมดแล้ว (5-7) -> ซ่อนไว้ (Soft Delete) ---
            cursor.execute("UPDATE `option` SET is_deleted = 1 WHERE id = %s", (option_id,))
            conn.commit()
            flash('ซ่อน Option เรียบร้อยแล้ว (ข้อมูลในสัญญาเก่าจะยังคงอยู่)', 'success')

    except Exception as e:
        if conn: conn.rollback()
        flash(f'เกิดข้อผิดพลาด: {e}', 'danger')
    finally:
        if cursor: cursor.close()
        if conn: conn.close()
        
    return redirect(url_for('manage_option'))


# ---------------------- MANAGE METER ----------------------
@app.route('/manage_meter', methods=['GET', 'POST'])
def manage_meter():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    user_id = session.get('user', {}).get('user_id')
    now = datetime.now()

    # --- 1. โหลดไฟล์ Config JSON ---
    config_dir = os.path.join(os.path.dirname(__file__), 'config_meter')
    elec_json_data = {}
    water_json_data = {}
    try:
        with open(os.path.join(config_dir, 'model.json'), 'r', encoding='utf-8') as f:
            elec_json_data = json.load(f)
        with open(os.path.join(config_dir, 'model_water.json'), 'r', encoding='utf-8') as f:
            water_json_data = json.load(f)
    except Exception as e:
        print(f"JSON Load Error: {e}")

    # PART A: จัดการการบันทึกข้อมูล (POST)
    if request.method == 'POST':
        try:
            if session.get('role') not in  ['admin', 'manager']:
                return jsonify({'status': 'error', 'massage': 'คุณไม่มีสิทธิ์แก้ไขข้อมูลมิเตอร์ (เฉพาะ Admin เเละ ผู้จัดการเท่านั้น เท่านั้น)'})
            unit_id = request.form.get('unit_id')       
            # รับค่า Flag การเปลี่ยนมิเตอร์และเลขมิเตอร์เก่า
            change_elec = request.form.get('change_elec_flag') == 'on'
            change_water = request.form.get('change_water_flag') == 'on'
            old_elec_last = float(request.form.get('old_elec_last_reading') or 0)
            old_water_last = float(request.form.get('old_water_last_reading') or 0)
            e_port_raw = request.form.get('electricity_port')
            elec_port = int(e_port_raw) if e_port_raw and e_port_raw.strip() else None
            w_port_raw = request.form.get('water_port')
            water_port = int(w_port_raw) if w_port_raw and w_port_raw.strip() else None
            elec_base_url = request.form.get('elec_base_url')
            elec_api_token = request.form.get('elec_api_token')
            water_base_url = request.form.get('water_base_url')
            water_api_token = request.form.get('water_api_token')
            
            # --- จัดการไฟฟ้า ---
            cursor.execute("SELECT id, serial_meter FROM meter WHERE unit_id=%s LIMIT 1", (unit_id,))
            m_exists = cursor.fetchone()

            elec_data = (
                request.form.get('electricity_serial_meter') or None, 
                request.form.get('electricity_slave_id') or None,
                request.form.get('electricity_module'), 
                datetime.strptime(request.form.get('installdate_elec'), '%d/%m/%Y').strftime('%Y-%m-%d') if request.form.get('installdate_elec') else None,
                request.form.get('electricity_comport')or None, 
                request.form.get('electricity_ip')or None,
                elec_port,
                elec_base_url,  
                elec_api_token, 
                request.form.get('electricity_status')or None,
                request.form.get('elec_unit_key')or None, 
                now, user_id
            )

            if not m_exists:
                # --- เคสที่ 1: เพิ่มมิเตอร์ครั้งแรก (INSERT) ---
                cursor.execute("""
                    INSERT INTO meter (
                        serial_meter, slave_id, module, installdate, 
                        comport, ip, port, base_url, api_auth_token, 
                        status, unit_key, created_at, created_by, 
                        unit_id, current_reading
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 0.00)
                """, elec_data + (unit_id,))
                
                m_id = cursor.lastrowid
                add_audit_log(cursor, 'METER', 'INSERT', f'เพิ่มมิเตอร์ไฟฟ้าใหม่ ID: {m_id} ห้อง: {unit_id}', user_id)

            else:
                # --- เคสที่ 2: มีมิเตอร์อยู่แล้ว (UPDATE) ---
                m_id = m_exists['id']

                if change_elec:
                    # 1. คำนวณหน่วยค้างจ่าย (ดัก NoneType กรณีพึ่งเคยมีบิลครั้งแรก)
                    cursor.execute("""
                        SELECT current_electricity_reading FROM invoices 
                        WHERE unit_id = %s AND status != 'draft' 
                        ORDER BY invoice_id DESC LIMIT 1
                    """, (unit_id,))
                    last_bill = cursor.fetchone()
                    
                    # ถ้า last_bill เป็น None (ห้องใหม่พึ่งติดมิเตอร์) ให้ใช้ 0
                    last_bill_val = float(last_bill['current_electricity_reading'] or 0) if last_bill else 0
                    pending_calc = float(old_elec_last or 0) - last_bill_val

                    # 2. บันทึกประวัติการถอดมิเตอร์เก่า
                    cursor.execute("""
                        INSERT INTO meter_history (unit_id, type, old_serial, final_reading, pending_units, created_at) 
                        VALUES (%s, 'elec', %s, %s, %s, NOW())
                    """, (unit_id, m_exists['serial_meter'], old_elec_last, pending_calc))

                    # 3. อัปเดตพร้อมรีเซ็ตเลขจดเป็น 0.00 (เพราะเปลี่ยนตัวใหม่)
                    cursor.execute("""
                        UPDATE meter SET 
                            serial_meter=%s, slave_id=%s, module=%s, installdate=%s, 
                            comport=%s, ip=%s, port=%s, base_url=%s, api_auth_token=%s, 
                            status=%s, unit_key=%s, 
                            updated_at=%s, updated_by=%s, current_reading = 0.00 
                        WHERE id=%s
                    """, elec_data + (m_id,))
                    
                    # อัปเดตค่าเริ่มต้นในตาราง unit
                    cursor.execute("UPDATE unit SET electricity_start = 0 WHERE unit_id=%s", (unit_id,))
                    log_msg = f'เปลี่ยนมิเตอร์ไฟฟ้าใหม่ ID: {m_id}'
                else:
                    # 4. อัปเดตข้อมูลทั่วไป (แก้ไขชื่อ/สาย/IP)
                    cursor.execute("""
                        UPDATE meter SET 
                            serial_meter=%s, slave_id=%s, module=%s, installdate=%s, 
                            comport=%s, ip=%s, port=%s, base_url=%s, api_auth_token=%s, 
                            status=%s, unit_key=%s, 
                            updated_at=%s, updated_by=%s 
                        WHERE id=%s
                    """, elec_data + (m_id,))
                    log_msg = f'อัปเดตข้อมูลมิเตอร์ไฟฟ้า ID: {m_id}'

                add_audit_log(cursor, 'METER', 'UPDATE', log_msg, user_id)

            # --- จัดการน้ำ ---
            cursor.execute("SELECT id, serial_meter FROM meter_water WHERE unit_id=%s LIMIT 1", (unit_id,))
            mw_exists = cursor.fetchone()
            
            water_data = (
                request.form.get('water_serial') or None, 
                request.form.get('water_slave_id') or None,
                request.form.get('water_module'),
                datetime.strptime(request.form.get('installdate_water'), '%d/%m/%Y').strftime('%Y-%m-%d') if request.form.get('installdate_water') else None,
                request.form.get('water_comport') or None, 
                request.form.get('water_ip') or None, 
                water_port,
                water_base_url,  
                water_api_token,
                request.form.get('water_status') or None,
                request.form.get('water_unit_key') or None, 
                now, user_id
            )

            if not mw_exists:
                # --- เคส B: เพิ่มมิเตอร์น้ำครั้งแรก (INSERT) ---
                cursor.execute("""
                    INSERT INTO meter_water (
                        serial_meter, slave_id, module, installdate, 
                        comport, ip, port, base_url, api_auth_token, 
                        status, unit_key, created_at, created_by, 
                        unit_id, current_reading
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 0.00)
                """, water_data + (unit_id,))
                
                mw_id = cursor.lastrowid
                add_audit_log(cursor, 'METER_WATER', 'INSERT', f'เพิ่มมิเตอร์น้ำใหม่ ID: {mw_id} ห้อง: {unit_id}', user_id)

            else:
                # --- เคส A: มีมิเตอร์น้ำเดิมอยู่แล้ว (UPDATE) ---
                mw_id = mw_exists['id']

                if change_water:
                    # 1. คำนวณหน่วยค้างจ่ายน้ำ (ดัก NoneType สำหรับห้องใหม่)
                    cursor.execute("""
                        SELECT current_water_reading FROM invoices 
                        WHERE unit_id = %s AND status != 'draft' 
                        ORDER BY invoice_id DESC LIMIT 1
                    """, (unit_id,))
                    last_bill_w = cursor.fetchone()
                    
                    last_bill_val_w = float(last_bill_w['current_water_reading'] or 0) if last_bill_w else 0
                    pending_calc_w = float(old_water_last or 0) - last_bill_val_w

                    # 2. บันทึกประวัติการถอดมิเตอร์น้ำเก่า
                    cursor.execute("""
                        INSERT INTO meter_history (unit_id, type, old_serial, final_reading, pending_units, created_at) 
                        VALUES (%s, 'water', %s, %s, %s, NOW())
                    """, (unit_id, mw_exists['serial_meter'], old_water_last, pending_calc_w))

                    # 3. อัปเดตพร้อมรีเซ็ตเลขจดน้ำ
                    cursor.execute("""
                        UPDATE meter_water SET 
                            serial_meter=%s, slave_id=%s, module=%s, installdate=%s, 
                            comport=%s, ip=%s, port=%s, base_url=%s, api_auth_token=%s, 
                            status=%s, unit_key=%s, 
                            updated_at=%s, updated_by=%s, current_reading = 0.00 
                        WHERE id=%s
                    """, water_data + (mw_id,))
                    
                    cursor.execute("UPDATE unit SET water_start = 0 WHERE unit_id=%s", (unit_id,))
                    log_msg_w = f'เปลี่ยนมิเตอร์น้ำใหม่ ID: {mw_id}'
                else:
                    # 4. อัปเดตข้อมูลทั่วไป
                    cursor.execute("""
                        UPDATE meter_water SET 
                            serial_meter=%s, slave_id=%s, module=%s, installdate=%s, 
                            comport=%s, ip=%s, port=%s, base_url=%s, api_auth_token=%s, 
                            status=%s, unit_key=%s, 
                            updated_at=%s, updated_by=%s 
                        WHERE id=%s
                    """, water_data + (mw_id,))
                    log_msg_w = f'อัปเดตข้อมูลมิเตอร์น้ำ ID: {mw_id}'

                add_audit_log(cursor, 'METER_WATER', 'UPDATE', log_msg_w, user_id)

            cursor.execute("UPDATE unit SET meter_id=%s, meter_water_id=%s WHERE unit_id=%s", (m_id, mw_id, unit_id))
            conn.commit()
            return jsonify({'status': 'success', 'massage': 'บันทึกข้อมูลเรียบร้อย'})

        except Exception as e:
            conn.rollback()
            return jsonify({'status': 'error', 'massage': str(e)})
        finally:
            cursor.close()
            conn.close()

    # แสดงหน้าจอปกติ (GET) - **ต้องอยู่นอก block POST**
    try:
        cursor.execute("""
            SELECT u.*, 
                   m.id AS mid, m.serial_meter AS electricity_serial_meter, m.slave_id AS electricity_slave_id, m.module AS electricity_module,
                   m.installdate AS electricity_installdate, m.comport AS electricity_comport, m.ip AS electricity_ip, 
                   m.port AS electricity_port, m.status AS electricity_status, m.unit_key AS elec_unit_key,
                   m.base_url AS elec_base_url, m.api_auth_token AS elec_api_token,
                   mw.id AS mwid, mw.serial_meter AS water_serial, mw.slave_id AS water_slave_id, mw.module AS water_module,
                   mw.installdate AS water_installdate, mw.comport AS water_comport, mw.ip AS water_ip, 
                   mw.port AS water_port, mw.status AS water_status, mw.unit_key AS water_unit_key,
                   mw.base_url AS water_base_url, mw.api_auth_token AS water_api_token 
            FROM unit u
            LEFT JOIN meter m ON u.meter_id = m.id
            LEFT JOIN meter_water mw ON u.meter_water_id = mw.id
            WHERE is_deleted=0
        """)
        units = cursor.fetchall()
        
        import serial.tools.list_ports
        ports = [p.device for p in serial.tools.list_ports.comports()] + ["TCP/IP"] + ["API"]

        # คืนค่าหน้าเว็บเสมอเมื่อเป็น GET
        return render_template('manage_meter.html', 
                               units=units, ports=ports, 
                               elec_modules=sorted(elec_json_data.keys()), 
                               water_modules=sorted(water_json_data.keys()),
                               elec_models_json=elec_json_data, 
                               water_models_json=water_json_data)
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/get_latest_iot_reading/<int:unit_id>/<string:meter_type>')
def get_latest_iot_reading(unit_id, meter_type):
    
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    db_type = 'electricity' if meter_type == 'elec' else 'water'
    try:
        cursor.execute("SELECT current_reading FROM meter_reading WHERE unit_id=%s AND meter_type=%s ORDER BY read_date DESC LIMIT 1", (unit_id, db_type))
        result = cursor.fetchone()
        return jsonify({"status": "success", "reading": float(result['current_reading']) if result else 0})
    finally:
        cursor.close()
        conn.close()

@app.route('/clear_meter/<int:unit_id>', methods=['POST'])
@role_required(['admin','manager'])
def clear_meter(unit_id):
    conn = get_db_connection()
    cursor = conn.cursor()

    # ดึง meter_id และ meter_water_id จาก unit
    cursor.execute(
        "SELECT meter_id, meter_water_id FROM unit WHERE unit_id = %s", (unit_id,))
    row = cursor.fetchone()
    meter_id, water_id = row if row else (None, None)

    # ลบจาก unit
    cursor.execute(
        "UPDATE unit SET meter_id = NULL, meter_water_id = NULL WHERE unit_id = %s", (unit_id,))

    # ลบจาก meter ไฟ
    if meter_id:
        cursor.execute("DELETE FROM meter WHERE id = %s", (meter_id,))

    # ลบจาก meter_water น้ำ
    if water_id:
        cursor.execute("DELETE FROM meter_water WHERE id = %s", (water_id,))

    conn.commit()
    cursor.close()
    conn.close()

    flash("ลบข้อมูลมิเตอร์สำเร็จ", "warning")
    return redirect(url_for('manage_meter'))


# ---------------------- MANAGE DOCUMENTS ----------------------
@app.route('/manage_doc', methods=['GET', 'POST'])
def manage_doc():
    # กำหนด Path ตรงๆ
    base_path = 'uploaded_docs'
    temp_path = os.path.join(base_path, 'templates')

    # สร้างโฟลเดอร์ถ้ายังไม่มี
    for p in [base_path, temp_path]:
        if not os.path.exists(p):
            os.makedirs(p)

    if request.method == 'POST':
        file = request.files.get('file')
        category = request.form.get('category')
        
        if file and file.filename != '' and category:
            ext = file.filename.rsplit('.', 1)[1].lower()
            
            if category == 'template':
                # ✅ บังคับ .docx และชื่อไฟล์เดียวเพื่อใช้ Gen สัญญา
                if ext != 'docx':
                    flash('Template สำหรับ Gen สัญญาต้องเป็น .docx เท่านั้น', 'danger')
                    return redirect(url_for('manage_doc'))
                filename = "main_contract_template.docx"
                filepath = os.path.join(temp_path, filename)
            else:
                # ✅ ไฟล์ทั่วไปเก็บใน uploaded_docs
                # ใช้ timestamp เพื่อป้องกันชื่อไฟล์ซ้ำ
                filename = f"{category}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
                filepath = os.path.join(base_path, filename)

            file.save(filepath)
            flash(f'อัปโหลด {filename} สำเร็จแล้ว', 'success')
        else:
            flash('กรุณาเลือกไฟล์และประเภทเอกสาร', 'warning')
        return redirect(url_for('manage_doc'))

    # ----------------- อ่านไฟล์จากโฟลเดอร์สดๆ ------------------
    files = {'general': [], 'template': []}
    
    # 1. อ่านไฟล์ทั่วไปใน uploaded_docs
    if os.path.exists(base_path):
        for f in os.listdir(base_path):
            full_path = os.path.join(base_path, f)
            # กรองเอาเฉพาะไฟล์ (ไม่เอาโฟลเดอร์ templates มาโชว์)
            if os.path.isfile(full_path):
                files['general'].append(f)
    
    # 2. อ่านไฟล์ใน uploaded_docs/templates
    if os.path.exists(temp_path):
        files['template'] = [f for f in os.listdir(temp_path) 
                             if os.path.isfile(os.path.join(temp_path, f))]

    return render_template('manage_doc.html', files=files)


# ---------------------- MANAGE TRANETS -----------------------
@app.route('/manage_tenants', methods=['GET', 'POST'])
def manage_tenants():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        action = request.form.get('action')
        tenant_id = request.form.get('tenant_id')
        id_card_val = request.form.get('id_card')

        file = request.files.get('id_card_file')
        filename = None
        if file and file.filename != '' and allowed_file(file.filename):
            ext = file.filename.rsplit('.', 1)[1].lower()
 
            filename = secure_filename(f"id_card_tenant_{id_card_val}.{ext}")
            file.save(os.path.join(UPLOAD_ID_CARD_TENANTS, filename))

        # 1. ADD or EDIT
        if action in ['add', 'edit']:
            data = (
                request.form.get('id_card'), request.form.get('fname'),
                request.form.get('lname'), request.form.get('gender'),
                request.form.get('age'), request.form.get('bd'),
                request.form.get('tel'), request.form.get('address'),
                request.form.get('email')
            )

            if action == 'add':
                query = """INSERT INTO tenants (id_card, fname, lname, gender, age, bd, tel, address, email, id_card_img) 
                           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"""
                cursor.execute(query, data + (filename,))
                add_audit_log(
                    cursor, 
                    'TENANT', 
                    'INSERT', 
                    f'เพิ่มผู้เช่าใหม่ ID: {tenant_id}', 
                    session.get('user', {}).get('user_id')
                )
                flash("เพิ่มผู้เช่าใหม่เรียบร้อย", "success")
            else:
                query = """UPDATE tenants SET id_card=%s, fname=%s, lname=%s, gender=%s, 
                           age=%s, bd=%s, tel=%s, address=%s, email=%s, id_card_img=%s WHERE tenant_id=%s"""
                cursor.execute(query, data + (filename, tenant_id))
                add_audit_log(
                    cursor, 
                    'TENANT', 
                    'UPDATE', 
                    f'อัพเดทผู้เช่า ID: {tenant_id}', 
                    session.get('user', {}).get('user_id')
                )
                flash("แก้ไขข้อมูลเรียบร้อย", "info")

        # 2. DELETE (ตรวจสอบสัญญาก่อนลบ)
        elif action == 'delete':
            # 1. เช็คสัญญาก่อน และดึงชื่อไฟล์รูปภาพมาด้วยในคราวเดียว (ประหยัด Query)
            cursor.execute("""
                SELECT 
                    (SELECT COUNT(*) FROM contracts WHERE tenant_id = %s AND status IN (1, 2, 3, 4) AND is_deleted = 0) as active_contracts,
                    id_card_img 
                FROM tenants WHERE tenant_id = %s
            """, (tenant_id, tenant_id))
            
            result = cursor.fetchone()
            
            # เช็คว่าเจอข้อมูลไหม
            if not result:
                flash("ไม่พบข้อมูลผู้เช่า", "danger")
            else:
                # ดึงค่าแบบรองรับทั้ง Tuple และ Dict
                active_count = result['active_contracts'] if isinstance(result, dict) else result[0]
                id_card_img = result['id_card_img'] if isinstance(result, dict) else result[1]

                if active_count > 0:
                    flash("❌ ไม่สามารถลบได้: ผู้เช่ายังมีสัญญาที่ใช้งานอยู่ในระบบ", "danger")
                else:
                    try:
                        # --- ลบไฟล์จริงออกจากเครื่อง ---
                        if id_card_img:
                            file_path = os.path.join(UPLOAD_ID_CARD, id_card_img)
                            if os.path.exists(file_path):
                                os.remove(file_path)

                        # --- ทำ Soft Delete ใน DB ---
                        cursor.execute("""
                            UPDATE tenants 
                            SET is_deleted = 1, id_card = NULL, id_card_img = NULL,
                                address = '-', email = NULL, age = NULL, bd = NULL,tel = NULL, gender = '-',created_at = NOW()
                            WHERE tenant_id = %s
                        """, (tenant_id,))
                        
                        add_audit_log(
                            cursor, 
                            'TENANT', 
                            'DELETE', 
                            f'ลบผู้เช่า ID: {tenant_id}', 
                            session.get('user', {}).get('user_id')
                        )
                        
                        conn.commit()
                        flash("ลบข้อมูลและไฟล์เอกสารเรียบร้อยแล้ว", "success")
                    except Exception as e:
                        conn.rollback()
                        flash(f"เกิดข้อผิดพลาด: {str(e)}", "danger")
        conn.commit()
        return redirect(url_for('manage_tenants'))

    search = request.args.get('search', '')
    cursor.execute("""
        SELECT t.*, 
           (SELECT u.name FROM contracts c 
            JOIN unit u ON c.room_id = u.unit_id 
            WHERE c.tenant_id = t.tenant_id AND c.status IN (1,2,3,4) LIMIT 1) as current_room
        FROM tenants t
        WHERE (t.fname LIKE %s OR t.lname LIKE %s OR t.tel LIKE %s) AND t.is_deleted = 0
        ORDER BY t.tenant_id DESC
    """, (f"%{search}%", f"%{search}%", f"%{search}%"))
    tenants = cursor.fetchall()

    cursor.close()
    conn.close()
    return render_template('manage_tenants.html', tenants=tenants, search=search)

@app.route('/download_id_card/<filename>')
def download_id_card(filename):
    upload_path = os.path.join(app.root_path, 'uploads', 'id_card_tenants')
    try:
        return send_from_directory(upload_path, filename, as_attachment=True)
    except FileNotFoundError:
        abort(404)

@app.route('/download_doc/<path:filename>')
def download_doc(filename):
    upload_path = os.path.join(current_app.root_path, 'uploaded_docs')
    
    return send_from_directory(
        directory=upload_path, 
        path=filename, 
        as_attachment=True 
    )

@app.route('/delete_doc/<path:filename>', methods=['POST'])
@log_activity("DOC", "DELETE", "ลบไฟล์เอกสาร")
def delete_doc(filename):
    filepath = os.path.join('uploaded_docs', filename)
    if os.path.exists(filepath):
        os.remove(filepath)
        flash('ลบไฟล์เรียบร้อยแล้ว', 'success')
    else:
        flash('ไม่พบไฟล์ที่ต้องการลบ', 'danger')
    return redirect(url_for('manage_doc'))

@app.route('/generate_docx/<int:contract_id>')
def generate_docx(contract_id):

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # ข้อมูล business
    cursor.execute("SELECT * FROM business LIMIT 1")
    business = cursor.fetchone()

    # ข้อมูล contract + tenant + room
    cursor.execute("""
        SELECT c.*, 
            t.fname, t.lname, t.age, t.address, t.id_card, t.bd, c.pay_date,
            u.name AS unit_name, u.floor, u.building
        FROM contracts c
        JOIN tenants t ON c.tenant_id = t.tenant_id
        JOIN unit u ON c.room_id = u.unit_id
        WHERE c.contract_id = %s
        LIMIT 1
    """, (contract_id,))
    contract = cursor.fetchone()

    cursor.execute("SELECT setting_key, setting_value FROM settings")
    settings_list = cursor.fetchall()
    settings = {s['setting_key']: s['setting_value'] for s in settings_list}

    # ข้อมูล options
    cursor.execute("""
        SELECT o.name, o.price
        FROM contract_option co
        JOIN `option` o ON co.option_id = o.id
        WHERE co.contract_id = %s
    """, (contract_id,))
    options = cursor.fetchall()

    conn.close()

    if not contract:
        flash('ไม่พบข้อมูลสัญญา', 'danger')
        return redirect(url_for('dashboard'))

    # สร้าง dict สำหรับแทนค่า
    data_dict = prepare_placeholder_data(
        contract, business, options, settings=settings)

    # หา template docx
    template_folder = os.path.join('uploaded_docs', 'templates')
    if not os.path.exists(template_folder):
        os.makedirs(template_folder)
    template_file = next((f for f in os.listdir(template_folder) if f.endswith('.docx')), None)
    if not template_file:
        flash('ไม่พบเทมเพลตเอกสาร', 'warning')
        return redirect(url_for('dashboard'))

    template_path = os.path.join(template_folder, template_file)
    os.makedirs('generated_docs', exist_ok=True)
    output_filename = f"สัญญาห้อง_{contract['unit_name']}.docx"
    output_path = os.path.join('generated_docs', output_filename)

    # แทนค่า placeholder
    replace_placeholders(template_path, data_dict, output_path)

    # บันทึกชื่อไฟล์ใน database
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE contracts SET contracts_file=%s WHERE contract_id=%s",
                   (output_filename, contract_id))
    conn.commit()
    cursor.close()
    conn.close()

    return send_file(output_path, as_attachment=True, download_name=output_filename)

@app.route('/uploaded_docs/<path:filename>')
@log_activity("DOC", "INSERT", "")
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER_DOCS'], filename)

@app.route('/edit_doc/<filename>', methods=['GET', 'POST'])
def edit_doc(filename):
    filepath = os.path.join(app.config['UPLOAD_FOLDER_DOCS'], filename)
    ext = filename.rsplit('.', 1)[1].lower()

    if not os.path.exists(filepath):
        flash('ไม่พบไฟล์ที่ต้องการ', 'danger')
        return redirect(url_for('manage_doc'))

    if ext not in {'txt', 'html', 'md', 'docx'}:
        flash('ไม่รองรับการแก้ไขไฟล์ประเภทนี้', 'danger')
        return redirect(url_for('manage_doc'))

    # ✅ รับข้อมูล HTML ที่ได้จาก CKEditor
    if request.method == 'POST':
        new_content = request.form.get('content')

        try:
            if ext == 'docx':
                # เขียนเนื้อหาใหม่ลง docx
                doc = Document()
                for para in new_content.split('\n'):
                    doc.add_paragraph(para)
                doc.save(filepath)
            else:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(new_content)

            flash(f'บันทึกการแก้ไข {filename} สำเร็จ', 'success')
        except Exception as e:
            flash(f'เกิดข้อผิดพลาดในการบันทึก: {e}', 'danger')

        return redirect(url_for('manage_doc'))

    # ✅ อ่านเนื้อหาเดิมเพื่อนำไปแสดงใน CKEditor
    try:
        if ext == 'docx':
            doc = Document(filepath)
            content = '\n'.join([para.text for para in doc.paragraphs])
        else:
            with open(filepath, 'rb') as f:
                raw = f.read()
            try:
                content = raw.decode('utf-8')
            except UnicodeDecodeError:
                content = raw.decode('windows-874')
    except Exception as e:
        flash(f'ไม่สามารถอ่านไฟล์: {e}', 'danger')
        content = None

    return render_template('edit_doc.html', filename=filename, content=content)


# ---------------------- AGREEMENT ---------------------
@app.route('/leases_uploaded')
def leases_uploaded():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # --- 1. ดึงรายชื่อห้องทั้งหมดเพื่อไปทำ Dropdown ---
    cursor.execute("SELECT unit_id, name FROM unit WHERE is_deleted = 0 ORDER BY name")
    all_units = cursor.fetchall()

    # Mapping สถานะ (เหมือนเดิม)
    status_map = { 1: "Draft...", 2: "Pending...", 3: "Active...", 4: "Expiring...", 5: "Expired", 6: "Moved Out", 7: "Cancelled" }
    status_icons = { 1: "draft.png", 2: "pending.png", 3: "planning-activities.png", 4: "time.png", 5: "expired.png", 6: "moving.png", 7: "cancel-order.png" }

    selected_status = request.args.get('status')
    try:
        selected_status = int(selected_status) if selected_status else None
    except ValueError:
        selected_status = None

    # เปลี่ยนจาก room_filter เป็น selected_room (รับเป็น unit_id)
    selected_room = request.args.get('room') 

    # 2. ปรับ SQL หลัก
    sql = """
        SELECT c.contract_id, c.contract_start, c.contract_end, c.status, c.contracts_file, c.renew_count,
               CONCAT(t.fname, ' ', t.lname) AS tenant_name,
               u.name AS room_number
        FROM contracts c
        JOIN tenants t ON c.tenant_id = t.tenant_id
        JOIN unit u ON c.room_id = u.unit_id
        WHERE 1=1
    """
    params = []

    if selected_status is not None:
        sql += " AND c.status = %s"
        params.append(selected_status)

    if selected_room:
        # กรองด้วย unit_id แทนการ LIKE ชื่อห้องเพื่อความแม่นยำ
        sql += " AND u.unit_id = %s"
        params.append(selected_room)

    sql += " ORDER BY c.contract_start DESC"
    cursor.execute(sql, params)
    contracts = cursor.fetchall()

    # (การจัดการวนลูปสร้าง list leases เหมือนเดิม...)
    leases = []
    for c in contracts:
        leases.append({
            'contract_id': c['contract_id'],
            'tenant_name': c['tenant_name'],
            'room_number': c['room_number'],
            'contract_start': c['contract_start'].strftime('%d-%m-%Y') if c['contract_start'] else '',
            'contract_end': c['contract_end'].strftime('%d-%m-%Y') if c['contract_end'] else '',
            'status': c['status'],
            'contracts_file': c['contracts_file'],
            'renew_count' : c['renew_count']
        })

    conn.close()
    return render_template(
        'leases_uploaded.html',
        leases=leases,
        status_map=status_map,
        status_icons=status_icons,
        selected_status=selected_status,
        selected_room=selected_room, # ส่งค่าที่เลือกกลับไป
        all_units=all_units           # ส่งรายชื่อห้องทั้งหมดไป
    )

@app.route('/delete_contract/<int:contract_id>', methods=['POST'])
@role_required(['admin', 'manager'])
def delete_contract(contract_id):
    # ดึงค่าที่ส่งมาจาก Swal (JavaScript)
    input_password = request.form.get('user_password')
    
    # ดึง ID จาก session['user'] ตามโครงสร้าง Login ของคุณ
    user_data = session.get('user', {})
    u_id = user_data.get('user_id')

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        # 1️⃣ ตรวจสอบรหัสผ่านผู้ใช้งานปัจจุบัน
        cursor.execute("SELECT password FROM user WHERE user_id = %s", (u_id,))
        user_record = cursor.fetchone()

        if not user_record or not check_password_hash(user_record['password'], input_password):
            return jsonify({'status': 'error', 'message': 'รหัสผ่านยืนยันไม่ถูกต้อง!'})

        # 2️⃣ หาไฟล์สัญญาเพื่อลบออกจาก Disk
        cursor.execute("SELECT contracts_file FROM contracts WHERE contract_id = %s", (contract_id,))
        row = cursor.fetchone()

        if row and row['contracts_file']:
            file_name = os.path.basename(row['contracts_file']) 
            file_path = os.path.join(CONTRACTS_FOLDER, file_name) 
            
            if os.path.isfile(file_path):
                os.remove(file_path)

        # 3️⃣ อัปเดต Database ให้ค่าไฟล์เป็น NULL
        cursor.execute("UPDATE contracts SET contracts_file = NULL WHERE contract_id = %s", (contract_id,))

        add_audit_log(
            cursor, 
            'CONTRACT', 
            'DELETE_DOC', 
            f'ลบไฟล์เอกสารออกจากสัญญา ID: {contract_id}', 
            u_id
        )
        conn.commit()
        
        return jsonify({'status': 'success', 'message': 'ลบเอกสารสัญญาเรียบร้อยแล้ว'})

    except Exception as e:
        return jsonify({'status': 'error', 'message': f'เกิดข้อผิดพลาด: {str(e)}'})
    finally:
        conn.close()


# ---------------------- USER SETTINGS ----------------------
@app.route("/user_settings")
def user_settings():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT user_id, username, fname, lname, email, gender, profile_img, role_id, tel FROM user WHERE is_deleted = 0 ORDER BY user_id")
        users = cursor.fetchall()
    except Exception as e:
        flash(f"เกิดข้อผิดพลาด: {e}", 'danger')
        users = []
    finally:
        cursor.close()
        conn.close()

    # ส่ง user ปัจจุบันด้วย เพื่อเช็คไม่ให้แก้ไข หรือลบตัวเองได้ (สมมติ session['user'] = username)
    current_user = session['user']

    return render_template("user_settings.html", users=users, current_user=current_user)

@app.route("/edit_user/<int:user_id>", methods=["GET", "POST"])
def edit_user(user_id):
    current_user_id = session['user']['user_id']
    current_role_id = int(session['user']['role_id'])
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT * FROM user WHERE user_id = %s", (user_id,))
    user_data = cursor.fetchone()

    if not user_data:
        flash('ไม่พบผู้ใช้งานนี้ในระบบ', 'danger')
        return redirect(url_for('user_settings'))

    target_role_id = int(user_data['role_id'])

    if current_role_id == 3 and current_user_id != user_id:
        flash('พนักงาน (Staff) มีสิทธิ์แก้ไขได้เฉพาะข้อมูลของตนเองเท่านั้น', 'warning')
        return redirect(url_for('user_settings'))

    if current_role_id == 2 and target_role_id == 1:
        flash('Manager ไม่สามารถแก้ไขข้อมูลของ Admin ได้', 'danger')
        return redirect(url_for('user_settings'))

    if request.method == 'POST':
        fname = request.form.get('fname', '').strip()
        lname = request.form.get('lname', '').strip()
        email = request.form.get('email', '').strip()
        tel = request.form.get('tel', '').strip()
        role_id = request.form.get('role_id')
        new_password = request.form.get('password', '').strip()

        # --- 1. เช็คอีเมลซ้ำ (ต้องไม่ใช่ ID ตัวเอง) ---
        cursor.execute("SELECT user_id FROM user WHERE email = %s AND user_id != %s", (email, user_id))
        if cursor.fetchone():
            flash('อีเมลนี้ถูกใช้งานโดยผู้ใช้อื่นแล้ว', 'danger')
            return redirect(url_for('edit_user', user_id=user_id))

        try:
            # เตรียม SQL และ Parameter พื้นฐาน
            sql_parts = ["fname=%s", "lname=%s", "email=%s", "tel=%s", "role_id=%s"]
            params = [fname, lname, email, tel, role_id]

            # --- 2. เช็ครหัสผ่านใหม่ (ถ้ามีการกรอกมา) ---
            if new_password:
                if len(new_password) < 8 or not re.search("[0-9]", new_password):
                    flash('รหัสผ่านใหม่ต้องมี 8 ตัวขึ้นไปและมีตัวเลข', 'warning')
                    return redirect(url_for('edit_user', user_id=user_id))

                if not re.search("[A-Z]", new_password):
                    flash('รหัสผ่านต้องมีตัวพิมพ์ใหญ่อย่างน้อย 1 ตัว', 'warning')
                    return redirect(url_for('edit_user', user_id=user_id))
                
                sql_parts.append("password=%s")
                params.append(generate_password_hash(new_password))

            # --- 3. จัดการไฟล์รูปภาพ (บันทึกไฟล์ก่อนแล้วค่อยเก็บชื่อลง DB) ---
            file_profile = request.files.get('profile_img')
            if file_profile and file_profile.filename != '':
                ext = file_profile.filename.rsplit('.', 1)[1].lower()
                p_filename = f"profile_{user_id}_{int(datetime.now().timestamp())}.{ext}"
                file_profile.save(os.path.join(UPLOAD_PROFILE, p_filename))
                sql_parts.append("profile_img=%s")
                params.append(p_filename)

            file_id_card = request.files.get('id_card_file')
            if file_id_card and file_id_card.filename != '':
                ext = file_id_card.filename.rsplit('.', 1)[1].lower()
                c_filename = f"idcard_{user_id}_{int(datetime.now().timestamp())}.{ext}"
                file_id_card.save(os.path.join(UPLOAD_ID_CARD, c_filename))
                sql_parts.append("id_card_file=%s")
                params.append(c_filename)

            final_sql = f"UPDATE user SET {', '.join(sql_parts)} WHERE user_id=%s"
            params.append(user_id)
            cursor.execute(final_sql, tuple(params))
            
            # บันทึกประวัติการแก้ไข
            add_audit_log(cursor, 'USER', 'UPDATE', f'แก้ไขข้อมูลผู้ใช้ ID: {user_id}', session.get('user', {}).get('user_id'))

            conn.commit()
            flash('อัปเดตข้อมูลสำเร็จ', 'success')
            return redirect(url_for('user_settings'))

        except Exception as e:
            conn.rollback()
            flash(f'เกิดข้อผิดพลาด: {str(e)}', 'danger')
        finally:
            cursor.close()
            conn.close()

    # ส่วน GET: ดึงข้อมูลเดิมมาแสดงใน Form
    cursor.execute("SELECT * FROM user WHERE user_id = %s", (user_id,))
    user_data = cursor.fetchone()
    cursor.execute("SELECT * FROM role")
    roles = cursor.fetchall()
    
    cursor.close()
    conn.close()
    return render_template("edit_user.html", user=user_data, roles=roles)

@app.route('/remove_user_file/<file_type>/<int:user_id>', methods=['POST'])
def remove_user_file(file_type, user_id):

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        # 1. ดึงชื่อไฟล์เดิมจาก DB มาดูก่อนว่าไฟล์ชื่ออะไร
        cursor.execute("SELECT profile_img, id_card_file FROM user WHERE user_id = %s", (user_id,))
        user_data = cursor.fetchone()

        if not user_data:
            return jsonify({'status': 'error', 'message': 'ไม่พบข้อมูลผู้ใช้'}), 404

        filename_to_delete = None
        folder_path = ""
        column_name = ""

        # 2. ตรวจสอบประเภทที่จะลบ และตั้งค่า Path
        if file_type == 'profile':
            filename_to_delete = user_data['profile_img']
            folder_path = UPLOAD_PROFILE  # ตรวจสอบว่าตัวแปรนี้เก็บ path เช่น 'static/profile_user'
            column_name = "profile_img"
        elif file_type == 'idcard':
            filename_to_delete = user_data['id_card_file']
            folder_path = UPLOAD_ID_CARD
            column_name = "id_card_file"
        else:
            return jsonify({'status': 'error', 'message': 'ประเภทไฟล์ไม่ถูกต้อง'}), 400

        # 3. ลบไฟล์จริงออกจาก Server (ถ้ามีไฟล์อยู่จริง)
        if filename_to_delete:
            full_path = os.path.join(folder_path, filename_to_delete)
            if os.path.exists(full_path):
                os.remove(full_path) # ลบไฟล์ออกจากโฟลเดอร์

        # 4. อัปเดต Database ให้ค่านั้นเป็น NULL
        query = f"UPDATE user SET {column_name} = NULL WHERE user_id = %s"
        
        add_audit_log(
                cursor, 
                'USER', 
                'DELET_USER_FILE', 
                f'ลบไฟล์ผู้ใช้งาน ID: {user_id}', 
                session.get('user', {}).get('user_id')
            )

        cursor.execute(query, (user_id,))
        conn.commit()

        return jsonify({'status': 'success', 'message': 'ลบไฟล์ออกจากระบบเรียบร้อยแล้ว'})

    except Exception as e:
        conn.rollback()
        return jsonify({'status': 'error', 'message': f'เกิดข้อผิดพลาด: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()

@app.route("/delete_user/<int:user_id>")
def delete_user(user_id):
    current_user_id = session['user']['user_id']
    current_role_id = int(session['user']['role_id'])
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        cursor.execute("SELECT profile_img, id_card_file, role_id FROM user WHERE user_id = %s", (user_id,))
        user_files = cursor.fetchone()

        if not user_files:
            flash('ไม่พบผู้ใช้งานนี้ในระบบ', 'danger')
            return redirect(url_for('user_settings'))

        target_role_id = int(user_files['role_id'])

        if current_role_id == 3:
            if current_user_id != user_id:
                flash('พนักงาน (Staff) มีสิทธิ์ลบได้เฉพาะบัญชีของตนเองเท่านั้น', 'warning')
                return redirect(url_for('user_settings'))

        elif current_role_id == 2:
            if target_role_id == 1: 
                flash('Manager ไม่สามารถลบ Admin ได้', 'danger')
                return redirect(url_for('user_settings'))

        # 2. ลองสั่ง DELETE เลย
        try:
            cursor.execute("DELETE FROM user WHERE user_id = %s", (user_id,))
            add_audit_log(
                cursor, 'USER', 'DELETE', 
                f'ลบผู้ใช้งาน ID: {user_id} แบบถาวร', 
                session.get('user', {}).get('user_id')
            )
            conn.commit()
            
            # ตามไปลบไฟล์ใน Server
            for key in ['profile_img', 'id_card_file']:
                if user_files[key]:
                    folder = UPLOAD_PROFILE if key == 'profile_img' else UPLOAD_ID_CARD
                    path = os.path.join(folder, user_files[key])
                    if os.path.exists(path):
                        os.remove(path)

            if user_id == current_user_id:
                session.clear() 
                flash("บัญชีของคุณถูกลบเรียบร้อยแล้ว ระบบได้ทำการออกจากระบบอัตโนมัติ", "success")
                return redirect(url_for('login')) 
            
            flash(f'ลบผู้ใช้งาน ID: {user_id} และไฟล์ทั้งหมดเรียบร้อยแล้ว', 'success')

        except Exception as db_err:
            # 3. ถ้าลบไม่ได้เพราะติด Foreign Key (Error 1451)
            conn.rollback() 
            # ตรวจสอบว่าเป็น Error เกี่ยวกับ Foreign Key หรือไม่ (MySQL Code 1451)
            if "1451" in str(db_err):
                cursor.execute("UPDATE user SET is_deleted = 1 WHERE user_id = %s", (user_id,))
                add_audit_log(
                    cursor, 'USER', 'DELETE', 
                    f'ปิดการใช้งานผู้ใช้งาน ID: {user_id} (มีข้อมูลเชื่อมโยง)', 
                    session.get('user', {}).get('user_id')
                )
                conn.commit()
                if user_id == current_user_id:
                    session.clear() 
                    flash("บัญชีของคุณถูกลบเรียบร้อยแล้ว ระบบได้ทำการออกจากระบบอัตโนมัติ", "success")
                    return redirect(url_for('login')) 
                
                flash(f'User นี้มีข้อมูลผูกไว้กับระบบอื่น จึงเปลี่ยนเป็น "ปิดการใช้งาน" แทนการลบถาวร', 'info')
            else:
                # Error อื่นๆ ที่ไม่ใช่ FK
                flash(f'ไม่สามารถลบได้: {str(db_err)}', 'danger')

    except Exception as e:
        conn.rollback()
        flash(f'เกิดข้อผิดพลาด: {str(e)}', 'danger')
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('user_settings'))


# ---------------------- PAYMENT ----------------------
@app.route('/confirm_payment/<int:invoice_id>', methods=['GET', 'POST'])
def confirm_payment(invoice_id):
    current_user = session['user']
    payee_id = current_user['user_id']

    today = get_now().date() if get_now else date.today()
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT * FROM business WHERE id = 1")
    biz = cursor.fetchone()

    cursor.execute("""
        SELECT i.*, u.name AS room_name, u.building, u.floor,u.status_id,
            t.fname AS tenant_fname, t.lname AS tenant_lname,
            c.status AS contract_status, c.contract_end 
        FROM invoices i
        JOIN unit u ON i.unit_id = u.unit_id
        LEFT JOIN tenants t ON i.tenant_id = t.tenant_id
        JOIN contracts c ON i.contract_id = c.contract_id
        WHERE i.invoice_id = %s
    """, (invoice_id,))

    invoice = cursor.fetchone()
    print(invoice['contract_status'])

    if not invoice:
        flash("ไม่พบใบแจ้งหนี้", "warning")
        return redirect(url_for('dashboard'))
    
    meter_save = invoice['meter_saved']
    if invoice['invoice_type'] != 'extra_bill':
        if meter_save == 0 or (invoice['current_electricity_reading'] is None and invoice['current_water_reading'] is None):
            flash("ยังไม่กรอกค่ามิเตอร์ ให้ครบ", "warning")
            return redirect(request.referrer)

     # กำหนดวันที่สำหรับแสดงบน invoice
    display_start =  invoice['billing_period_start']
    display_end = invoice['billing_period_end']

    # แปลง string เป็น date เผื่อ DB return เป็น string
    if isinstance(display_start, str):
        display_start = datetime.strptime(display_start, "%Y-%m-%d").date()

    # หลังจากดึง invoice มาแล้ว
    cursor.execute("""
        SELECT COUNT(*) AS prev_invoice_count
        FROM invoices
        WHERE contract_id = %s AND billing_period_start < %s
    """, (invoice['contract_id'], invoice['billing_period_start']))
    row = cursor.fetchone()

    invoice['is_extra_bill'] = (invoice['invoice_type'] == 'extra_bill')
    invoice['is_first_month'] = (invoice['invoice_type'] == 'first')

    if not invoice:
        flash("ไม่พบใบแจ้งหนี้", "warning")
        cursor.close()
        conn.close()
        return redirect(url_for('dashboard'))

    # ดึงรายการค่าใช้จ่ายจาก invoice_items
    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type in ('option','penalty','service')
    """, (invoice_id,))
    items = cursor.fetchall()

    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type = 'meter_adjustment'
    """, (invoice_id,))
    meter_adjustment = cursor.fetchall()

    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type = 'discount'
    """, (invoice_id,))
    discount = cursor.fetchall()
    
    cursor.execute(
        "SELECT setting_value FROM settings WHERE setting_key='late_fee_per_day'")
    row = cursor.fetchone()
    late_fee_per_day = float(row['setting_value']) if row else 100

    overdue_days = update_late_penalty(cursor, invoice_id)
    invoice['overdue_days'] = overdue_days

    if request.method == 'POST':
        payment_method = request.form.get('payment_method')

        file = request.files.get('slip')
        slip_filename = None
        if file and file.filename != '':
            if not os.path.exists(INCOME_UPLOAD_PATH):
                os.makedirs(INCOME_UPLOAD_PATH)
            slip_filename = generate_slip_filename(file, "slip_monthly", invoice_id)
    
            file.save(os.path.join(INCOME_UPLOAD_PATH, slip_filename))
            
        if invoice['contract_status'] == 2 and invoice['status_id'] != 6 and invoice['invoice_type'] != 'extra_bill':
            cursor.execute("""
                UPDATE invoices
                SET status = 'paid', payment_date = NOW(), payment_method = %s, payee = %s, slip_file = %s
                WHERE invoice_id = %s
            """, (payment_method, payee_id, slip_filename, invoice_id))

            cursor.execute("""
                UPDATE unit u
                JOIN invoices i ON u.unit_id = i.unit_id
                SET u.status_id = 2
                WHERE i.invoice_id = %s
            """, (invoice_id,))

            cursor.execute("""
                UPDATE contracts c
                JOIN invoices i ON c.contract_id = i.contract_id
                SET c.status = 3
                WHERE i.invoice_id = %s
            """, (invoice_id,))

            record_transaction(
                cursor, 
                amount=invoice['total_amount'],
                t_type='income',
                category='ค่าเช่ารายเดือน',
                ref_invoice_id=invoice_id,
                note=f"รับชำระค่าเช่าห้อง {invoice['room_name']}",
                created_by=payee_id
            )

        elif invoice['contract_status'] == 3 and (invoice['status_id'] != 6 or invoice['status_id'] == 7) and invoice['invoice_type'] != 'extra_bill':
            cursor.execute("""
                UPDATE invoices
                SET status = 'paid', payment_date = NOW(), payment_method = %s, payee = %s, slip_file = %s
                WHERE invoice_id = %s
            """, (payment_method, payee_id, slip_filename, invoice_id))

            cursor.execute("""
                UPDATE contracts c
                JOIN invoices i ON c.contract_id = i.contract_id
                SET c.status = 3
                WHERE i.invoice_id = %s
            """, (invoice_id,))

            record_transaction(
                cursor, 
                amount=invoice['total_amount'],
                t_type='income',
                category='ค่าเช่ารายเดือน',
                ref_invoice_id=invoice_id,
                note=f"รับชำระค่าเช่าห้อง {invoice['room_name']}",
                created_by=payee_id
            )

        elif invoice['invoice_type'] == 'final' or invoice['contract_status'] == 4 or invoice['status_id'] == 6 :
            # กรณีหมดอายุแล้ว → ห้องว่าง, สัญญาปิด
            cursor.execute("""
                UPDATE invoices
                SET status = 'paid', payment_date = NOW(), payment_method = %s, payee = %s,slip_file = %s
                WHERE invoice_id = %s
            """, (payment_method, payee_id, slip_filename, invoice_id))

            if invoice['total_amount'] < 0:
                type='expense'
                total_amount = invoice['total_amount'] * -1
                t_note=f"ค่ารายจ่ายเพิ่มเติม ห้อง {invoice['room_name']}"
            elif invoice['total_amount'] >= 0:
                type='income'
                total_amount = invoice['total_amount']
                t_note=f"ค่าบริการเพิ่มเติม ห้อง {invoice['room_name']}"

            cursor.execute("""
                UPDATE unit
                SET status_id = 6
                WHERE unit_id = %s
            """, (invoice['unit_id'],))
            
            record_transaction(
                cursor, 
                amount=total_amount,
                t_type=type,
                category='ค่าเช่าบิดสุดท้าย',
                ref_invoice_id=invoice_id,
                note=t_note,
                created_by=payee_id
            )

        elif invoice['invoice_type'] == 'extra_bill':
            cursor.execute("""
                UPDATE invoices
                SET status = 'paid', payment_date = NOW(), payment_method = %s, payee = %s, slip_file = %s
                WHERE invoice_id = %s
            """, (payment_method, payee_id, slip_filename, invoice_id))

            if invoice['total_amount'] < 0:
                type='expense'
                total_amount = invoice['total_amount'] * -1
                t_note=f"รายจ่ายบิลเพิ่มเติมจาก ห้อง {invoice['room_name']}"
            elif invoice['total_amount'] >= 0:
                type='income'
                t_note=f"ค่าบริการเพิ่มเติม ห้อง {invoice['room_name']}"

            record_transaction(
                cursor, 
                amount=invoice['total_amount'],
                t_type=type,
                category='บิลเพิ่มเติม',
                ref_invoice_id=invoice_id,
                note=t_note,
                created_by=payee_id
            )

        else:
            flash("ไม่พบใบแจ้งหนี้", "warning")
            cursor.close()
            conn.close()
            return redirect(url_for('dashboard'))

        add_audit_log(
            cursor, 
            'INVOICE', 
            'PAYMENT', 
            f'ชำระเงินใบแจ้งหนี้ ID: {invoice_id}', 
            session.get('user', {}).get('user_id')
        )

        conn.commit()
        cursor.close()
        conn.close()

        flash("ชำระเงินเรียบร้อยแล้ว", "success")
        return redirect(url_for('print_receipt', invoice_id=invoice_id))

    cursor.close()
    conn.close()

    return render_template('confirm_payment.html', invoice=invoice, items=items, today=today, late_fee_per_day=late_fee_per_day, discount=discount, display_start=display_start,
    display_end=display_end,biz=biz,meter_adjustment=meter_adjustment)

@app.route('/confirm_daily_payment/<int:invoice_id>', methods=['GET', 'POST'])
def confirm_daily_payment(invoice_id):

    current_user = session['user']
    payee_id = current_user['user_id']

    today = get_now or date.today()
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT * FROM business WHERE id = 1")
    biz = cursor.fetchone()

    # ดึง invoice รายวัน
    cursor.execute("""
        SELECT i.*, u.name AS room_name, u.building, u.floor
        FROM invoices i
        JOIN unit u ON i.unit_id = u.unit_id
        WHERE i.invoice_id = %s AND (i.invoice_type='daily' OR i.invoice_type='extra_bill')
    """, (invoice_id,))
    invoice = cursor.fetchone()

    cursor.execute("""
        SELECT t.price_daily
        FROM type t
        JOIN unit u ON u.type_unit_id = t.type_id
        JOIN invoices i ON i.unit_id = u.unit_id
        WHERE i.invoice_id = %s
    """, (invoice_id,))
    price_daily_list = cursor.fetchall()
    price_daily = price_daily_list[0]['price_daily'] if price_daily_list else 0

     # ดึงรายการค่าใช้จ่ายจาก invoice_items
    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type in ('option','penalty','service')
    """, (invoice_id,))
    items = cursor.fetchall()

    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type = 'discount'
    """, (invoice_id,))
    discount = cursor.fetchall()

    cursor.execute(
        "SELECT setting_value FROM settings WHERE setting_key='late_fee_per_day'")
    row = cursor.fetchone()
    late_fee_per_day = float(row['setting_value']) if row else 100

    overdue_days = update_late_penalty(cursor, invoice_id)
    invoice['overdue_days'] = overdue_days

    print(invoice['billing_period_start'])

    if not invoice:
        flash("ไม่พบใบแจ้งหนี้รายวันนี้", "warning")
        cursor.close()
        conn.close()
        return redirect(url_for('dashboard'))

    meter_save = invoice['meter_saved']
    if invoice['invoice_type'] != 'extra_bill':
        if meter_save == 0 or (invoice['current_electricity_reading'] is None and invoice['current_water_reading'] is None):
            flash("ยังไม่กรอกค่ามิเตอร์ ให้ครบ", "warning")
            return redirect(request.referrer)

    if request.method == 'POST':
        payment_method = request.form.get('payment_method', 'cash')
        file = request.files.get('slip')
        slip_filename = None

        # --- 📂 ส่วนบันทึกไฟล์สลิป ---
        if file and file.filename != '':
            if not os.path.exists(INCOME_UPLOAD_PATH):
                os.makedirs(INCOME_UPLOAD_PATH)
            slip_filename = generate_slip_filename(file, "slip_daily", invoice_id)
    
            file.save(os.path.join(INCOME_UPLOAD_PATH, slip_filename))

        try:
            # 1. อัปเดตสถานะบิล และชื่อไฟล์สลิป (ถ้ามี)
            cursor.execute("""
                UPDATE invoices
                SET status='paid', payment_date=NOW(), 
                    payment_method=%s, slip_file=%s
                WHERE invoice_id=%s
            """, (payment_method, slip_filename, invoice_id))

            # 2. อัปเดตสถานะห้องเป็น 'มีผู้เช่า' (occupied = 2)
            cursor.execute("UPDATE unit SET status_id=2 WHERE unit_id=%s", (invoice['unit_id'],))

            # 3. บันทึกบัญชีรายรับ
            record_transaction(
                cursor, 
                amount=invoice['total_amount'],
                t_type='income',
                category='ค่าเช่ารายวัน',
                ref_invoice_id=invoice_id,
                note=f"รับชำระค่าเช่ารายวัน ห้อง {invoice['room_name']}",
                created_by=payee_id
            )

            add_audit_log(
                cursor, 
                'INVOICE', 
                'PAYMENT_DAILY', 
                f'ชำระเงินใบแจ้งหนี้รายวัน ID: {invoice_id}', 
                session.get('user', {}).get('user_id')
            )

            conn.commit()
            flash("ชำระเงินเรียบร้อยแล้ว", "success")
            return redirect(url_for('print_receipt', invoice_id=invoice_id))

        except Exception as e:
            conn.rollback()
            flash(f"เกิดข้อผิดพลาด: {str(e)}", "danger")

    cursor.close()
    conn.close()
    return render_template('confirm_daily_payment.html', invoice=invoice, today=today, price_daily=price_daily ,
                           items=items ,discount=discount ,late_fee_per_day=late_fee_per_day,biz=biz)

@app.route('/add_expense', methods=['POST'])
def add_expense():
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        current_now = get_now().strftime('%Y-%m-%d %H:%M:%S')
        file = request.files.get('slip_file')
        filename = None
        if file and file.filename != '':
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            original_name = secure_filename(file.filename)
            
            filename = f"{timestamp}_{original_name}"
        
            file.save(os.path.join(EXPENSE_UPLOAD_PATH, filename))

        expense_date_raw = request.form.get('expense_date')
        try:
            expense_date = datetime.strptime(expense_date_raw, '%d/%m/%Y').strftime('%Y-%m-%d')
        except:
            expense_date = datetime.now().strftime('%Y-%m-%d')
        
        category = request.form.get('category')
        description = request.form.get('description')
        amount = float(request.form.get('amount'))
        created_by = session['user']['user_id']

        query = """
            INSERT INTO expenses (expense_date, category, description, amount, created_by, created_at, slip_file)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """
        cursor.execute(query, (expense_date, category, description, amount, created_by, current_now, filename))
        expense_id = cursor.lastrowid

        record_transaction(
                cursor, 
                amount=amount,
                t_type='expense',
                category=category,
                ref_expense_id=expense_id,
                note=f"บันทึกรายจ่าย {category}",
                created_by=created_by
            )

        add_audit_log(
            cursor,
            'EXPENSE',
            'INSERT',
            f'บันทึกรายจ่ายใหม่ ID: {expense_id}, หมวดหมู่: {category}, จำนวน: {amount}',
            created_by
        )
        conn.commit()
        
        return jsonify({"status": "success", "message": "บันทึกรายจ่ายแล้ว"})
    
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/delete_expense/<int:expense_id>', methods=['POST'])
def delete_expense(expense_id):
    if session.get('role') not in ['admin', 'manager']:
        return jsonify({"status": "error", "message": "สิทธิ์ไม่เพียงพอ"}), 403

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        cursor.execute("SELECT * FROM expenses WHERE expense_id = %s", (expense_id,))
        expense = cursor.fetchone()

        if not expense:
            return jsonify({"status": "error", "message": "ไม่พบข้อมูลรายจ่าย"}), 404

        if expense.get('slip_file'):
            file_to_delete = os.path.join(EXPENSE_UPLOAD_PATH, expense['slip_file'])
            try:
                if os.path.exists(file_to_delete):
                    os.remove(file_to_delete)
                    print(f"ลบไฟล์สำเร็จ: {file_to_delete}")
                else:
                    print(f"ไม่พบไฟล์บน Server: {file_to_delete}")
            except Exception as e:
                print(f"เกิดข้อผิดพลาดขณะลบไฟล์: {e}")

        cursor.execute("DELETE FROM transactions WHERE ref_expense_id = %s", (expense_id,))
        cursor.execute("DELETE FROM expenses WHERE expense_id = %s", (expense_id,))

        # 4. บันทึกประวัติ
        add_audit_log(
            cursor,
            'EXPENSE',
            'DELETE',
            f'ลบรายจ่าย ID: {expense_id}, หมวด: {expense["category"]}, ยอด: {expense["amount"]}',
            session['user']['user_id']
        )
        
        conn.commit()
        return jsonify({"status": "success", "message": "ลบข้อมูลและไฟล์หลักฐานเรียบร้อยแล้ว"})
    
    except Exception as e:
        conn.rollback() # ย้อนกลับหากลบไม่สำเร็จ
        print(f"Error deleting expense: {e}")
        return jsonify({"status": "error", "message": "เกิดข้อผิดพลาดภายในระบบ"}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/billing')
def billing():
    if session.get('role') not in ['admin', 'manager']:
        flash("คุณไม่มีสิทธิ์เข้าถึงหน้านี้", "danger")
        return redirect(url_for('dashboard'))

    search_room = request.args.get('room', '')
    start_date = request.args.get('start_date', '')
    end_date = request.args.get('end_date', '')

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    data_list = []

    is_date_filtered = bool(start_date or end_date)

    # --- ส่วนที่ 1: ดึงรายรับ (Invoices) ---
    inv_sql = """
        SELECT 
            i.invoice_id as id, 
            i.created_at as date, 
            'income' as type, 
            u.name as room, 
            i.total_amount as amount, 
            i.status, 
            i.invoice_type,
            CONCAT(IFNULL(t.fname, i.guest_fname), ' ', IFNULL(t.lname, i.guest_lname)) as name,
            i.slip_file
        FROM invoices i
        JOIN unit u ON i.unit_id = u.unit_id
        LEFT JOIN tenants t ON i.tenant_id = t.tenant_id
        WHERE 1=1
    """
    inv_params = []
    if search_room:
        inv_sql += " AND u.name LIKE %s"
        inv_params.append(f"%{search_room}%")
    if start_date:
        inv_sql += " AND i.created_at >= %s"
        inv_params.append(f"{start_date} 00:00:00")
    if end_date:
        inv_sql += " AND i.created_at <= %s"
        inv_params.append(f"{end_date} 23:59:59")

    if not is_date_filtered:
        inv_sql += " ORDER BY i.created_at DESC LIMIT 50"

    cursor.execute(inv_sql, inv_params)
    data_list.extend(cursor.fetchall())

    # --- ส่วนที่ 2: ดึงรายจ่าย (Expenses) ---
    exp_sql = """
        SELECT 
            expense_id as id, 
            created_at as date, 
            'expense' as type,
            category as room, 
            amount, 
            'expense' as status, 
            description as name,
            slip_file
        FROM expenses
        WHERE 1=1
    """
    exp_params = []
    if search_room:
        exp_sql += " AND (category LIKE %s OR description LIKE %s)"
        exp_params.extend([f"%{search_room}%", f"%{search_room}%"])
    if start_date:
        exp_sql += " AND expense_date >= %s"
        exp_params.append(start_date)
    if end_date:
        exp_sql += " AND expense_date <= %s"
        exp_params.append(end_date)

    if not is_date_filtered:
        exp_sql += " ORDER BY created_at DESC LIMIT 50"

    cursor.execute(exp_sql, exp_params)
    data_list.extend(cursor.fetchall())

    def sort_key(x):
        d = x.get('date')
        if d is None:
            return datetime(1900, 1, 1).date() 
        if isinstance(d, datetime):
            return d.date()
        return d

    data_list.sort(key=sort_key, reverse=True)

    if not is_date_filtered:
        data_list = data_list[:50]
    cursor.close()
    conn.close()

    return render_template('billing.html', 
                           invoices=data_list, 
                           search_room=search_room, 
                           start_date=start_date, 
                           end_date=end_date)


# ---------------------- print_receipt ----------------------
@app.route('/print_receipt/<int:invoice_id>')
def print_receipt(invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT i.*, 
            u.name AS room_name, u.building, u.floor, ty.price_daily,
            t.fname AS tenant_fname, t.lname AS tenant_lname
        FROM invoices i
        LEFT JOIN unit u ON i.unit_id = u.unit_id
        LEFT JOIN type ty ON u.type_unit_id = ty.type_id
        LEFT JOIN tenants t ON i.tenant_id = t.tenant_id
        WHERE i.invoice_id = %s
    """, (invoice_id,))
    invoice = cursor.fetchone()
    

    display_start = invoice['billing_period_start']
    display_end =  invoice['billing_period_end']

    # หลังจากดึง invoice มาแล้ว
    cursor.execute("""
        SELECT COUNT(*) AS prev_invoice_count
        FROM invoices
        WHERE contract_id = %s AND billing_period_start < %s
    """, (invoice['contract_id'], invoice['billing_period_start']))

    row = cursor.fetchone()
    
    invoice['is_extra_bill'] = (invoice['invoice_type'] == 'extra_bill')
    invoice['is_first_month'] = (invoice['invoice_type'] == 'first')
    invoice['is_daily'] = (invoice['invoice_type'] == 'daily')

    if not invoice:
        flash("ไม่พบใบเสร็จ", "warning")
        cursor.close()
        conn.close()
        return redirect(url_for('dashboard'))

     # ดึงรายการค่าใช้จ่ายจาก invoice_items
    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type in ('option','penalty','service')
    """, (invoice_id,))
    items = cursor.fetchall()

    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type = 'meter_adjustment'
    """, (invoice_id,))
    meter_adjustment = cursor.fetchall()

    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type = 'discount'
    """, (invoice_id,))
    discount = cursor.fetchall()

    # ดึง late_fee_per_day
    cursor.execute(
        "SELECT setting_value FROM settings WHERE setting_key='late_fee_per_day'")
    row = cursor.fetchone()
    late_fee_per_day = float(row['setting_value']) if row else 100

    overdue_days = update_late_penalty(cursor, invoice_id)
    invoice['overdue_days'] = overdue_days

    cursor.close()
    conn.close()
    return render_template('print_receipt.html', invoice=invoice, items=items, late_fee_per_day=late_fee_per_day, 
                           discount=discount,display_start=display_start,
                           display_end=display_end,meter_adjustment=meter_adjustment)


# ---------------------- CONTRACT ----------------------
@app.route('/confirm_contract/<int:contract_id>', methods=['GET', 'POST'])
def confirm_contract(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        file = request.files.get('file')
        if file and allowed_file(file.filename):
            # ใช้ f-string จัดการชื่อไฟล์ให้สื่อความหมาย
            filename = f"signed_{contract_id}_{secure_filename(file.filename)}"
            
            # ✅ เปลี่ยนมาใช้ CONTRACTS_FOLDER ที่เราสร้างไว้ใน PATHS
            filepath = os.path.join(CONTRACTS_FOLDER, filename)
            file.save(filepath)

            try:
                # อัปเดตชื่อไฟล์ลงใน Database
                cursor.execute(
                    "UPDATE contracts SET contracts_file = %s, status = 2 WHERE contract_id = %s", 
                    (filename, contract_id)
                )
                
                # 🟢 เพิ่ม Audit Log บันทึกประวัติ
                add_audit_log(
                    cursor, 
                    'CONTRACT', 
                    'CONFIRM', 
                    f'ยืนยันสัญญาและอัปโหลดไฟล์ (ID: {contract_id})', 
                    session.get('user', {}).get('user_id')
                )
                
                conn.commit()
                flash("ยืนยันสัญญาและอัปโหลดไฟล์สำเร็จ", "success")
            except Exception as e:
                conn.rollback()
                flash(f"เกิดข้อผิดพลาด: {e}", "danger")
        else:
            flash("กรุณาเลือกไฟล์ให้ถูกต้อง", "danger")

        return redirect(url_for('confirm_contract', contract_id=contract_id))

    # GET - Load data
    cursor.execute("SELECT * FROM contracts WHERE contract_id = %s", (contract_id,))
    contract = cursor.fetchone()
    
    cursor.close()
    conn.close()
    return render_template('confirm_contract.html', contract=contract)

@app.route('/approve_contract/<int:contract_id>', methods=['POST'])
def approve_contract(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE contracts SET status = 2 WHERE contract_id = %s", (contract_id,))
    
    add_audit_log(
        cursor, 'CONTRACT', 'APPROVE', 
        f'อนุมัติสัญญา ID: {contract_id}', 
        session.get('user', {}).get('user_id')
    )

    conn.commit()
    cursor.close()
    conn.close()
    flash("ยืนยันสัญญาเรียบร้อย", "success")
    return redirect(url_for('dashboard'))

@app.route('/upload_signed_contract/<int:contract_id>', methods=['POST'])
def upload_signed_contract(contract_id):
    if not request.files:
        flash('กรุณาเลือกไฟล์', 'danger')
        return redirect(url_for('dashboard'))

    file = next(iter(request.files.values()))
    if file.filename == '':
        flash('กรุณาเลือกไฟล์', 'danger')
        return redirect(url_for('dashboard'))

    # 1. จัดการชื่อไฟล์และ Path (ใช้ตัวแปรกลางที่คุณตั้งไว้)
    now = datetime.now()
    filename = f"signed_{contract_id}_{now.strftime('%d-%m-%Y')}_{secure_filename(file.filename)}"
    
    # ✅ ใช้ CONTRACTS_FOLDER ที่ประกาศไว้ตอนเริ่มแอป
    filepath = os.path.join(CONTRACTS_FOLDER, filename)
    
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True) # ใช้ dictionary เพื่อความง่าย

    try:
        # บันทึกไฟล์ลง Disk
        file.save(filepath)

        # 2. อัปเดตสถานะสัญญา
        cursor.execute("""
            UPDATE contracts 
            SET contracts_file = %s, status = 2, created_at = NOW()
            WHERE contract_id = %s
        """, (filename, contract_id))

        # 3. ดึง room_id และอัปเดตสถานะห้อง
        cursor.execute("SELECT room_id FROM contracts WHERE contract_id = %s", (contract_id,))
        contract_data = cursor.fetchone()
        
        if contract_data:
            room_id = contract_data['room_id']
            cursor.execute("UPDATE unit SET status_id = 5 WHERE unit_id = %s", (room_id,))

        # 4. 🟢 เพิ่ม Audit Log (ใช้ add_audit_log ที่เราคุยกัน)
        add_audit_log(
            cursor, 
            'CONTRACT', 
            'UPLOAD_SIGNED', 
            f'อัปโหลดสัญญาเซ็นแล้ว ห้อง ID: {room_id if "room_id" in locals() else "Unknown"}', 
            session.get('user', {}).get('user_id')
        )

        conn.commit()
        flash('✅ อัปโหลดสัญญาเซ็นแล้วสำเร็จ', 'success')

    except Exception as e:
        if conn: conn.rollback()
        # ถ้าพังและไฟล์ถูกเซฟไปแล้ว อาจจะพิจารณาลบไฟล์ทิ้งเพื่อความสะอาด
        if os.path.exists(filepath): os.remove(filepath)
        flash(f'เกิดข้อผิดพลาด: {str(e)}', 'danger')
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

    return redirect(url_for('dashboard'))

@app.route('/download_signed_contract/<filename>')
def download_signed_contract(filename):
    try:
        safe_filename = os.path.basename(filename)
        
        return send_from_directory(
            CONTRACTS_FOLDER, 
            safe_filename, 
            as_attachment=True
        )
    except FileNotFoundError:
        abort(404)

@app.route('/save_contract', methods=['POST'])
def save_contract():
    conn = get_db_connection()
    cursor = conn.cursor()

    unit_id = request.form['unit_id']
    tenant_id = request.form.get('tenant_id')  # กรณีคุณส่งจาก hidden input
    start_date = request.form['start_date']
    end_date = request.form['end_date']
    monthly_price = request.form['monthly_price']
    deposit = request.form.get('security_deposit', 0)
    template_id = request.form['template_id']
    option_ids = request.form.getlist('option_ids')  # list[] จาก checkbox

    try:
        # 1. Insert into contracts
        cursor.execute("""
            INSERT INTO contracts 
                (tenant_id, room_id, start_date, end_date, monthly_price, deposit, template_id, status, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, 'รอการชำระ', %s)
        """, (
            tenant_id, unit_id, start_date, end_date, monthly_price,
            deposit, template_id, datetime.now()
        ))

        contract_id = cursor.lastrowid  # ดึง contract_id ที่เพิ่มล่าสุด

        # 2. Insert options
        for opt_id in option_ids:
            cursor.execute("""
                INSERT INTO contract_option (contract_id, option_id)
                VALUES (%s, %s)
            """, (contract_id, opt_id))

        conn.commit()
        flash("บันทึกสัญญาเรียบร้อย", "success")
        return redirect(url_for('preview_contract', contract_id=contract_id))

    except Exception as e:
        conn.rollback()
        flash(f"เกิดข้อผิดพลาด: {e}", "danger")
        return redirect(url_for('dashboard'))

    finally:
        cursor.close()
        conn.close()
       
@app.route('/contracts/<filename>')
def contracts(filename):
    safe_filename = os.path.basename(filename)
    try:
        return send_from_directory(CONTRACTS_FOLDER, safe_filename)
    except FileNotFoundError:
        abort(404)

@app.route('/edit_signed_contract/<int:contract_id>', methods=['POST'])
def edit_signed_contract(contract_id):
    # 1. เช็คว่ามีไฟล์ส่งมาไหม
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': 'กรุณาเลือกไฟล์'})

    file = request.files['file']
    if file.filename == '':
        return jsonify({'status': 'error', 'message': 'กรุณาเลือกไฟล์'})

    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        # 2. ดึงไฟล์เดิมเพื่อลบออก (ป้องกันไฟล์ขยะเต็ม Server)
        cursor.execute("SELECT contracts_file FROM contracts WHERE contract_id=%s", (contract_id,))
        contract = cursor.fetchone()

        if contract and contract['contracts_file']:
            old_files = contract['contracts_file'].split(',')
            for f_name in old_files:
                # ✅ ใช้ CONTRACTS_FOLDER ที่เป็น Absolute Path
                old_path = os.path.join(CONTRACTS_FOLDER, f_name.strip())
                if os.path.exists(old_path):
                    os.remove(old_path)

        old_file = contract['contracts_file'] if contract else None
        folder = 'contracts_file'

        if old_file:
            old_path = os.path.join(folder, old_file)
            if os.path.exists(old_path):
                os.remove(old_path)

        # 3. เซฟไฟล์ใหม่ (ตั้งชื่อให้ไม่ซ้ำ)
        now = datetime.now()
        filename = f"contract_signed_{contract_id}_{now.strftime('%d-%m-%Y')}_{secure_filename(file.filename)}"
        os.makedirs(folder, exist_ok=True)
        file.save(os.path.join(folder, filename))

        # 4. อัปเดตชื่อไฟล์ใน Database
        cursor.execute("UPDATE contracts SET contracts_file = %s WHERE contract_id = %s", (filename, contract_id))

        add_audit_log(
            cursor, 
            'CONTRACT',
            'UPDATE',
            f'แก้ไขไฟล์สัญญาเซ็นแล้ว ID: {contract_id}', 
            session.get('user', {}).get('user_id')
        )

        conn.commit()
        
        return jsonify({
            'status': 'success', 
            'message': 'อัปโหลดไฟล์สัญญาใหม่เรียบร้อยแล้ว'
        })

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})
    finally:
        if 'conn' in locals():
            conn.close()


# ---------------------- SETTINGS ----------------------
@app.route('/settings', methods=['GET', 'POST'])
@role_required(['admin','manager'])
def update_settings():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == 'POST':
        keys = [
            'invoice_due_day', 'late_fee_per_day', 'electricity_rate','water_rate',
            'receipt_page_limit',
            'min_lease_months', 'min_stay_before_early_move',
            'auto_delete_contract_files','auto_generate_bill','enable_late_penalty'
        ]

        for key in keys:
            value = request.form.get(key)

            # ถ้าเป็น toggle และไม่ได้ถูกติ๊ก → ตั้งค่าเป็น 0
            if key == 'auto_delete_contract_files' and value != '1':
                value = '0'

            if key == 'auto_generate_bill' and value != '1':
                value = '0'

            if key == 'enable_late_penalty' and value != '1':
                value = '0'

            # ✅ แปลงค่าเป็น string และกำหนด default ถ้าเป็น None
            if value is None or value.strip() == '':
                value = '0'  # หรือค่า defaultที่คุณต้องการ

            cursor.execute("""
                UPDATE settings SET setting_value=%s WHERE setting_key=%s
            """, (value, key))

        add_audit_log(
            cursor, 
            'SETTINGS', 
            'UPDATE',
            f'อัพเดท setting ในระบบ',
            session.get('user', {}).get('user_id')
        )
        conn.commit()
        flash('บันทึกการตั้งค่าสำเร็จ ✅', 'success')
        return redirect(url_for('update_settings'))

    # GET: ดึงค่าปัจจุบันมาแสดง
    cursor.execute("SELECT setting_key, setting_value FROM settings")
    rows = cursor.fetchall()
    settings = {row['setting_key']: row['setting_value'] for row in rows}

    cursor.close()
    conn.close()
    return render_template('settings.html', settings=settings)

@app.route('/print_invoice/<int:invoice_id>')
def print_invoice(invoice_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # ดึงข้อมูล invoice + ห้อง + ผู้เช่า
    cursor.execute("""
        SELECT i.*, u.name AS room_name, u.building, u.floor,
               t.fname AS tenant_fname, t.lname AS tenant_lname
        FROM invoices i
        JOIN unit u ON i.unit_id = u.unit_id
        LEFT JOIN tenants t ON i.tenant_id = t.tenant_id
        WHERE i.invoice_id = %s
    """, (invoice_id,))
    invoice = cursor.fetchone()

    # กำหนดวันที่สำหรับแสดงบน invoice
    display_start = invoice['billing_period_start']
    display_end = invoice['billing_period_end']

    # หลังจากดึง invoice มาแล้ว
    cursor.execute("""
        SELECT COUNT(*) AS prev_invoice_count
        FROM invoices
        WHERE contract_id = %s AND billing_period_start < %s
    """, (invoice['contract_id'], invoice['billing_period_start']))

    row = cursor.fetchone()
    
    invoice['is_extra_bill'] = (invoice['invoice_type'] == 'extra_bill')
    invoice['is_first_month'] = (invoice['invoice_type'] == 'first')
    invoice['is_daily'] = (invoice['invoice_type'] == 'daily')

    if not invoice:
        flash("ไม่พบใบเสร็จ", "warning")
        return redirect(url_for('dashboard'))
    
    if invoice['invoice_type'] != 'extra_bill':
        meter_save = invoice['meter_saved']
        if meter_save == 0 or (invoice['current_electricity_reading'] is None and invoice['current_water_reading'] is None):
            flash("ยังไม่กรอกค่ามิเตอร์ ให้ครบ", "warning")
            return redirect(request.referrer)
    
    # กำหนดชื่อผู้เช่า/guest ตามประเภท invoice
    if invoice['invoice_type'] == 'daily':
        tenant_name = f"{invoice.get('guest_fname','')} {invoice.get('guest_lname','')}".strip(
        )
    else:
        tenant_name = f"{invoice.get('tenant_fname','')} {invoice.get('tenant_lname','')}".strip(
        )

    # ดึงรายการค่าใช้จ่ายจาก invoice_items
    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type in ('option','penalty','service')
    """, (invoice_id,))
    items = cursor.fetchall()

    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type = 'meter_adjustment'
    """, (invoice_id,))
    meter_adjustment = cursor.fetchall()


    cursor.execute("""
        SELECT description, unit_price, quantity, total_price, type
        FROM invoice_items
        WHERE invoice_id = %s and type = 'discount'
    """, (invoice_id,))
    discount = cursor.fetchall()

    cursor.execute(
        "SELECT setting_value FROM settings WHERE setting_key='late_fee_per_day'")
    row = cursor.fetchone()
    late_fee_per_day = float(row['setting_value']) if row else 100

    cursor.execute("""
        SELECT t.price_daily
        FROM type t
        JOIN unit u ON u.type_unit_id = t.type_id
        JOIN invoices i ON i.unit_id = u.unit_id
        WHERE i.invoice_id = %s
    """, (invoice_id,))
    price_daily_list = cursor.fetchall()
    price_daily = price_daily_list[0]['price_daily'] if price_daily_list else 0

    overdue_days = update_late_penalty(cursor, invoice_id)
    invoice['overdue_days'] = overdue_days
    e_rate = float(get_setting('electricity_rate', 7))
    w_rate = float(get_setting('water_rate', 18))

    cursor.close()
    conn.close()

    return render_template(
        "print_invoice.html",
        invoice=invoice,
        items=items,
        discount=discount,
        tenant_name=tenant_name,
        e_rate=e_rate,
        w_rate=w_rate,
        late_fee_per_day=late_fee_per_day,
        price_daily=price_daily,
        display_start=display_start,
        display_end=display_end,
        meter_adjustment=meter_adjustment
    )


# ---------------------- NOTICE MOVE OUT ----------------------
@app.route('/notice_move_out/<int:contract_id>', methods=['POST'])
def notice_move_out(contract_id):
    notice_date_str = request.form.get('notice_date')  # รับ string จากฟอร์ม
    try:
        notice_date = datetime.strptime(notice_date_str, '%d-%m-%Y').date()
    except ValueError:
        flash("❌ วันที่ไม่ถูกต้อง", "danger")
        return redirect(url_for('dashboard'))

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        UPDATE contracts 
        SET notice_move_out_date = %s
        WHERE contract_id = %s
    """, (notice_date, contract_id))

    cursor.execute("""
        UPDATE unit 
        SET status_id = 7 
        WHERE unit_id = (SELECT room_id FROM contracts WHERE contract_id = %s)
    """, (contract_id,))

    add_audit_log(
        cursor, 
        'TENANT', 
        'NOTICE_MOVE_OUT', 
        f'แจ้งย้ายออกสัญญา ID: {contract_id}, วันที่แจ้ง: {notice_date_str}', 
        session.get('user', {}).get('user_id')
    )

    conn.commit()
    cursor.close()
    conn.close()
    flash("แจ้งย้ายออกเรียบร้อย", "warning")
    return redirect(url_for('dashboard'))

@app.route('/create_invoice_move_out/<int:contract_id>', methods=['POST'])
def create_invoice_move_out(contract_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)   

    try:
        # ดึงข้อมูลสัญญา + ห้อง + ผู้เช่า
        cursor.execute("""
            SELECT c.contract_id, c.price, c.premiums, c.room_id, c.electricity_start, c.water_start,
                   t.tenant_id, t.fname, t.lname, c.notice_move_out_date, c.contract_start, c.contract_end
            FROM contracts c
            JOIN tenants t ON c.tenant_id = t.tenant_id
            WHERE c.contract_id = %s
        """, (contract_id,))
        contract = cursor.fetchone()

        if not contract:
            flash("ไม่พบสัญญา", "warning")
            return redirect(url_for('dashboard'))

        move_out_date = contract['notice_move_out_date']   
        print(move_out_date)
        print(today)

        if today < move_out_date:
            flash(f"❌ ยังไม่ถึงกำหนดวันแจ้งย้ายออก (กำหนดคือ {move_out_date.strftime('%d/%m/%Y')})", "danger")
            return redirect(url_for('dashboard'))

        # 3. ดึง invoice ล่าสุดที่ชำระแล้ว
        cursor.execute("""
            SELECT billing_period_end
            FROM invoices
            WHERE contract_id=%s AND status in ('paid','overdue')
            ORDER BY billing_period_end DESC
            LIMIT 1
        """, (contract_id,))
        last_invoice_paid = cursor.fetchone()

        if last_invoice_paid:
            start_date = last_invoice_paid['billing_period_end'] + timedelta(days=1)
        else:
            start_date = contract['contract_start']

        # 4. ดึง electricity_rate, water_rate
        cursor.execute("""
            SELECT setting_key, setting_value
            FROM settings
            WHERE setting_key IN ('electricity_rate', 'water_rate')
        """)
        settings = cursor.fetchall()
        settings_dict = {s['setting_key']: float(s['setting_value']) for s in settings}
        electricity_rate = settings_dict.get('electricity_rate', 0)
        water_rate = settings_dict.get('water_rate', 0)

        # 5. ดึง previous electricity/water
        cursor.execute("""
            SELECT current_electricity_reading, current_water_reading
            FROM invoices
            WHERE unit_id=%s
            AND contract_id=%s
            AND invoice_type IN ('monthly', 'first')
            AND status NOT IN ('cancelled','void')
            ORDER BY billing_period_start DESC
            LIMIT 1
        """, (contract['room_id'], contract_id))
        last_reading = cursor.fetchone()
        
        if last_reading:
            prev_elec = last_reading['current_electricity_reading'] or 0
            prev_water = last_reading['current_water_reading'] or 0
        else:
            prev_elec = contract['electricity_start'] or 0
            prev_water = contract['water_start'] or 0

        # 6. คำนวณค่าใช้จ่าย (เพิ่มการดึงหนี้เก่า)
        cursor.execute("""
            SELECT invoice_id, total_amount, billing_period_start, billing_period_end 
            FROM invoices 
            WHERE contract_id = %s AND status IN ('unpaid', 'overdue')
        """, (contract_id,))
        unpaid_invoices = cursor.fetchall()
        unpaid_amount = sum(inv['total_amount'] for inv in unpaid_invoices)

        rent = contract['price']
        reimburse = contract['premiums'] or 0
        penalty = 0 

        # ยอดรวม = ค่าเช่าใหม่ + หนี้เก่า + ค่าปรับ - เงินมัดจำ
        total_amount = rent + unpaid_amount + penalty - reimburse

        # 7. สร้าง final invoice
        cursor.execute("""
            INSERT INTO invoices (
                unit_id, tenant_id, contract_id,
                invoice_type, billing_period_start, billing_period_end,
                issue_date, due_date, rent_amount, previous_electricity_reading, electricity_rate,
                previous_water_reading, water_rate, reimburse, total_amount, status, created_by, created_at, premiums
            ) VALUES (%s, %s, %s, 'final', %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'draft', %s, %s, 0)
        """, (
            contract['room_id'], contract['tenant_id'], contract_id,
            start_date, move_out_date,
            today, move_out_date,
            rent, prev_elec, electricity_rate, prev_water, water_rate,
            reimburse, total_amount, session.get('user', {}).get('user_id'), datetime.now()
        ))
        invoice_id = cursor.lastrowid  

        # เพิ่มรายการหนี้เก่าลงใน invoice_items
        for inv in unpaid_invoices:
            desc = f"ยอดค้างชำระจากบิล #{inv['invoice_id']} ({inv['billing_period_start']} ถึง {inv['billing_period_end']})"
            cursor.execute("""
                INSERT INTO invoice_items (invoice_id, description, unit_price, quantity, total_price)
                VALUES (%s, %s, %s, %s, %s)
            """, (invoice_id, desc, inv['total_amount'], 1, inv['total_amount']))

        # 8. เพิ่ม option ลง invoice_items (แก้ไขให้ดึง unit_name มาด้วยป้องกัน Error)
        cursor.execute("""
            SELECT o.id, o.name, o.price
            FROM contract_option co
            JOIN `option` o ON co.option_id = o.id
            WHERE co.contract_id = %s AND o.is_deleted = 0
        """, (contract_id,))
        options = cursor.fetchall()

        for opt in options:
            opt_name = opt['name']
            opt_price = opt['price']

            desc_option = f"{opt_name}" 
            cursor.execute("""
                INSERT INTO invoice_items (invoice_id, description, unit_price, quantity, total_price, option_id, type)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (invoice_id, desc_option, opt_price, 1, opt_price, opt['id'], 'option'))

        # 9. อัปเดตสถานะห้องรอชําระบิลสุดท้าย (status_id=6)
        cursor.execute("""
            UPDATE unit
            SET status_id=6
            WHERE unit_id=%s
        """, (contract['room_id'],))

        # 10. เปลี่ยนสถานะบิลเก่าเป็น 'cancelled' เพื่อไม่ให้ยอดซ้ำซ้อน
        cursor.execute("""
            UPDATE invoices SET status = 'cancelled' 
            WHERE contract_id = %s AND status IN ('unpaid', 'overdue') AND invoice_id != %s
        """, (contract_id, invoice_id))

        # บันทึก Audit Log
        add_audit_log(
            cursor, 
            'TENANT', 
            'CREATE_FINAL_INVOICE', 
            f'สร้างบิลย้ายออกสัญญา ID: {contract_id}, บิล ID: {invoice_id}, ยอดรวม: {total_amount}', 
            session.get('user', {}).get('user_id')
        )

        # ยืนยันรายการทั้งหมด
        conn.commit()
        flash("บันทึกย้ายออกและสร้างบิลสุดท้ายเรียบร้อยแล้ว", "success")

    except Exception as e:
        # หากเกิด Error ใดๆ ให้ยกเลิกสิ่งที่ทำค้างไว้ทั้งหมดทันที
        if conn:
            conn.rollback()
        print(f"❌ Error mark_invoices_overdue: {str(e)}")
        flash(f"เกิดข้อผิดพลาดในการสร้างบิล: {str(e)}", "danger")

    finally:
        # ปิดการเชื่อมต่อเสมอเพื่อป้องกัน Lock ค้าง
        if cursor:
            cursor.close()
        if conn:
            conn.close()

    return redirect(url_for('dashboard'))


# --------------------- UPDATE METER ---------------------
@app.route('/update_meter', methods=['POST'])
def update_meter():
    unit_id = request.form.get('unit_id')
    electricity = request.form.get('electricity', type=float) or 0.0
    water = request.form.get('water', type=float) or 0.0
    
    elec_old_units = request.form.get('elec_old_units', type=float) or 0.0
    water_old_units = request.form.get('water_old_units', type=float) or 0.0

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    try:
        # 1. ดึงข้อมูลบิล Draft และตรวจสอบประเภทบิล
        cursor.execute("""
            SELECT invoice_id, invoice_type, previous_electricity_reading, previous_water_reading 
            FROM invoices 
            WHERE unit_id=%s AND status='draft' 
            ORDER BY created_at DESC LIMIT 1
        """, (unit_id,))
        invoice = cursor.fetchone()
        
        if not invoice: 
            return jsonify({"status": "error", "message": "ไม่พบบิล draft สำหรับห้องนี้"})

        # 2. ดึงเรทน้ำไฟจาก Settings
        cursor.execute("SELECT setting_key, setting_value FROM settings WHERE setting_key IN ('electricity_rate', 'water_rate')")
        rows = cursor.fetchall()
        configs = {row['setting_key']: float(row['setting_value']) for row in rows}
        e_rate = configs.get('electricity_rate', 7.0) 
        w_rate = configs.get('water_rate', 0)

        # 3. กำหนดเลขมิเตอร์ก่อนหน้า (Previous Reading)
        prev_el_db = float(invoice['previous_electricity_reading'] or 0.0)
        prev_wt_db = float(invoice['previous_water_reading'] or 0.0)

        # --- ⚡ จัดการไฟฟ้า (Logic หน่วยค้าง) ---
        if elec_old_units > 0:
            actual_prev_el = 0.0  # มิเตอร์ใหม่เริ่มที่ 0
            cursor.execute("DELETE FROM invoice_items WHERE invoice_id=%s AND type='meter_adjustment' AND description LIKE '%ไฟ%'", (invoice['invoice_id'],))
            cursor.execute("""
                INSERT INTO invoice_items (invoice_id, description, unit_price, quantity, total_price, type)
                VALUES (%s, %s, %s, %s, %s, 'meter_adjustment')
            """, (invoice['invoice_id'], f"หน่วยไฟค้างจากมิเตอร์ตัวเก่าห้อง {unit_id}", e_rate, elec_old_units, elec_old_units * e_rate))
        else:
            actual_prev_el = prev_el_db

        # --- 💧 จัดการน้ำ (Logic หน่วยค้าง) ---
        if water_old_units > 0:
            actual_prev_wt = 0.0 
            cursor.execute("DELETE FROM invoice_items WHERE invoice_id=%s AND type='meter_adjustment' AND description LIKE '%น้ำ%'", (invoice['invoice_id'],))
            cursor.execute("""
                INSERT INTO invoice_items (invoice_id, description, unit_price, quantity, total_price, type)
                VALUES (%s, %s, %s, %s, %s, 'meter_adjustment')
            """, (invoice['invoice_id'], f"หน่วยน้ำค้างจากมิเตอร์ตัวเก่าห้อง {unit_id}", w_rate, water_old_units, water_old_units * w_rate))
        else:
            actual_prev_wt = prev_wt_db

        # 4. [สำคัญ] คำนวณ Usage แยกตามประเภทบิล
        if invoice['invoice_type'] == 'first':
            # เดือนแรก: เซ็ตให้ Usage เป็น 0 (เลขก่อนหน้า = เลขปัจจุบัน) เพื่อตั้งต้นมิเตอร์
            final_prev_el = electricity
            final_prev_wt = water
            elec_usage = 0.0
            water_usage = 0.0
        else:
            # เดือนปกติ หรือ รายวัน หรือ เดือนสุดท้าย: คำนวณตามจริง
            final_prev_el = actual_prev_el
            final_prev_wt = actual_prev_wt
            elec_usage = electricity - final_prev_el
            water_usage = water - final_prev_wt

        # 5. Validation: ป้องกันเลขมิเตอร์ถอยหลัง (ยกเว้นเดือนแรกที่เลขเท่ากันได้)
        if invoice['invoice_type'] != 'first':
            if electricity < final_prev_el:
                return jsonify({"status": "error", "message": f"เลขมิเตอร์ไฟ ({electricity}) น้อยกว่าครั้งก่อน ({final_prev_el})"})
            if water < final_prev_wt:
                return jsonify({"status": "error", "message": f"เลขมิเตอร์น้ำ ({water}) น้อยกว่าครั้งก่อน ({final_prev_wt})"})

        # 6. ล้างสถานะบันทึกเก่าใน meter_history
        cursor.execute("UPDATE meter_history SET is_billed = 1 WHERE unit_id = %s AND is_billed = 0", (unit_id,))

        # 7. อัปเดตข้อมูลมิเตอร์ลงในบิล (invoices)
        cursor.execute("""
            UPDATE invoices SET 
                previous_electricity_reading=%s, previous_water_reading=%s,
                current_electricity_reading=%s, current_water_reading=%s,
                electricity_usage=%s, water_usage=%s,
                electricity_rate=%s, water_rate=%s, meter_saved=1
            WHERE invoice_id=%s
        """, (final_prev_el, final_prev_wt, electricity, water, elec_usage, water_usage, e_rate, w_rate, invoice['invoice_id']))

        # 8. บันทึกประวัติมิเตอร์ (meter_reading)
        for m_type, val in [('electricity', electricity), ('water', water)]:
            table = "meter" if m_type == 'electricity' else "meter_water"
            cursor.execute(f"SELECT id, serial_meter FROM {table} WHERE unit_id = %s", (unit_id,))
            meter_detail = cursor.fetchone()

            if meter_detail:
                cursor.execute("""
                    INSERT INTO meter_reading (unit_id, meter_type, meter_id, serial_meter, current_reading, source, invoice_id, created_by, read_date) 
                    VALUES (%s, %s, %s, %s, %s, 'manual', %s, %s, NOW())
                """, (unit_id, m_type, meter_detail['id'], meter_detail['serial_meter'], val, invoice['invoice_id'], session['user']['user_id']))

        # 9. อัปเดตค่าเริ่มต้นห้อง (unit) เพื่อใช้เป็นเลขตั้งต้นในเดือนถัดไป
        cursor.execute("""
            UPDATE unit 
            SET electricity_start = %s, water_start = %s 
            WHERE unit_id = %s
        """, (electricity, water, unit_id))

        # 10. คำนวณยอดรวมสุทธิใหม่ (จะไปเช็ค Logic invoice_type ในนี้ต่อ)
        refresh_invoice_total(cursor, invoice['invoice_id'])

        add_audit_log(
            cursor, 
            'METER', 
            'UPDATE', 
            f'อัปเดตมิเตอร์ห้อง ID: {unit_id} - ไฟ: {electricity} (Usage: {elec_usage}), น้ำ: {water} (Usage: {water_usage})', 
            session.get('user', {}).get('user_id')
        )
        
        conn.commit()
        return jsonify({"status": "success", "message": "บันทึกมิเตอร์เรียบร้อยแล้ว"})

    except Exception as e:
        if conn: conn.rollback()
        return jsonify({"status": "error", "message": f"เกิดข้อผิดพลาด: {str(e)}"})
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

@app.route('/api/get_unit_meter_data/<int:unit_id>')
def get_unit_meter_data(unit_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        # ดึงประวัติการเปลี่ยนมิเตอร์ล่าสุดที่ยังไม่ได้ปิดบิล
        sql = """
            SELECT type, pending_units, final_reading 
            FROM meter_history 
            WHERE unit_id = %s AND is_billed = 0
        """
        cursor.execute(sql, (unit_id,))
        rows = cursor.fetchall()

        res = {
            "status": "success",
            "adjustments": {"electricity": 0, "water": 0},
            "history": {"electricity": "", "water": ""}
        }

        for r in rows:
            m_type = "electricity" if 'elec' in str(r['type']).lower() else "water"
            res["adjustments"][m_type] = float(r['pending_units'] or 0)
            res["history"][m_type] = r['final_reading'] # เลขล่าสุดที่เคยอ่านได้

        return jsonify(res)
    finally:
        cursor.close()
        conn.close()

# ----------------------TOOL METER -------------------
@app.route('/tool_meter')
@role_required(['admin','manager'])
def tool_meter():
    json_path_electric = os.path.join(app.root_path, 'config_meter', 'model.json')
    json_path_water = os.path.join(app.root_path, 'config_meter', 'model_water.json')
    
    # 1. โหลดข้อมูลจาก DB
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT u.name AS room_name, m.serial_meter
        FROM meter m
        LEFT JOIN unit u ON m.unit_id = u.unit_id
        WHERE m.serial_meter IS NOT NULL
    """)
    meters = cursor.fetchall()
    cursor.close()
    conn.close()

    # 2. รายการ Serial Ports
    serial_ports = [p.device for p in serial.tools.list_ports.comports()]

    # 3. โหลดและรวมข้อมูล JSON
    modules_json = {}

    # โหลดไฟล์มิเตอร์ไฟฟ้า (ถ้ามี)
    if os.path.exists(json_path_electric):
        with open(json_path_electric, 'r', encoding='utf-8') as f:
            modules_json.update(json.load(f))

    # โหลดไฟล์มิเตอร์น้ำ (ถ้ามี) แล้วนำมารวมกัน
    if os.path.exists(json_path_water):
        with open(json_path_water, 'r', encoding='utf-8') as f:
            water_data = json.load(f)
            modules_json.update(water_data) # รวมข้อมูลเข้าด้วยกัน

    return render_template("tool_meter.html", 
                           meters=meters, 
                           modules_json=modules_json,  
                           serial_ports=serial_ports)

@app.route('/read_meter', methods=['POST'])
def ajax_read_meter():
    data = request.get_json()
    model_name = data.get("model_name")
    parameters = data.get("parameters")
    connType = data.get("connType")
    function_code = data.get("function_code")
    selected_keys = data.get("keys")
    api_base_url = data.get("api_base_url") 
    api_token = data.get("api_token")
    start_t = data.get("start_t")
    end_t = data.get("end_t")
    register_key = data.get("register_key")
    data_type = data.get("data_type")
    gain = data.get("gain")
    port = data.get("port")
    ip = data.get("ip")
    count = data.get("count")
    serial_ports = data.get("serial_ports")
    address= data.get("address")
    unit_id=data.get("unit_id")
    baudrate=data.get("baudrate")

    value = read_meter_tool(model_name=model_name, register_key=register_key, port=port, ip=ip, 
                            serial_ports=serial_ports, address=address, unit_id=unit_id, 
                            api_base_url=api_base_url,api_token=api_token,
                            baudrate=baudrate, function_code=function_code, parameters=parameters, 
                            connType=connType, count=count, gain=gain,data_type=data_type,
                            requested_keys=selected_keys,start_t=start_t,end_t=end_t)

    if isinstance(value, (int, float)):
        return jsonify({"success": True, "value": value})

    elif isinstance(value, dict):
        return jsonify({"success": True, "value": value})
    
    elif isinstance(value, list):
        return jsonify({"success": True, "value": value})
    
    elif isinstance(value, str):
        return jsonify({"success": False, "message": value})
    
    else:
        return jsonify({"success": False, "message": "อุปกรณ์ไม่ตอบสนอง (Timeout)"})

@app.route('/write_meter', methods=['POST'])
def ajax_write_meter():
    data = request.get_json()
    model_name = data.get("model_name")
    parameters = data.get("parameters")
    connType = data.get("connType")
    function_code = data.get("function_code")
    register_key = data.get("register_key")
    data_type = data.get("data_type")
    gain = data.get("gain")
    word_swap = data.get("word_swap")
    port = data.get("port")
    ip = data.get("ip")
    count = data.get("count")
    serial_ports = data.get("serial_ports")
    address= data.get("address")
    unit_id=data.get("unit_id")
    baudrate=data.get("baudrate")
    data_dec=data.get("data_dec")

    value = write_meter_tool(model_name, register_key=register_key, port=port, ip=ip, serial_ports=serial_ports, address=address, unit_id=unit_id, baudrate=baudrate
                        , data_dec=data_dec, data_type=data_type, gain=gain, parameters=parameters, function_code=function_code
                        , word_swap=word_swap, count=count, connType=connType)

    # ถ้าผลลัพธ์เป็น Dictionary และมี success=True หรือเป็นคำว่า "success"
    if (isinstance(value, dict) and value.get("success")) or str(value).lower() == "success":
        return jsonify({"success": True, "message": "เขียนข้อมูลสำเร็จ"})
    
    # กรณีล้มเหลว
    msg = value.get("message") if isinstance(value, dict) else str(value)
    return jsonify({"success": False, "message": msg or "ไม่สามารถเขียนข้อมูลได้"})

@app.route('/meter_config_manager')
def meter_config_manager():
    elec_data = read_config(PATH_MODEL_ELEC)
    water_data = read_config(PATH_MODEL_WATER)
    return render_template('meter_config_manager.html', elec=elec_data, water=water_data)

@app.route('/api/manage_config', methods=['POST'])
def manage_config():
    try:
        data = request.json
        action = data.get('action')
        target = data.get('target')  # 'elec' หรือ 'water'
        model_name = data.get('model_name')
        config = data.get('config')

        # 1. กำหนดชื่อไฟล์ตามโครงสร้างที่คุณต้องการ
        config_dir = 'config_meter'
        if target == 'elec':
            filename = os.path.join(config_dir, "model.json")
        else:
            filename = os.path.join(config_dir, "model_water.json")

        # ตรวจสอบว่ามีโฟลเดอร์หรือยัง ถ้าไม่มีให้สร้าง
        if not os.path.exists(config_dir):
            os.makedirs(config_dir)

        # 2. อ่านไฟล์เดิม
        current_data = {}
        if os.path.exists(filename):
            with open(filename, 'r', encoding='utf-8') as f:
                try:
                    current_data = json.load(f)
                except json.JSONDecodeError:
                    current_data = {}

        # 3. จัดการข้อมูล (Save / Delete)
        if action == 'save':
            current_data[model_name] = config
            message = f"บันทึกรุ่น {model_name} ลงใน {os.path.basename(filename)} เรียบร้อย"
        elif action == 'delete':
            if model_name in current_data:
                del current_data[model_name]
                message = f"ลบรุ่น {model_name} ออกแล้ว"
            else:
                return jsonify({"status": "error", "message": "ไม่พบข้อมูล"}), 404

        # 4. เขียนไฟล์กลับ
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(current_data, f, indent=4, ensure_ascii=False)

        return jsonify({"status": "success", "message": message})

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/download_manual')
@role_required(['admin', 'manager'])
def download_manual():
    directory = os.path.join(current_app.root_path, 'static/manual_doc')
    filename = 'คู่มือ.docx'
    
    return send_from_directory(
        directory, 
        filename, 
        as_attachment=True  # บังคับให้ Browser ดาวน์โหลดแทนการพยายามเปิดอ่าน
    )

# ---------------------- RUN APP ----------------------
if __name__ == '__main__':
    try:
        app.run(host='0.0.0.0', port=5000, debug=False)
    except (KeyboardInterrupt, SystemExit):
        scheduler.shutdown()
   